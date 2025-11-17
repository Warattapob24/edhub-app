# app/teacher/routes.py
from collections import Counter, defaultdict
from datetime import date, datetime, timedelta
from tempfile import template
from typing import Optional
from flask import abort, current_app, flash, json, jsonify, redirect, render_template, request, send_file, url_for, render_template_string
from flask_login import current_user, login_required
import pandas as pd
from sqlalchemy import and_, or_
from wtforms import IntegerField, StringField, SubmitField, TextAreaField
from wtforms.validators import DataRequired, Length, Optional
from app.auth.decorators import initial_setup_required
from sqlalchemy.orm import joinedload, selectinload, aliased, contains_eager
from app import db
from sqlalchemy import func
# Ensure all necessary models are imported
from app.models import (AcademicYear, AttendanceRecord, AssessmentDimension, AssessmentItem, AssessmentTemplate, AssessmentTopic,
                        AttendanceWarning, AuditLog, Classroom, CourseGrade, Enrollment, GradedItem, Indicator, LearningStrand,
                        LessonPlanConstraint, PostTeachingLog, Room, RubricLevel, Score, Semester, Course, LearningUnit,
                        LessonPlan, Setting, Standard, Student, StudentGroup, SubUnit, SubjectGroup, TimetableEntry, User,
                        Subject, QualitativeScore, GroupScore, WeeklyScheduleSlot, Notification)
from app.teacher.forms import LearningUnitForm
from app.teacher import bp
from flask_wtf import FlaskForm
# Ensure all necessary services are imported
from app.services import (calculate_final_grades_for_course, check_and_create_attendance_warnings,
                          get_lesson_plan_export_data, get_pator05_data, resolve_active_attendance_warning,
                          copy_lesson_plan, create_blank_lesson_plan) # Added copy_lesson_plan and create_blank_lesson_plan
import logging
import docx
import numpy as np
import io # For handling in-memory files
from weasyprint import HTML, LOGGER # The PDF generation library
from docx import Document # Add this import for Word export
from docx.shared import Pt, Inches # Add this for setting font size in Word
import google.oauth2.credentials
import googleapiclient.discovery
from google.auth.transport.requests import Request as GoogleRequest
import jwt # 👈 สำหรับสร้าง Token ที่ปลอดภัย
import uuid # 👈 สำหรับสร้าง ID ที่ไม่ซ้ำกัน

@bp.route('/dashboard')
@login_required
def dashboard():
    if not current_user.has_role('Teacher'):
        abort(403)
    
    semester = Semester.query.filter_by(is_current=True).first_or_404()
    today = date(2025, 9, 5) 
    # today = date.today()  # <-- 1. กำหนดค่าให้ today ก่อน
    today_weekday = today.isoweekday() # <-- 2. จากนั้นจึงนำ today ไปใช้งาน

    thai_day = ["จันทร์", "อังคาร", "พุธ", "พฤหัสบดี", "ศุกร์", "เสาร์", "อาทิตย์"]
    thai_month = ["มกราคม", "กุมภาพันธ์", "มีนาคม", "เมษายน", "พฤษภาคม", "มิถุนายน", "กรกฎาคม", "สิงหาคม", "กันยายน", "ตุลาคม", "พฤศจิกายน", "ธันวาคม"]

    # today.weekday() เริ่มจาก 0 = จันทร์
    day_str = thai_day[today.weekday()] 
    # today.month เริ่มจาก 1
    month_str = thai_month[today.month - 1] 
    # แปลงเป็น พ.ศ.
    year_str = today.year + 543 

    today_formatted = f"วัน{day_str}ที่ {today.day} {month_str} พ.ศ. {year_str}"

    entries = TimetableEntry.query.join(
        WeeklyScheduleSlot
    ).join(
        Course
    ).filter(
        WeeklyScheduleSlot.semester_id == semester.id,
        WeeklyScheduleSlot.day_of_week == today_weekday,
        Course.teachers.any(id=current_user.id)
    ).options(
        joinedload(TimetableEntry.slot),
        joinedload(TimetableEntry.course).selectinload(Course.subject),
        joinedload(TimetableEntry.course).selectinload(Course.classroom),
        joinedload(TimetableEntry.course).selectinload(Course.room)
    ).order_by(WeeklyScheduleSlot.period_number).all()
    # Aliases for self-join
    Slot1 = aliased(WeeklyScheduleSlot)
    Slot2 = aliased(WeeklyScheduleSlot)
    Entry1 = aliased(TimetableEntry)
    Entry2 = aliased(TimetableEntry)

    # Base query for today's entries
    entries_query = db.session.query(
        Entry1, 
        Entry2.id # Get the ID of the next entry if it exists
    ).join(
        Slot1, Entry1.weekly_schedule_slot_id == Slot1.id
    ).join(
        Course, Entry1.course_id == Course.id
    ).outerjoin( # Use outerjoin to find the potential next entry
        Entry2, 
        and_(
            Entry1.course_id == Entry2.course_id, # Must be the same course
            Entry2.weekly_schedule_slot_id == db.session.query(Slot2.id).filter(
                Slot1.semester_id == Slot2.semester_id,
                Slot1.day_of_week == Slot2.day_of_week,
                Slot2.period_number == Slot1.period_number + 1, # Next period
                Slot1.grade_level_id == Slot2.grade_level_id
            ).scalar_subquery() # Find the slot ID for the next period
        )
    ).filter(
        Slot1.semester_id == semester.id,
        Slot1.day_of_week == today_weekday,
        Course.teachers.any(id=current_user.id)
    ).options(
        # Eager load necessary data for Entry1
        joinedload(Entry1.slot),
        joinedload(Entry1.course).selectinload(Course.subject),
        joinedload(Entry1.course).selectinload(Course.classroom),
        joinedload(Entry1.course).selectinload(Course.room)
    ).order_by(Slot1.period_number)

    # Process results: Create a list of dictionaries
    entries_with_next = []
    for entry, next_id in entries_query.all():
        entries_with_next.append({
            'entry': entry,
            'next_entry_id': next_id # Will be None if no consecutive entry
        })
    return render_template('teacher/dashboard.html', 
                           title="ห้องเรียนวันนี้",
                           entries=entries,
                           entries_with_next=entries_with_next,
                           today_formatted=today_formatted,
                           today=today)

@bp.route('/lesson-plans')
@login_required
# @initial_setup_required
def lesson_plans():
    current_semester = Semester.query.filter_by(is_current=True).options(
        joinedload(Semester.academic_year)
    ).first()

    grouped_plans = defaultdict(lambda: {'plan': None, 'courses': []})

    teacher_courses = [] # Initialize teacher_courses outside the if block

    if current_semester:
        # --- START REVISED QUERY ---
        teacher_courses = Course.query.options(
            joinedload(Course.lesson_plan).joinedload(LessonPlan.subject).joinedload(Subject.subject_group),
            joinedload(Course.classroom)
        ).filter(
            Course.semester_id == current_semester.id,
            Course.teachers.any(id=current_user.id)
        ).all()

        for course in teacher_courses:
            if course.lesson_plan:
                plan_id = course.lesson_plan.id
                if not grouped_plans[plan_id]['plan']:
                    grouped_plans[plan_id]['plan'] = course.lesson_plan
                grouped_plans[plan_id]['courses'].append(course)
        # --- END REVISED QUERY ---

    # +++++++++++++ START DEBUGGING +++++++++++++
    print("-" * 30)
    print(f"[DEBUG] Teacher Dashboard - User ID: {current_user.id}")
    print(f"[DEBUG] Current Semester ID: {current_semester.id if current_semester else 'None'}")
    print(f"[DEBUG] Found {len(teacher_courses)} courses assigned to teacher:")
    # Print details of each course found
    for c in teacher_courses:
         print(f"  - Course ID: {c.id}, Subject: {c.subject.subject_code if c.subject else 'N/A'}, Classroom: {c.classroom.name if c.classroom else 'N/A'}, Plan ID: {c.lesson_plan_id}")

    print("[DEBUG] Grouped Plans Structure sent to template:")
    for plan_id, data in grouped_plans.items():
        subject_name = data['plan'].subject.name if data.get('plan') and data['plan'].subject else "N/A"
        # Print only the classroom names for clarity
        classroom_names = sorted([c.classroom.name for c in data.get('courses', []) if c.classroom])
        print(f"  - Plan ID {plan_id} ({subject_name}): Classrooms = {classroom_names}")
    print("-" * 30)
    # +++++++++++++ END DEBUGGING +++++++++++++

    return render_template('teacher/lesson_plans.html',
                            title='แผนการสอน',
                            grouped_plans=grouped_plans,
                            current_semester=current_semester
                        )

@bp.route('/plan/<int:plan_id>/workspace')
@login_required
def workspace(plan_id):
    plan = LessonPlan.query.options(
        joinedload(LessonPlan.subject)
    ).get_or_404(plan_id)

    if not any(current_user in c.teachers for c in plan.courses):
        abort(403)

    units = sorted(plan.learning_units, key=lambda unit: unit.sequence)
    form = FlaskForm()
    
    # ดึงข้อมูลเงื่อนไข (constraints) ที่มีอยู่
    constraints_obj = LessonPlanConstraint.query.filter_by(lesson_plan_id=plan.id).all()
    # แปลงเป็น Dictionary เพื่อให้ Template ใช้งานง่าย
    constraints = {c.constraint_type: c.value for c in constraints_obj}

    # ใช้ set comprehension เพื่อดึง classroom ที่ไม่ซ้ำกัน
    unique_classrooms = {course.classroom for course in plan.courses if course.classroom}
    # เรียงลำดับตามชื่อห้องเรียน
    classrooms = sorted(list(unique_classrooms), key=lambda cr: cr.name)

    return render_template(
        'teacher/course_workspace.html',
        title=f"จัดการแผนการสอน: {plan.subject.name}",
        plan=plan,
        units=units,
        form=form,
        classrooms=classrooms,
        constraints=constraints
    )

class LearningUnitForm(FlaskForm):
    title = StringField('ชื่อหน่วยการเรียนรู้', validators=[DataRequired(), Length(max=255)])
    topic = StringField('หัวข้อเรื่อง', validators=[Optional(), Length(max=255)])
    hours = IntegerField('จำนวนชั่วโมง', validators=[Optional()])
    learning_objectives = TextAreaField('จุดประสงค์การเรียนรู้', validators=[Optional()])
    core_concepts = TextAreaField('สาระสำคัญ/ความคิดรวบยอด', validators=[Optional()])
    activities = TextAreaField('กิจกรรมการเรียนรู้', validators=[Optional()])
    media_sources = TextAreaField('สื่อและแหล่งเรียนรู้', validators=[Optional()])
    submit = SubmitField('บันทึกแผนการสอน')

# --- แก้ไข API สำหรับสร้างหน่วยการเรียนรู้ ---
@bp.route('/api/plan/<int:plan_id>/units', methods=['POST'])
@login_required
def create_learning_unit(plan_id):
    plan = LessonPlan.query.get_or_404(plan_id)

    # ตรวจสอบสิทธิ์ว่าครูสอนในแผนนี้จริงหรือไม่
    if not any(current_user in c.teachers for c in plan.courses):
        return jsonify({'status': 'error', 'message': 'Unauthorized'}), 403

    data = request.get_json()
    title = data.get('title')

    if not title:
        return jsonify({'status': 'error', 'message': 'Missing title'}), 400

    max_seq = db.session.query(db.func.max(LearningUnit.sequence))\
                .filter_by(lesson_plan_id=plan.id).scalar()
    next_seq = (max_seq or 0) + 1

    new_unit = LearningUnit(lesson_plan_id=plan.id, title=title, sequence=next_seq)
    db.session.add(new_unit)
    db.session.commit()

    return jsonify({
        'status': 'success',
        'unit': {
            'id': new_unit.id,
            'title': new_unit.title,
            'sequence': new_unit.sequence
        }
    })

@bp.route('/api/units/<int:unit_id>', methods=['DELETE'])
@login_required
def delete_learning_unit(unit_id):
    """API endpoint to delete a learning unit."""
    unit = LearningUnit.query.get_or_404(unit_id)

    # --- Permission Check (สำคัญมาก) ---
    # ตรวจสอบว่าผู้ใช้ปัจจุบันเป็นครูที่สอนในรายวิชานี้จริงหรือไม่
    is_teacher_of_course = any(current_user in course.teachers for course in unit.lesson_plan.courses)
    if not is_teacher_of_course:
        abort(403) # Forbidden

    try:
        db.session.delete(unit)
        db.session.commit()
        return jsonify({'status': 'success', 'message': 'ลบหน่วยการเรียนรู้เรียบร้อยแล้ว'})
    except Exception as e:
        db.session.rollback()
        return jsonify({'status': 'error', 'message': str(e)}), 500

@bp.route('/api/units/<int:unit_id>/plan')
@login_required
def get_lesson_plan(unit_id):
    unit = LearningUnit.query.options(
        joinedload(LearningUnit.indicators).joinedload(Indicator.standard)
    ).get_or_404(unit_id)
    
    form = LearningUnitForm(obj=unit) 

    # --- สร้างข้อมูล Indicators เริ่มต้นในรูปแบบ JSON ---
    initial_indicators = []
    for ind in unit.indicators:
        initial_indicators.append({
            "id": ind.id,
            "text": f"[{ind.standard.code} {ind.code}] {ind.description}",
            "indicator_code": ind.code,
            "indicator_desc": ind.description,
            "standard_id": ind.standard.id,
            "standard_code": ind.standard.code,
            "standard_desc": ind.standard.description
        })
    initial_indicators_json = json.dumps(initial_indicators)
    
    return render_template(
        'teacher/_lesson_plan_tab.html', 
        unit=unit, 
        form=form,
        initial_indicators_json=initial_indicators_json # ส่ง JSON String ไปให้ Template
    )

@bp.route('/api/units/<int:unit_id>/plan/save', methods=['POST'])
@login_required
def save_unit_plan(unit_id):
    unit = LearningUnit.query.get_or_404(unit_id)
    plan = unit.lesson_plan
    
    # Get all unique teachers associated with this lesson plan through its courses
    plan_teachers = {teacher for course in plan.courses for teacher in course.teachers}
    
    # Check if the current logged-in user is in that set of teachers
    if current_user not in plan_teachers:
        return jsonify({'status': 'error', 'message': 'Permission denied'}), 403

    data = request.get_json()
    if not data:
        return jsonify({'status': 'error', 'message': 'Invalid data'}), 400

    # Update basic text fields
    unit.title = data.get('title', unit.title)
    unit.topic = data.get('topic')
    unit.hours = data.get('hours')
    unit.learning_objectives = data.get('learning_objectives')
    unit.core_concepts = data.get('core_concepts')
    unit.activities = data.get('activities')
    unit.learning_content = data.get('learning_content')       # <-- ฟิลด์ที่ขาดหายไป
    unit.learning_activities = data.get('learning_activities') # <-- แก้ไขชื่อ
    unit.media_sources = data.get('media_resources')           # <-- แก้ไขชื่อ    
    
    # Update indicators relationship
    indicator_ids = data.get('indicators', [])
    if indicator_ids:
        # Fetch the actual Indicator objects from the database
        indicators = Indicator.query.filter(Indicator.id.in_(indicator_ids)).all()
        # Assign the list of objects directly to the relationship
        unit.indicators = indicators
    else:
        # If the list is empty, assign an empty list to clear the relationship
        unit.indicators = []

    try:
        db.session.commit()
        return jsonify({
            'status': 'success', 
            'message': 'บันทึกแผนการสอนเรียบร้อยแล้ว',
            'new_title': unit.title
        })
    except Exception as e:
        db.session.rollback()
        current_app.logger.error(f"Error saving unit plan for unit {unit_id}: {e}")
        return jsonify({'status': 'error', 'message': 'เกิดข้อผิดพลาดในการบันทึกข้อมูล'}), 500

@bp.route('/api/units/<int:unit_id>/hours', methods=['POST'])
@login_required
def save_unit_hours(unit_id):
    unit = LearningUnit.query.get_or_404(unit_id)
    plan = unit.lesson_plan
    if not any(current_user in c.teachers for c in plan.courses):
        abort(403)

    data = request.get_json()
    try:
        # แปลงค่าที่รับมาและบันทึกลงฐานข้อมูลโดยตรง
        new_hours = int(data.get('hours', 0))
        unit.hours = new_hours
        db.session.commit()
        
        return jsonify({'status': 'success', 'message': 'บันทึกคาบเรียนสำเร็จ'})

    except (ValueError, TypeError):
        db.session.rollback()
        return jsonify({'status': 'error', 'message': 'จำนวนคาบไม่ถูกต้อง'}), 400
    except Exception as e:
        db.session.rollback()
        current_app.logger.error(f"Error saving unit hours for unit {unit_id}: {e}")
        return jsonify({'status': 'error', 'message': 'เกิดข้อผิดพลาดในการบันทึก'}), 500

@bp.route('/api/units/<int:unit_id>/assessment-setup')
@login_required
def get_assessment_setup(unit_id):
    active_unit = LearningUnit.query.get_or_404(unit_id)
    if not any(current_user in c.teachers for c in active_unit.lesson_plan.courses):
        abort(403)

    all_units_in_plan = LearningUnit.query.filter_by(
        lesson_plan_id=active_unit.lesson_plan_id
    ).options(
        joinedload(LearningUnit.graded_items).joinedload(GradedItem.dimension)
    ).order_by(LearningUnit.sequence).all()
    unit_ids = [unit.id for unit in all_units_in_plan]
    
    # --- OPTIMIZATION: Pre-fetch all necessary data in bulk ---
    # 1. Fetch all assessment items for all units in the plan at once.
    all_assessment_items_in_plan = AssessmentItem.query.filter(
        AssessmentItem.learning_unit_id.in_(unit_ids)
    ).all()
    
    # 2. Group items by unit_id for fast lookup without hitting the DB in a loop.
    items_by_unit = defaultdict(list)
    for item in all_assessment_items_in_plan:
        items_by_unit[item.learning_unit_id].append(item)
        
    # 3. Get all unique topic IDs from the pre-fetched items.
    all_topic_ids = {item.assessment_topic_id for item in all_assessment_items_in_plan}
    
    # 4. Fetch all topic objects (and their parents) related to these items at once.
    all_topics_map = {}
    if all_topic_ids:
        all_topics_map = {
            t.id: t for t in AssessmentTopic.query.filter(AssessmentTopic.id.in_(all_topic_ids))
            .options(joinedload(AssessmentTopic.parent)).all()
        }

    all_templates = AssessmentTemplate.query.order_by(AssessmentTemplate.id).all()
    dimensions = AssessmentDimension.query.order_by(AssessmentDimension.id).all()
    form = FlaskForm()

    # A simple helper class to avoid SQLAlchemy session issues when building the tree
    class TopicNode:
        def __init__(self, topic, is_selected):
            self.id, self.name, self.parent_id, self.template_id = \
                topic.id, topic.name, topic.parent_id, topic.template_id
            self.is_selected = is_selected
            self.selected_children = []

    units_assessment_data = {}
    # --- Main loop now uses pre-fetched data, avoiding database queries ---
    for unit in all_units_in_plan:
        structured_topics = defaultdict(list)
        # Use the pre-fetched dictionary instead of a new query
        current_assessment_items = items_by_unit[unit.id]
        selected_topic_ids = {item.assessment_topic_id for item in current_assessment_items}
        
        if selected_topic_ids:
            # Use the pre-fetched map of topic objects
            all_selected_topics = [all_topics_map[tid] for tid in selected_topic_ids if tid in all_topics_map]
            
            tree_nodes = {}
            for topic in all_selected_topics:
                if topic.id not in tree_nodes:
                    tree_nodes[topic.id] = TopicNode(topic, is_selected=True)
                if topic.parent and topic.parent.id not in tree_nodes:
                    tree_nodes[topic.parent.id] = TopicNode(topic.parent, is_selected=False)

            root_nodes = []
            for node in tree_nodes.values():
                if node.parent_id and node.parent_id in tree_nodes:
                    tree_nodes[node.parent_id].selected_children.append(node)
                else:
                    root_nodes.append(node)

            for node in tree_nodes.values():
                node.selected_children.sort(key=lambda x: x.id)
            root_nodes.sort(key=lambda x: x.id)

            for topic in root_nodes:
                structured_topics[topic.template_id].append(topic)
        
        units_assessment_data[unit.id] = {'structured_topics': structured_topics}

    return render_template('teacher/_assessment_setup_tab.html',
                           units=all_units_in_plan,
                           active_unit_id=active_unit.id,
                           templates=all_templates,
                           units_assessment_data=units_assessment_data,
                           dimensions=dimensions,
                           form=form)

@bp.route('/api/units/<int:unit_id>/selected-topics')
@login_required
def get_selected_topics_for_unit(unit_id):
    """
    API endpoint ที่คืนค่าเฉพาะ ID ของหัวข้อที่ถูกเลือกสำหรับหน่วยการเรียนรู้ที่ระบุ
    """
    unit = LearningUnit.query.get_or_404(unit_id)
    if not any(current_user in c.teachers for c in unit.lesson_plan.courses):
        abort(403)
    
    # Get a list of all selected topic IDs for this unit
    selected_ids = [item.assessment_topic_id for item in unit.assessment_items]
    
    # **Crucial Fix**: Return the list wrapped in a dictionary 
    # with the key "selected_ids" that the JavaScript expects.
    return jsonify({'selected_ids': selected_ids})

@bp.route('/api/templates/<int:template_id>/topics-for-selection')
@login_required
def get_topics_for_selection(template_id):
    template = AssessmentTemplate.query.get_or_404(template_id)
    
    # Helper function to format the topic and its children recursively
    def format_topic(topic):
        # เรียงลำดับ children ตาม id หรือชื่อถ้าต้องการ
        children = sorted(topic.children, key=lambda x: x.id)
        return {
            'id': topic.id,
            'name': topic.name,
            'children': [format_topic(child) for child in children]
        }
        
    top_level_topics = AssessmentTopic.query.filter_by(template_id=template.id, parent_id=None).order_by(AssessmentTopic.id).all()
    topic_tree = [format_topic(topic) for topic in top_level_topics]
    
    return jsonify(topic_tree)

@bp.route('/api/units/<int:unit_id>/assessment-items', methods=['POST'])
@login_required
def update_assessment_items(unit_id):
    unit = LearningUnit.query.get_or_404(unit_id)
    if not any(current_user in c.teachers for c in unit.lesson_plan.courses):
        abort(403)

    data = request.get_json()
    if not data:
        return jsonify({'status': 'error', 'message': 'Invalid data'}), 400

    template_id = data.get('template_id')
    selected_ids = data.get('topic_ids', [])

    if template_id is None:
        return jsonify({'status': 'error', 'message': 'Missing template_id'}), 400    

    # 1. ค้นหา ID ของ Topic ทั้งหมดที่อยู่ภายใต้ Template นี้
    #    เราจะใช้ ID เหล่านี้เป็นเงื่อนไขในการลบของเก่า
    topic_ids_in_template = db.session.query(AssessmentTopic.id).filter_by(template_id=template_id).scalar_subquery()

    # 2. สั่งลบ AssessmentItem เก่าทั้งหมดของ "หน่วยการเรียนรู้" นี้
    #    ที่เชื่อมโยงกับ Topic ภายใน Template ที่กำลังแก้ไขอยู่เท่านั้น
    AssessmentItem.query.filter(
        AssessmentItem.learning_unit_id == unit_id,
        AssessmentItem.assessment_topic_id.in_(topic_ids_in_template)
    ).delete(synchronize_session=False)
    
    # 3. สร้าง AssessmentItem ใหม่ตามรายการที่ถูกเลือกส่งมา
    new_items = []
    for topic_id in selected_ids:
        # ตรวจสอบให้แน่ใจว่า topic_id ที่ส่งมาเป็นตัวเลขที่ถูกต้อง
        try:
            item = AssessmentItem(learning_unit_id=unit.id, assessment_topic_id=int(topic_id))
            new_items.append(item)
        except (ValueError, TypeError):
            # ข้าม ID ที่ไม่ถูกต้องไปเงียบๆ หรือจะ log ไว้ก็ได้
            pass
            
    if new_items:
        db.session.bulk_save_objects(new_items)
    
    # 4. บันทึกการเปลี่ยนแปลงทั้งหมดลงฐานข้อมูล
    db.session.commit()

    return jsonify({'status': 'success', 'message': 'บันทึกการตั้งค่าการประเมินเรียบร้อย'})

@bp.route('/search-indicators')
@login_required
def search_indicators():
    """
    API endpoint for TomSelect to search for indicators.
    Results are now sorted by standard and indicator code.
    """
    query_str = request.args.get('q', '', type=str)
    
    if not query_str or len(query_str) < 1:
        return jsonify([])

    search_term = f"%{query_str}%"
    
    indicators = Indicator.query.join(Standard).filter(
        Indicator.creator_type == 'ADMIN',
        or_(
            Indicator.code.ilike(search_term),
            Indicator.description.ilike(search_term),
            Standard.code.ilike(search_term)
        )
    # --- THIS IS THE FIX: Add sorting to the query ---
    ).order_by(Standard.code, Indicator.code).limit(50).all()
    # --- END OF FIX ---

    results = [
        {
            "id": i.id,
            "text": f"[{i.standard.code} {i.code}] {i.description}",
            "indicator_code": i.code,
            "indicator_desc": i.description,
            "standard_id": i.standard.id,
            "standard_code": i.standard.code,
            "standard_desc": i.standard.description
        }
        for i in indicators
    ]
    return jsonify(results)

@bp.route('/api/indicators/add-custom', methods=['POST'])
@login_required
def add_custom_indicator():
    """
    Creates a new custom indicator scoped to the current user and lesson plan.
    """
    data = request.get_json()
    code = data.get('code')
    description = data.get('description')
    plan_id = data.get('plan_id')

    if not all([code, description, plan_id]):
        return jsonify({'status': 'error', 'message': 'Missing required data'}), 400

    # Find or create a dedicated "Standard" for teacher-created indicators
    # This prevents polluting the main standards list
    custom_standard = Standard.query.join(LearningStrand).filter(
        Standard.code == "กำหนดเอง",
        LearningStrand.name == "ผลการเรียนรู้"
    ).first()

    if not custom_standard:
        # Find a generic subject group to attach to, or create one if needed
        # NOTE: This assumes at least one SubjectGroup exists.
        generic_group = SubjectGroup.query.first()
        if not generic_group:
            # This is an edge case, but good to handle
            return jsonify({'status': 'error', 'message': 'No subject groups found to create custom standard'}), 500

        custom_strand = LearningStrand.query.filter_by(name="ผลการเรียนรู้", subject_group_id=generic_group.id).first()
        if not custom_strand:
            custom_strand = LearningStrand(name="ผลการเรียนรู้", subject_group=generic_group)
            db.session.add(custom_strand)
            db.session.flush()

        custom_standard = Standard(code="กำหนดเอง", description="ผลการเรียนรู้", learning_strand=custom_strand)
        db.session.add(custom_standard)
        db.session.flush()

    indicator = Indicator(
        code=code,
        description=description,
        standard_id=custom_standard.id,
        creator_type='TEACHER',
        creator_id=current_user.id,
        lesson_plan_id=plan_id
    )
    db.session.add(indicator)
    db.session.commit()
    
    # Return the new indicator in the same rich format
    return jsonify({
        'status': 'success',
        'indicator': {
            'id': indicator.id,
            'text': f"[{indicator.standard.code} {indicator.code}] {indicator.description}",
            'indicator_code': indicator.code,
            'indicator_desc': indicator.description,
            'standard_id': indicator.standard.id,
            'standard_code': indicator.standard.code,
            'standard_desc': indicator.standard.description
        }
    }), 201

@bp.route('/api/units/<int:unit_id>/graded-items', methods=['POST'])
@login_required
def add_graded_item(unit_id):
    unit = LearningUnit.query.get_or_404(unit_id)
    if not any(current_user in c.teachers for c in unit.lesson_plan.courses):
        abort(403)

    data = request.get_json()
    if not data:
        return jsonify({'status': 'error', 'message': 'Invalid data'}), 400

    # Basic validation
    name = data.get('name')
    max_score = data.get('max_score')
    dimension_id = data.get('assessment_dimension_id')
    if not all([name, max_score, dimension_id]):
        return jsonify({'status': 'error', 'message': 'ข้อมูลไม่ครบถ้วน'}), 400

    new_item = GradedItem(
        name=name,
        max_score=float(max_score),
        indicator_type=data.get('indicator_type', 'FORMATIVE'),
        learning_unit_id=unit_id,
        assessment_dimension_id=int(dimension_id)
    )
    db.session.add(new_item)
    db.session.commit()

    # Return the created item's data for dynamic update on the page
    return jsonify({
        'status': 'success',
        'message': 'บันทึกรายการสำเร็จ',
        'item': {
            'id': new_item.id,
            'name': new_item.name,
            'max_score': new_item.max_score,
            'indicator_type': new_item.indicator_type,
            'dimension_code': new_item.dimension.code
        }
    })

@bp.route('/api/graded-items/<int:item_id>', methods=['DELETE'])
@login_required
def delete_graded_item(item_id):
    item = GradedItem.query.get_or_404(item_id)
    unit = item.learning_unit
    
    # Permission Check
    if not any(current_user in c.teachers for c in unit.lesson_plan.courses):
        abort(403)
        
    try:
        db.session.delete(item)
        db.session.commit()
        return jsonify({'status': 'success', 'message': 'ลบรายการสำเร็จ'})
    except Exception as e:
        db.session.rollback()
        return jsonify({'status': 'error', 'message': str(e)}), 500

@bp.route('/api/graded-items/<int:item_id>', methods=['GET'])
@login_required
def get_graded_item(item_id):
    item = GradedItem.query.get_or_404(item_id)
    # Permission Check
    if not any(current_user in c.teachers for c in item.learning_unit.lesson_plan.courses):
        abort(403)
    
    return jsonify({
        'id': item.id,
        'name': item.name,
        'max_score': item.max_score,
        'assessment_dimension_id': item.assessment_dimension_id,
        'indicator_type': item.indicator_type
    })

@bp.route('/api/graded-items/<int:item_id>', methods=['PUT'])
@login_required
def update_graded_item(item_id):
    item = GradedItem.query.get_or_404(item_id)
    # Permission Check
    if not any(current_user in c.teachers for c in item.learning_unit.lesson_plan.courses):
        abort(403)

    data = request.get_json()
    if not data:
        return jsonify({'status': 'error', 'message': 'Invalid data'}), 400

    # Update fields
    item.name = data.get('name', item.name)
    item.max_score = float(data.get('max_score', item.max_score))
    item.assessment_dimension_id = int(data.get('assessment_dimension_id', item.assessment_dimension_id))
    item.indicator_type = data.get('indicator_type', item.indicator_type)
    
    db.session.commit()

    return jsonify({
        'status': 'success', 
        'message': 'อัปเดตรายการสำเร็จ',
        'item': {
            'id': item.id,
            'name': item.name,
            'max_score': item.max_score,
            'dimension_code': item.dimension.code,
            'indicator_type': item.indicator_type
        }
    })

@bp.route('/api/units/<int:unit_id>/graded-items-for-selection', methods=['GET'])
@login_required
def get_graded_items_for_selection(unit_id):
    """API endpoint to get all graded items for a specific unit for selection."""
    unit = LearningUnit.query.get_or_404(unit_id)
    # Permission Check
    if not any(current_user in c.teachers for c in unit.lesson_plan.courses):
        abort(403)
    
    graded_items = GradedItem.query.filter_by(learning_unit_id=unit.id).order_by(GradedItem.id).all()
    
    results = [
        {
            "id": item.id,
            "name": f"{item.name} ({item.max_score or 0:g} คะแนน)"
        } 
        for item in graded_items
    ]
    return jsonify(results)

@bp.route('/api/units/<int:unit_id>/exam-scores', methods=['POST'])
@login_required
def save_exam_scores(unit_id):
    unit = LearningUnit.query.get_or_404(unit_id)
    # Permission Check
    if not any(current_user in c.teachers for c in unit.lesson_plan.courses):
        abort(403)

    data = request.get_json()
    if data is None:
        return jsonify({'status': 'error', 'message': 'Invalid JSON data'}), 400

    # This new logic handles the payload { "midterm_score": ..., "final_score": ... }
    # .get() will return None if a key is missing, which is the desired behavior.
    unit.midterm_score = data.get('midterm_score')
    unit.final_score   = data.get('final_score')

    try:
        db.session.commit()
        return jsonify({'status': 'success', 'message': 'บันทึกคะแนนสอบเรียบร้อย'})
    except Exception as e:
        db.session.rollback()
        current_app.logger.error(f"Error saving exam scores for unit {unit_id}: {e}")
        return jsonify({'status': 'error', 'message': 'เกิดข้อผิดพลาดในการบันทึก'}), 500

@bp.route('/api/enrollments/save-exam-score-bulk', methods=['POST'])
@login_required
def save_exam_scores_bulk():
    data = request.get_json()
    scores_to_update = data.get('scores', [])
    course_id = data.get('course_id')

    if not scores_to_update or not course_id:
        abort(400)

    student_ids = [s.get('student_id') for s in scores_to_update]
    
    existing_grades = CourseGrade.query.filter(
        CourseGrade.student_id.in_(student_ids),
        CourseGrade.course_id == course_id
    ).all()
    grades_map = {cg.student_id: cg for cg in existing_grades}

    for score_data in scores_to_update:
        student_id = score_data.get('student_id')
        exam_type = score_data.get('exam_type')
        score_value = score_data.get('score')
        
        grade_obj = grades_map.get(student_id)
        if not grade_obj:
            grade_obj = CourseGrade(student_id=student_id, course_id=course_id)
            db.session.add(grade_obj)
            grades_map[student_id] = grade_obj

        score = float(score_value) if score_value is not None and str(score_value).strip() != '' else None
        if exam_type == 'midterm':
            grade_obj.midterm_score = score
        elif exam_type == 'final':
            grade_obj.final_score = score
    
    try:
        db.session.commit()
        return jsonify({'status': 'success', 'message': 'บันทึกคะแนนสอบสำเร็จ'})
    except Exception as e:
        db.session.rollback()
        current_app.logger.error(f"Bulk exam score save error: {e}")
        return jsonify({'status': 'error', 'message': str(e)}), 500

@bp.route('/api/enrollments/save-exam-score', methods=['POST'])
@login_required
def save_enrollment_exam_score():
    """
    Saves a single exam score (midterm or final) for a student in a specific course.
    This endpoint has been fixed to correctly query the enrollment record and return messages in Thai.
    """
    data = request.get_json()
    student_id = data.get('student_id')
    course_id = data.get('course_id')
    exam_type = data.get('exam_type')  # 'midterm' or 'final'
    score = data.get('score')

    if not all([student_id, course_id, exam_type]):
        return jsonify({'status': 'error', 'message': 'ข้อมูลไม่ครบถ้วน (ต้องการ student_id, course_id, exam_type)'}), 400

    # ค้นหาหรือสร้าง CourseGrade record ที่ถูกต้อง
    course_grade = CourseGrade.query.filter_by(student_id=student_id, course_id=course_id).first()
    if not course_grade:
        course_grade = CourseGrade(student_id=student_id, course_id=course_id)
        db.session.add(course_grade)
        
    # Update score based on the specified exam type
    try:
        # แปลงค่า score ที่รับมาให้เป็น float หรือ None
        score_value = float(score) if score is not None and str(score).strip() != '' else None

        if exam_type == 'midterm':
            course_grade.midterm_score = score_value
        elif exam_type == 'final':
            course_grade.final_score = score_value
        else:
            return jsonify({'status': 'error', 'message': 'ประเภทการสอบไม่ถูกต้อง'}), 400
        
        # คอมมิตการเปลี่ยนแปลงลงฐานข้อมูล
        db.session.commit()
        return jsonify({'status': 'success', 'message': 'บันทึกคะแนนสอบเรียบร้อย'})
    except Exception as e:
        db.session.rollback()
        # current_app.logger.error(f"Error saving exam score for enrollment {enrollment.id}: {e}")
        return jsonify({'status': 'error', 'message': 'เกิดข้อผิดพลาดภายในระบบขณะบันทึกคะแนน'}), 500
    
bp.route('/api/units/<int:unit_id>/view')
@login_required
def get_unified_unit_view(unit_id):
    """
    Endpoint to get the single, unified view for a learning unit.
    It combines data fetching for both lesson plan and assessment setup.
    """
    unit = LearningUnit.query.options(
        joinedload(LearningUnit.indicators).joinedload(Indicator.standard),
        joinedload(LearningUnit.graded_items).joinedload(GradedItem.dimension),
        joinedload(LearningUnit.assessment_items)
    ).get_or_404(unit_id)

    # --- Permission Check ---
    if not any(current_user in c.teachers for c in unit.lesson_plan.courses):
        abort(403)

    # --- Data Fetching (Combined from previous endpoints) ---
    form = FlaskForm() # For CSRF token
    all_templates = AssessmentTemplate.query.order_by(AssessmentTemplate.id).all()
    dimensions = AssessmentDimension.query.order_by(AssessmentDimension.id).all()
    selected_topic_ids = {item.assessment_topic_id for item in unit.assessment_items}

    structured_selected_topics = defaultdict(list)
    if selected_topic_ids:
        selected_topics_q = AssessmentTopic.query.filter(AssessmentTopic.id.in_(selected_topic_ids)).all()
        for topic in selected_topics_q:
            if topic.parent_id is None or topic.parent_id not in selected_topic_ids:
                structured_selected_topics[topic.template_id].append(topic)
    
    return render_template(
        'teacher/_unified_unit_view.html',
        unit=unit,
        form=form,
        all_templates=all_templates,
        dimensions=dimensions,
        selected_topic_ids=list(selected_topic_ids),
        structured_selected_topics=structured_selected_topics
    )

@bp.route("/api/plan/<int:plan_id>/ratio-target", methods=["GET", "POST", "DELETE"])
@login_required
def plan_ratio_target(plan_id):
    plan = LessonPlan.query.get_or_404(plan_id)

    # Permission Check
    if not any(current_user in c.teachers for c in plan.courses):
        abort(403)

    if request.method == "POST":
        data = request.get_json()
        plan.target_mid_ratio = data.get("mid_ratio")
        plan.target_final_ratio = data.get("final_ratio")
        db.session.commit()
        return jsonify({"status": "success", "message": "บันทึกสัดส่วนสำเร็จ"})

    if request.method == "DELETE":
        plan.target_mid_ratio = None
        plan.target_final_ratio = None
        db.session.commit()
        return jsonify({"status": "success", "message": "ล้างค่าเป้าหมายสัดส่วนแล้ว"})

    # GET request
    return jsonify({
        "mid_ratio": plan.target_mid_ratio,
        "final_ratio": plan.target_final_ratio
    })

@bp.route('/api/units/<int:unit_id>/sub_units', methods=['POST'])
@login_required
def create_sub_unit(unit_id):
    unit = LearningUnit.query.get_or_404(unit_id)
    if not any(current_user in c.teachers for c in unit.lesson_plan.courses):
        abort(403)
    
    # คำนวณลำดับชั่วโมงถัดไป
    max_seq = db.session.query(db.func.max(SubUnit.hour_sequence)).filter_by(learning_unit_id=unit.id).scalar()
    next_seq = (max_seq or 0) + 1
    
    sub_unit = SubUnit(
        learning_unit_id=unit.id,
        title=f"แผนการสอนชั่วโมงที่ {next_seq}",
        hour_sequence=next_seq,
        activities="ระบุกิจกรรมการเรียนการสอน"
    )
    db.session.add(sub_unit)
    db.session.commit()
    
    return jsonify({
        'id': sub_unit.id,
        'title': sub_unit.title,
        'hour_sequence': sub_unit.hour_sequence,
        'activities': sub_unit.activities
    }), 201

@bp.route('/api/sub_units/<int:sub_unit_id>', methods=['GET'])
@login_required
def get_sub_unit(sub_unit_id):
    sub_unit = SubUnit.query.get_or_404(sub_unit_id)
    if not any(current_user in c.teachers for c in sub_unit.learning_unit.lesson_plan.courses):
        abort(403)
    
    return jsonify({
        'id': sub_unit.id,
        'title': sub_unit.title,
        'activities': sub_unit.activities,
        'indicator_ids': [ind.id for ind in sub_unit.indicators],
        'graded_item_ids': [gi.id for gi in sub_unit.graded_items],
    })

@bp.route('/api/sub_units/<int:sub_unit_id>', methods=['PUT'])
@login_required
def update_sub_unit(sub_unit_id):
    sub_unit = SubUnit.query.get_or_404(sub_unit_id)
    if not any(current_user in c.teachers for c in sub_unit.learning_unit.lesson_plan.courses):
        abort(403)

    data = request.get_json()
    sub_unit.title = data.get('title', sub_unit.title)
    sub_unit.activities = data.get('activities', sub_unit.activities)
    
    # อัปเดตความสัมพันธ์
    if 'indicator_ids' in data:
        sub_unit.indicators = Indicator.query.filter(Indicator.id.in_(data['indicator_ids'])).all()
    if 'graded_item_ids' in data:
        sub_unit.graded_items = GradedItem.query.filter(GradedItem.id.in_(data['graded_item_ids'])).all()
    
    db.session.commit()
    
    # ซิงค์ข้อมูลสองทางไปยังหน่วยหลัก
    # _sync_parent_unit_relations(sub_unit.learning_unit)
    
    return jsonify({'status': 'success', 'message': 'อัปเดตแผนรายชั่วโมงสำเร็จ'})

@bp.route('/api/sub_units/<int:sub_unit_id>', methods=['DELETE'])
@login_required
def delete_sub_unit(sub_unit_id):
    sub_unit = SubUnit.query.get_or_404(sub_unit_id)
    parent_unit = sub_unit.learning_unit # เก็บหน่วยหลักไว้ก่อนลบ
    if not any(current_user in c.teachers for c in parent_unit.lesson_plan.courses):
        abort(403)
    
    db.session.delete(sub_unit)
    db.session.commit()

    # ซิงค์ข้อมูลสองทางไปยังหน่วยหลักหลังการลบ
    # _sync_parent_unit_relations(parent_unit)
    
    return jsonify({'status': 'success', 'message': 'ลบแผนรายชั่วโมงสำเร็จ'})

# --- Helper _update_student_alerts (Might be moved to services.py later) ---
def _update_student_alerts(enrollment, plan=None): # Added default None for plan
    # ... (code as provided, ensure 'plan' variable is correctly referenced if used) ...
    # NOTE: The provided code for _update_student_alerts uses 'plan.id' but receives 'lesson_plan'
    # This might need correction depending on where it's called from or moved to.
    # For now, assuming it's called with the correct object named 'plan'.
    if not enrollment or not plan: # Adjusted check
        return {}

    student = enrollment.student
    current_alerts = {}

    all_graded_items = GradedItem.query.join(LearningUnit).filter(
        LearningUnit.lesson_plan_id == plan.id # Assuming 'plan' is the LessonPlan object
    ).order_by(LearningUnit.sequence, GradedItem.id).all()

    student_scores = Score.query.filter(
        Score.student_id == student.id,
        Score.graded_item_id.in_([item.id for item in all_graded_items])
    ).all()

    # --- START: NEW LOGIC FOR UNSCORED STUDENTS ---
    if not student_scores:
        current_alerts['กรอกคะแนน'] = 'ยังไม่มีการให้คะแนน'
        if enrollment.alerts != current_alerts:
            enrollment.alerts = current_alerts
        return current_alerts
    # --- END: NEW LOGIC ---

    total_score = sum(s.score or 0 for s in student_scores)
    total_max_score = sum(item.max_score or 0 for item in all_graded_items if item.max_score)

    if total_max_score > 0 and (total_score / total_max_score) < 0.5:
        current_alerts['0'] = f"คะแนนรวม {total_score:.1f}/{total_max_score:.1f} (ไม่ถึง 50%)"

    summative_items = [item for item in all_graded_items if item.indicator_type == 'SUMMATIVE']
    incomplete_summative_items = []

    if summative_items and student_scores:
        scores_map = {s.graded_item_id: s.score for s in student_scores}
        for item in summative_items:
            score = scores_map.get(item.id)
            if score is None: # Check specifically for None (meaning never scored)
                incomplete_summative_items.append(item.name)

    if incomplete_summative_items:
        current_alerts['ร'] = f"งานปลายภาคไม่สมบูรณ์: {', '.join(incomplete_summative_items)}"

    if enrollment.alerts != current_alerts:
        enrollment.alerts = current_alerts

    return current_alerts

@bp.route('/api/course/<int:course_id>/gradebook-data')
@login_required
def get_gradebook_data(course_id):
    """
    [FINAL UNIFIED VERSION] This route now correctly calls the central grading service
    AND fetches all necessary structural data (like qualitative assessments) for the UI.
    """
    course = Course.query.options(joinedload(Course.lesson_plan)).get_or_404(course_id)
    if current_user not in course.teachers:
        abort(403)
    
    classroom_id = request.args.get('classroom_id')
    if not classroom_id:
        return jsonify({'status': 'error', 'message': 'Missing classroom_id parameter'}), 400

    plan = course.lesson_plan
    if not plan:
        return jsonify({'status': 'error', 'message': 'Lesson plan not found for this course'}), 404

    # --- Step 1: Call the SINGLE SOURCE OF TRUTH to get all calculated grade data ---
    # This part remains the same, ensuring grading consistency.
    student_grades_data_from_service, max_scores_info = calculate_final_grades_for_course(course)

    # --- Step 2: Re-format student data for the Gradebook's specific needs ---
    enrollments = Enrollment.query.filter_by(classroom_id=classroom_id).all()
    enrollment_map = {en.student_id: en for en in enrollments}
    
    students_data = []
    for s_data in student_grades_data_from_service:
        student = s_data['student']
        enrollment = enrollment_map.get(student.id)
        if enrollment: # Only include students who are actually in the selected classroom
            students_data.append({
                'id': student.id,
                'roll_number': enrollment.roll_number,
                'student_id': student.student_id,
                'name_prefix': student.name_prefix,
                'first_name': student.first_name,
                'last_name': student.last_name,
                'status': student.status,
                'has_ms_status': s_data['has_ms_status'],
                'midterm_score': s_data['midterm_score'],
                'final_score': s_data['final_score'],
                'alerts': enrollment.alerts or {},
                'has_active_warning': False 
            })
    student_ids = [s['id'] for s in students_data]

    # --- Step 3: Fetch raw UI data that the Gradebook needs ---
    # 3.1 Groups
    all_groups = StudentGroup.query.filter_by(lesson_plan_id=plan.id, course_id=course_id).options(joinedload(StudentGroup.enrollments)).all()
    groups_json = [{'id': g.id, 'name': g.name} for g in all_groups]
    student_group_map = {en.student_id: group.id for group in all_groups for en in group.enrollments}

    # 3.2 Graded Items Structure
    units_data, units_map = [], {}
    all_graded_items = GradedItem.query.join(LearningUnit).filter(
        LearningUnit.lesson_plan_id == plan.id
    ).order_by(LearningUnit.sequence, GradedItem.id).all()

    for item in all_graded_items:
        unit = item.learning_unit
        if unit.id not in units_map:
            units_map[unit.id] = {"unit_id": unit.id, "title": unit.title, "items": []}
            units_data.append(units_map[unit.id])
        units_map[unit.id]["items"].append({
            'id': item.id, 'name': item.name, 'max_score': item.max_score,
            'learning_unit_id': item.learning_unit_id,
            'is_group_assignment': item.is_group_assignment,
            'indicator_type': item.indicator_type,
            'google_form_id': item.google_form_id,
            'google_sheet_id': item.google_sheet_id
        })

    # 3.3 Individual Scores
    scores_data = {}
    if student_ids:
        numeric_scores = Score.query.filter(Score.student_id.in_(student_ids), Score.graded_item_id.isnot(None)).all()
        for s in numeric_scores:
            scores_data[f"{s.student_id}-{s.graded_item_id}"] = {'score': s.score}
        
        qualitative_scores = QualitativeScore.query.filter(QualitativeScore.student_id.in_(student_ids), QualitativeScore.course_id == course_id).all()
        for qs in qualitative_scores:
            scores_data[f"{qs.student_id}-q-{qs.assessment_topic_id}"] = {'score': qs.score_value}
    
    # --- Step 4: [RESTORED] Fetch Qualitative Assessment Structure ---
    qualitative_assessment_data = []
    learning_unit_ids = {unit.id for unit in plan.learning_units}
    selected_items_q = AssessmentItem.query.filter(AssessmentItem.learning_unit_id.in_(learning_unit_ids)).all()
    selected_topic_ids = {item.assessment_topic_id for item in selected_items_q}
    topic_to_unit_map = {item.assessment_topic_id: item.learning_unit_id for item in selected_items_q}

    if selected_topic_ids:
        active_templates = AssessmentTemplate.query.join(
            AssessmentTopic, AssessmentTopic.template_id == AssessmentTemplate.id
        ).filter(
            AssessmentTopic.id.in_(selected_topic_ids)
        ).distinct().order_by(AssessmentTemplate.display_order.asc()).options(
            joinedload(AssessmentTemplate.rubric_levels),
            joinedload(AssessmentTemplate.topics).joinedload(AssessmentTopic.children)
        ).all()

        for template in active_templates:
            template_data = {
                'template_name': template.name, 'template_id': template.id, 'main_topics': [],
                'rubrics': [{'label': r.label, 'value': r.value} for r in sorted(template.rubric_levels, key=lambda x: x.value, reverse=True)]
            }
            main_topics_in_template = [t for t in template.topics if t.parent_id is None]
            
            for main_topic in sorted(main_topics_in_template, key=lambda t: t.id):
                selected_sub_topics = [{'id': st.id, 'name': st.name} for st in main_topic.children if st.id in selected_topic_ids]

                if main_topic.id in selected_topic_ids or selected_sub_topics:
                    unit_id_for_topic = topic_to_unit_map.get(main_topic.id) or (topic_to_unit_map.get(selected_sub_topics[0]['id']) if selected_sub_topics else None)
                    template_data['main_topics'].append({
                        'main_topic_id': main_topic.id, 'main_topic_name': main_topic.name,
                        'learning_unit_id': unit_id_for_topic,
                        'selected_sub_topics': sorted(selected_sub_topics, key=lambda x: x['id'])
                    })
            
            if template_data['main_topics']:
                qualitative_assessment_data.append(template_data)

    # --- Step 5: Assemble and return the complete JSON payload ---
    return jsonify({
        'course_id': course_id,
        'students': students_data,
        'units_data': units_data,
        'scores': scores_data,
        'student_group_map': student_group_map,
        'groups': groups_json,
        'qualitative_assessment_data': qualitative_assessment_data,
        'midterm_google_form_id': course.midterm_google_form_id,
        'final_google_form_id': course.final_google_form_id,
        'is_midterm_enabled': max_scores_info['midterm'] > 0,
        'is_final_enabled': max_scores_info['final'] > 0,
        'grand_max_score': max_scores_info['grand_total'],
        'total_midterm_max_score': max_scores_info['midterm'],
        'total_final_max_score': max_scores_info['final'],
        'summative_item_ids': max_scores_info['summative_item_ids']
    })

@bp.route('/api/score/save', methods=['POST'])
@login_required
def save_score():
    data = request.get_json()
    student_id = data.get('student_id')
    item_id = data.get('item_id')
    score_value = data.get('score')

    if not all([student_id, item_id]):
        return jsonify({'status': 'error', 'message': 'ข้อมูลไม่ครบถ้วน'}), 400

    # ตรวจสอบความเป็นเจ้าของ item ผ่าน graded_item -> learning_unit -> lesson_plan -> course -> teacher
    item = GradedItem.query.get_or_404(item_id)
    # (ในระบบที่ซับซ้อนขึ้นควรมีการตรวจสอบสิทธิ์ที่รัดกุมกว่านี้)

    try:
        # แปลง score_value เป็น float ถ้าไม่เป็นค่าว่าง
        score_float = float(score_value) if score_value not in [None, ''] else None
    except (ValueError, TypeError):
        return jsonify({'status': 'error', 'message': 'รูปแบบคะแนนไม่ถูกต้อง'}), 400

    # Logic: Update if exists, Insert if not (Upsert)
    score_obj = Score.query.filter_by(student_id=student_id, graded_item_id=item_id).first()

    if score_obj:
        score_obj.score = score_float
    else:
        score_obj = Score(
            student_id=student_id,
            graded_item_id=item_id,
            score=score_float
        )
        db.session.add(score_obj)

    try:
        db.session.commit()
        return jsonify({'status': 'success', 'message': 'บันทึกคะแนนเรียบร้อย'})
    except Exception as e:
        db.session.rollback()
        return jsonify({'status': 'error', 'message': str(e)}), 500

@bp.route('/api/scores/save-bulk', methods=['POST'])
@login_required
def save_scores_bulk():
    data = request.get_json()
    scores_to_update = data.get('scores', [])
    course_id = data.get('course_id')

    if not all([scores_to_update, course_id]):
        print("--- SAVE FAILED: Missing scores or course_id ---")
        abort(400, "Missing scores or course_id")

    student_ids = list(set([s['student_id'] for s in scores_to_update]))
    item_ids = list(set([s['graded_item_id'] for s in scores_to_update]))

    print(f"--- Processing scores for Course ID: {course_id} ---")
    print(f"--- Student IDs: {student_ids} ---")
    print(f"--- Item IDs: {item_ids} ---")

    try:
        existing_scores = Score.query.filter(
            Score.student_id.in_(student_ids),
            Score.graded_item_id.in_(item_ids)
        ).all()
        existing_map = {f"{s.student_id}-{s.graded_item_id}": s for s in existing_scores}
        print(f"--- Found {len(existing_scores)} existing scores ---")

        objects_to_commit = []
        audit_logs_to_commit = [] # Separate list for logs

        for score_data in scores_to_update:
            student_id = score_data['student_id']
            item_id = score_data['graded_item_id']
            score_value = score_data.get('score')
            score_float = float(score_value) if score_value not in [None, ''] else None
            
            key = f"{student_id}-{item_id}"
            old_score_obj = existing_map.get(key)
            old_score_text = "ยังไม่มีคะแนน" # Default for audit

            if old_score_obj:
                if old_score_obj.score != score_float:
                    print(f"--- Preparing UPDATE for {key}: {old_score_obj.score} -> {score_float} ---")
                    old_score_text = str(old_score_obj.score) if old_score_obj.score is not None else "ยังไม่มีคะแนน" # Get old value BEFORE updating
                    old_score_obj.score = score_float
                    objects_to_commit.append(old_score_obj)

                    # --- RE-ENABLE AuditLog ---
                    audit = AuditLog(
                        user_id=current_user.id, action="Update Score", model_name="Score",
                        record_id=old_score_obj.id, old_value=old_score_text, new_value=str(score_float)
                    )
                    audit_logs_to_commit.append(audit) # Add to separate list
                    # --- END RE-ENABLE ---
                else:
                    print(f"--- No change needed for {key} (Score: {score_float}) ---")
            else:
                print(f"--- Preparing INSERT for {key} (Score: {score_float}) ---")
                new_score = Score(
                    student_id=student_id,
                    graded_item_id=item_id,
                    score=score_float
                )
                objects_to_commit.append(new_score)
                db.session.flush() # Flush to get the ID for the audit log

                # --- RE-ENABLE AuditLog ---
                audit = AuditLog(
                    user_id=current_user.id, action="Create Score", model_name="Score",
                    record_id=new_score.id, old_value=old_score_text, new_value=str(score_float)
                )
                audit_logs_to_commit.append(audit) # Add to separate list
                # --- END RE-ENABLE ---

        if objects_to_commit:
            print(f"--- Attempting to ADD {len(objects_to_commit)} score objects and {len(audit_logs_to_commit)} logs to session... ---")
            db.session.add_all(objects_to_commit + audit_logs_to_commit) # Add both lists
            print("--- Objects added. Attempting COMMIT... ---")
            db.session.commit()
            print("--- COMMIT SUCCEEDED (Python, with AuditLog) ---")
            return jsonify({'status': 'success', 'message': 'บันทึกข้อมูลรวบยอดสำเร็จ'})
        else:
            print("--- No score changes detected. Nothing to commit. ---")
            return jsonify({'status': 'success', 'message': 'ไม่มีข้อมูลคะแนนที่เปลี่ยนแปลง'})
    
    except Exception as e:
        db.session.rollback()
        print(f"--- ERROR DURING SAVE/COMMIT (with AuditLog): {e} ---")
        current_app.logger.error(f"Error in save_scores_bulk (with AuditLog): {e}")
        return jsonify({'status': 'error', 'message': str(e)}), 500

@bp.route('/api/plan/<int:plan_id>/classrooms')
@login_required
def get_plan_classrooms(plan_id):
    plan = LessonPlan.query.get_or_404(plan_id)
    if not any(current_user in c.teachers for c in plan.courses):
        abort(403)
    
    # สร้าง dictionary เพื่อเก็บข้อมูลห้องเรียนและ course_id ที่ไม่ซ้ำกัน
    classroom_data = {}
    for course in plan.courses:
        if course.classroom:
            # เก็บ course_id คู่กับ classroom_id เพื่อป้องกันข้อมูลซ้ำซ้อน
            if course.classroom.id not in classroom_data:
                classroom_data[course.classroom.id] = {
                    'id': course.classroom.id,
                    'name': course.classroom.name,
                    'course_id': course.id # เพิ่ม course_id เข้าไป
                }
    
    # แปลงเป็น list และเรียงลำดับตามชื่อ
    results = sorted(classroom_data.values(), key=lambda cr: cr['name'])
    return jsonify(results)

@bp.route('/api/plan/<int:plan_id>/gradebook-ui')
@login_required
def get_gradebook_ui(plan_id):
    plan = LessonPlan.query.options(
        joinedload(LessonPlan.courses).joinedload(Course.classroom)
    ).get_or_404(plan_id)
    if not any(current_user in c.teachers for c in plan.courses):
        abort(403)
    
    # --- START OF FIX ---
    # สร้าง Map เพื่อเก็บข้อมูลห้องเรียนและ course_id ที่ไม่ซ้ำกัน
    classroom_map = {}
    for course in plan.courses:
        # ตรวจสอบว่าคอร์สนี้มีห้องเรียน และครูที่ login อยู่เป็นผู้สอน
        if course.classroom and current_user in course.teachers:
            # ใช้ classroom.id เป็น key เพื่อป้องกันข้อมูลซ้ำ
            if course.classroom.id not in classroom_map:
                classroom_map[course.classroom.id] = {
                    'classroom': course.classroom,
                    'course_id': course.id
                }
    
    # แปลง Map เป็น List และเรียงลำดับตามชื่อห้องเรียน
    classrooms_data = sorted(classroom_map.values(), key=lambda x: x['classroom'].name)
    # --- END OF FIX ---
    
    # ส่งตัวแปรที่แก้ไขแล้ว (classrooms_data) เข้าไปใน template
    return render_template('teacher/_gradebook_tab.html', plan=plan, classrooms_data=classrooms_data)

@bp.route('/api/units/<int:unit_id>/reflection-tab')
@login_required
def get_reflection_tab(unit_id):
    """
    Renders the main container template for the reflection tab.
    The actual data will be fetched by a separate API call from the frontend.
    """
    unit = LearningUnit.query.get_or_404(unit_id)
    plan = unit.lesson_plan
    if not any(current_user in c.teachers for c in plan.courses):
        abort(403)
    
    # We only need to pass the unit for context; the rest is loaded via API.
    return render_template('teacher/_reflection_tab.html', unit=unit)

@bp.route('/api/units/<int:unit_id>/performance-dashboard')
@login_required
def get_unit_performance_dashboard(unit_id):
    """
    Calculates and returns the aggregated performance data for all classrooms
    for a specific learning unit.
    """
    unit = LearningUnit.query.get_or_404(unit_id)
    plan = unit.lesson_plan
    if not any(current_user in c.teachers for c in plan.courses):
        abort(403)
    
    # 1. Get all graded items and total max score for this unit
    graded_items_in_unit = GradedItem.query.filter_by(learning_unit_id=unit.id).all()
    item_ids = [item.id for item in graded_items_in_unit]
    unit_max_score = sum(item.max_score for item in graded_items_in_unit if item.max_score) or 0

    # 2. Get all relevant classrooms and enrollments
    classrooms = {c.classroom for c in plan.courses if c.classroom and current_user in c.teachers}
    sorted_classrooms = sorted(list(classrooms), key=lambda cr: cr.name)
    classroom_ids = [cr.id for cr in sorted_classrooms]
    
    enrollments = Enrollment.query.filter(Enrollment.classroom_id.in_(classroom_ids)).all()
    enrollment_map = {e.id: e for e in enrollments}
    student_ids = [e.student_id for e in enrollments]

    # 3. Get all scores for these students on these items
    scores_q = db.session.query(
        Score.student_id,
        func.sum(Score.score).label('total_score')
    ).filter(
        Score.student_id.in_(student_ids),
        Score.graded_item_id.in_(item_ids)
    ).group_by(Score.student_id).subquery()

    # 4. Join everything to calculate stats per classroom
    results = {}
    all_scores_list = []
    for classroom in sorted_classrooms:
        students_in_class = Student.query.join(Enrollment).filter(Enrollment.classroom_id == classroom.id).all()
        student_ids_in_class = [s.id for s in students_in_class]

        if not student_ids_in_class:
            continue

        scores_in_class = db.session.query(scores_q.c.total_score).filter(scores_q.c.student_id.in_(student_ids_in_class)).all()
        
        # Filter out None values that may result from students with no scores
        valid_scores = [s[0] for s in scores_in_class if s[0] is not None]
        all_scores_list.extend(valid_scores)

        total_students = len(student_ids_in_class)
        avg_score = sum(valid_scores) / len(valid_scores) if valid_scores else 0
        
        # Standard Deviation calculation
        mean = avg_score
        variance = sum([((x - mean) ** 2) for x in valid_scores]) / len(valid_scores) if valid_scores else 0
        sd = variance ** 0.5

        # Score Distribution
        dist = {'excellent': 0, 'good': 0, 'fair': 0, 'improve': 0}
        if unit_max_score > 0:
            for score in valid_scores:
                percent = (score / unit_max_score) * 100
                if percent >= 80: dist['excellent'] += 1
                elif percent >= 70: dist['good'] += 1
                elif percent >= 50: dist['fair'] += 1
                else: dist['improve'] += 1
        
        # Incomplete work
        scored_students_count = len(valid_scores)
        incomplete_count = total_students - scored_students_count

        results[classroom.id] = {
            'id': classroom.id,
            'name': classroom.name,
            'student_count': total_students,
            'avg_score': round(avg_score, 2),
            'sd': round(sd, 2),
            'distribution': dist,
            'incomplete': incomplete_count
        }

    # 5. Fetch existing logs for this unit
    overall_log = PostTeachingLog.query.filter_by(learning_unit_id=unit.id, teacher_id=current_user.id, classroom_id=None).first()
    per_room_logs_q = PostTeachingLog.query.filter(
        PostTeachingLog.learning_unit_id == unit.id,
        PostTeachingLog.teacher_id == current_user.id,
        PostTeachingLog.classroom_id.isnot(None)
    ).all()
    per_room_logs = {log.classroom_id: {'log_content': log.log_content} for log in per_room_logs_q}


    return jsonify({
        'dashboard_data': list(results.values()),
        'unit_max_score': unit_max_score,
        'logs': {
            'overall': {
                'log_content': overall_log.log_content if overall_log else '',
                'problems_obstacles': overall_log.problems_obstacles if overall_log else '',
                'solutions': overall_log.solutions if overall_log else ''
            },
            'per_room': per_room_logs
        }
    })

@bp.route('/api/qualitative-scores/save', methods=['POST'])
@login_required
def save_qualitative_score():
    data = request.get_json()
    student_id = data.get('student_id')
    topic_id = data.get('topic_id')
    score_value = data.get('score')
    course_id = data.get('course_id')

    if not all([student_id, topic_id, course_id, score_value is not None]):
        return jsonify({'status': 'error', 'message': 'ข้อมูลไม่ครบถ้วน'}), 400

    try:
        # ใช้โมเดล QualitativeScore ตามที่คุณแนะนำ ซึ่งเหมาะสมกว่า
        qs = QualitativeScore.query.filter_by(
            student_id=student_id,
            assessment_topic_id=topic_id,
            course_id=course_id
        ).first()

        # แปลงค่า score ที่รับมาเป็นตัวเลข หากเป็นไปได้
        if score_value is None or score_value == '':
            # ถ้าค่าที่ส่งมาเป็นค่าว่าง และมีข้อมูลเดิมอยู่ ให้ลบทิ้ง
            if qs:
                db.session.delete(qs)
        else:
            # ถ้ามีค่าส่งมา ให้ทำการบันทึกหรืออัปเดต
            score_val_int = int(float(score_value))
            if not qs:
                # สร้างใหม่เมื่อยังไม่มีข้อมูล
                qs = QualitativeScore(
                    student_id=student_id,
                    assessment_topic_id=topic_id,
                    course_id=course_id,
                    score_value=score_val_int
                )
                db.session.add(qs)
            else:
                # อัปเดตข้อมูลเดิม
                qs.score_value = score_val_int

        db.session.commit()
        return jsonify({'status': 'success', 'message': 'บันทึกคะแนนสำเร็จ'})
        
    except Exception as e:
        db.session.rollback()
        current_app.logger.error(f"Error in save_qualitative_score: {e}")
        return jsonify({'status': 'error', 'message': 'เกิดข้อผิดพลาดในการบันทึกข้อมูลเชิงคุณภาพ'}), 500
    
@bp.route('/api/qualitative-scores/save-bulk', methods=['POST'])
@login_required
def save_qualitative_scores_bulk():
    data = request.get_json()
    scores_to_update = data.get('scores', [])
    course_id = data.get('course_id')

    if not scores_to_update or not course_id:
        abort(400)

    try:
        student_ids = {s.get('student_id') for s in scores_to_update}
        topic_ids = {s.get('topic_id') for s in scores_to_update}

        # ค้นหาข้อมูลที่มีอยู่แล้วทั้งหมดในครั้งเดียว
        existing_scores = QualitativeScore.query.filter(
            QualitativeScore.student_id.in_(student_ids),
            QualitativeScore.assessment_topic_id.in_(topic_ids),
            QualitativeScore.course_id == course_id
        ).all()
        
        # สร้าง map เพื่อให้เข้าถึงง่าย: 'student_id-topic_id' -> score_object
        existing_map = {f"{s.student_id}-{s.assessment_topic_id}": s for s in existing_scores}
        
        for item in scores_to_update:
            key = f"{item.get('student_id')}-{item.get('topic_id')}"
            score_value = item.get('score')
            existing_score_obj = existing_map.get(key)

            if score_value is None or score_value == '':
                # ถ้าค่าใหม่เป็นค่าว่าง และมีข้อมูลเดิมอยู่ ให้ลบ
                if existing_score_obj:
                    db.session.delete(existing_score_obj)
            else:
                # ถ้ามีค่าใหม่ส่งมา
                score_val_int = int(float(score_value))
                if existing_score_obj:
                    # อัปเดตค่าเดิม
                    existing_score_obj.score_value = score_val_int
                else:
                    # สร้างรายการใหม่
                    new_score = QualitativeScore(
                        student_id=item.get('student_id'),
                        assessment_topic_id=item.get('topic_id'),
                        course_id=course_id,
                        score_value=score_val_int
                    )
                    db.session.add(new_score)

        db.session.commit()
        return jsonify({'status': 'success', 'message': 'บันทึกข้อมูลสำเร็จ'})

    except Exception as e:
        db.session.rollback()
        current_app.logger.error(f"Error in save_qualitative_scores_bulk: {e}")
        return jsonify({'status': 'error', 'message': 'เกิดข้อผิดพลาดในการบันทึกข้อมูลแบบ Bulk'}), 500

@bp.route('/api/plan/<int:plan_id>/groups', methods=['GET', 'POST'])
@login_required
def manage_student_groups(plan_id):
    # --- ส่วนที่ 1: การตั้งค่าเริ่มต้น (ทำงานทั้ง GET และ POST) ---
    plan = LessonPlan.query.get_or_404(plan_id)
    # Basic permission check
    if not any(current_user in c.teachers for c in plan.courses):
        abort(403)

    # --- ส่วนที่ 2: จัดการกับ POST Request (สำหรับบันทึกข้อมูล) ---
    if request.method == 'POST':
        data = request.get_json()
        groups_to_save = data.get('groups', [])
        course_id = data.get('course_id')
        if not course_id:
            abort(400, 'Missing course_id')
        
        course = Course.query.get_or_404(course_id)
        classroom_id = course.classroom_id

        # 2. ดึงกลุ่มเก่าทั้งหมดของ "หน่วยการเรียนรู้" นี้ และลบทิ้ง
        # SQLAlchemy จะจัดการความสัมพันธ์ (ลบ foreign key) ให้โดยอัตโนมัติ
        StudentGroup.query.filter_by(lesson_plan_id=plan_id, creator_id=current_user.id).delete()
        db.session.flush() # Apply the deletion

        # 1. ดึงข้อมูลการลงทะเบียนของนักเรียนในห้องนี้ทั้งหมด (ทำครั้งเดียว)
        enrollments_in_class = Enrollment.query.filter_by(classroom_id=classroom_id).all()
        enrollment_map = {e.id: e for e in enrollments_in_class}

        # 3. สร้างกลุ่มใหม่และกำหนดสมาชิกตามข้อมูลที่ส่งมา
        for group_data in groups_to_save:
            group_name = group_data.get('name', '').strip()

            if not group_name:
                continue

            new_group = StudentGroup(
                name=group_name, 
                lesson_plan_id=plan_id, 
                course_id=course_id,      # <-- 2.2 เพิ่ม course_id
                creator_id=current_user.id # <-- 2.3 เพิ่ม creator_id
            )
            member_ids = group_data.get('members', [])
            new_group.enrollments = [enrollment_map.get(int(m_id)) for m_id in member_ids if int(m_id) in enrollment_map]
            db.session.add(new_group)
            
        try:
            db.session.commit()
            return jsonify({'status': 'success', 'message': 'บันทึกการจัดกลุ่มเรียบร้อย'})
        except Exception as e:
            db.session.rollback()
            return jsonify({'status': 'error', 'message': str(e)}), 500

    # เพิ่มการกรอง course_id และ creator_id เพื่อความถูกต้อง
    course_id = request.args.get('course_id')
    if not course_id:
        return jsonify([]) # ถ้าไม่มี course_id ส่งมา ให้ return list ว่าง

    # GET Request
    groups = StudentGroup.query.filter_by(
        lesson_plan_id=plan_id, 
        course_id=course_id,
        creator_id=current_user.id
    ).options(joinedload(StudentGroup.enrollments)).all()

    groups_data = [{'id': g.id, 'name': g.name, 'enrollments': [en.id for en in g.enrollments]} for g in groups]
    return jsonify(groups_data)

@bp.route('/api/enrollments/assign-group', methods=['POST'])
@login_required
def assign_students_to_group():
    data = request.get_json()
    enrollment_ids = data.get('enrollment_ids', [])
    group_id = data.get('group_id') # ค่านี้อาจเป็น null เพื่อนำนักเรียนออกจากกลุ่ม

    # TODO: ควรมีการตรวจสอบสิทธิ์เพิ่มเติมว่าครูมีสิทธิ์แก้ไข enrollment_ids เหล่านี้
    
    # ใช้ .update() เพื่อประสิทธิภาพในการอัปเดตข้อมูลจำนวนมาก
    Enrollment.query.filter(Enrollment.id.in_(enrollment_ids)).update(
        {'student_group_id': group_id}, 
        synchronize_session=False
    )
    db.session.commit()

    return jsonify({'status': 'success', 'message': 'อัปเดตกลุ่มนักเรียนเรียบร้อย'})

@bp.route('/api/classrooms/<int:classroom_id>/enrollments', methods=['GET'])
@login_required
def get_enrollments_for_classroom(classroom_id):
    # (สามารถเพิ่มการตรวจสอบสิทธิ์เพื่อความปลอดภัยในอนาคต)
    # ค้นหาการลงทะเบียนทั้งหมดในห้องเรียนที่ระบุ และเรียงตามเลขที่
    enrollments = Enrollment.query.join(Student).filter(
        Enrollment.classroom_id == classroom_id
    ).order_by(Enrollment.roll_number).all()

    # จัดรูปแบบข้อมูลให้เป็น JSON ที่ Frontend ต้องการ
    enrollments_data = [{
        'id': en.id, # ID ของการลงทะเบียน (Enrollment ID)
        'student_id': en.student.id,
        'roll_number': en.roll_number,
        'name': f"{en.student.name_prefix or ''}{en.student.first_name} {en.student.last_name}"
    } for en in enrollments]
    
    return jsonify(enrollments_data)

@bp.route('/api/group-scores/save', methods=['POST'])
@login_required
def save_group_score():
    data = request.get_json()
    group_id = data.get('student_group_id')
    item_id = data.get('graded_item_id')
    score_value = data.get('score')

    group = StudentGroup.query.get_or_404(group_id)
    # Security check: Ensure the current user is the creator of the group
    if group.creator_id != current_user.id:
        abort(403)

    # Find existing score or create a new one
    group_score = GroupScore.query.filter_by(student_group_id=group_id, graded_item_id=item_id).first()

    if score_value is None or score_value == '':
        if group_score:
            db.session.delete(group_score)
    else:
        if not group_score:
            group_score = GroupScore(student_group_id=group_id, graded_item_id=item_id)
            db.session.add(group_score)
        group_score.score = float(score_value)

    db.session.commit()
    return jsonify({'status': 'success', 'message': 'บันทึกคะแนนกลุ่มสำเร็จ'})

@bp.route('/api/log/unit/<int:unit_id>', methods=['POST'])
@login_required
def manage_post_teaching_log(unit_id):
    """
    Handles saving for BOTH overall and per-room logs.
    The presence of 'classroom_id' in the payload determines the log type.
    """
    # Security check
    unit = LearningUnit.query.get_or_404(unit_id)
    if not any(current_user in c.teachers for c in unit.lesson_plan.courses):
        abort(403)

    data = request.get_json()
    classroom_id = data.get('classroom_id') # This can be None

    # Upsert logic
    log = PostTeachingLog.query.filter_by(
        learning_unit_id=unit_id,
        teacher_id=current_user.id,
        classroom_id=classroom_id
    ).first()

    if not log:
        log = PostTeachingLog(
            learning_unit_id=unit_id,
            teacher_id=current_user.id,
            classroom_id=classroom_id
        )
        db.session.add(log)

    # Populate data from payload
    log.log_content = data.get('log_content', '')
    if classroom_id is None: # Only overall log has these fields
        log.problems_obstacles = data.get('problems_obstacles')
        log.solutions = data.get('solutions')

    db.session.commit()
    return jsonify({'status': 'success', 'message': 'บันทึกสำเร็จ'})

@bp.route('/api/plan/<int:plan_id>/submit-for-review', methods=['POST'])
@login_required
def submit_plan_for_review(plan_id):
    """
    API endpoint for teachers to submit their lesson plan for review.
    """
    plan = LessonPlan.query.get_or_404(plan_id)
    
    # Security check: Ensure the current user is a teacher for this plan.
    if not any(current_user in c.teachers for c in plan.courses):
        return jsonify({'status': 'error', 'message': 'Unauthorized'}), 403

    # Update status to 'Pending Review' only if it's in a valid state to be sent
    if plan.status in ['ฉบับร่าง', 'ต้องการการแก้ไข']:
        plan.status = 'รอการตรวจสอบ'
        db.session.commit()
        return jsonify({'status': 'success', 'message': 'ส่งแผนการสอนเพื่อรับการตรวจสอบเรียบร้อยแล้ว'})
    else:
        return jsonify({'status': 'error', 'message': f'ไม่สามารถส่งแผนได้เนื่องจากสถานะปัจจุบันคือ "{plan.status}"'}), 400

# [START ADDITION] API Endpoint for copying a lesson plan
@bp.route('/api/plan/<int:source_plan_id>/copy', methods=['POST'])
@login_required
def copy_plan_api(source_plan_id):
    """ API endpoint to copy a lesson plan. """
    source_plan = db.session.get(LessonPlan, source_plan_id)

    # Basic permission check: Is the current user a teacher for this plan?
    if not source_plan or not any(current_user in c.teachers for c in source_plan.courses):
        return jsonify({'status': 'error', 'message': 'ไม่พบแผนการสอน หรือไม่มีสิทธิ์คัดลอก'}), 404

    data = request.get_json()
    target_academic_year_id = data.get('target_academic_year_id')

    if not target_academic_year_id:
        return jsonify({'status': 'error', 'message': 'ไม่ได้ระบุปีการศึกษาเป้าหมาย'}), 400

    success, result_data = copy_lesson_plan(
        source_plan_id=source_plan_id,
        target_academic_year_id=target_academic_year_id,
        current_user_id=current_user.id # Pass current user ID for custom indicators etc.
    )

    if success:
        return jsonify({'status': 'success', 'new_plan_id': result_data})
    else:
        # result_data contains the error message here
        return jsonify({'status': 'error', 'message': result_data}), 400 # Use 400 for logical errors like "already exists"
# [END ADDITION]

# [START ADDITION] API Endpoint to fetch academic years for the modal
@bp.route('/api/academic-years') # REMOVE '/teacher' prefix if moved to 'main' blueprint
@login_required
def get_academic_years():
    """ API endpoint to get all academic years for selection. """
    try:
        # Find the current academic year ID
        current_semester = Semester.query.filter_by(is_current=True).first()
        current_year_id = current_semester.academic_year_id if current_semester else None

        years = AcademicYear.query.order_by(AcademicYear.year.desc()).all()

        results = [{
            'id': year.id,
            'year': year.year,
            'is_current': (year.id == current_year_id)
        } for year in years]

        return jsonify(results)
    except Exception as e:
        current_app.logger.error(f"Error fetching academic years: {e}")
        return jsonify({'status': 'error', 'message': str(e)}), 500
# [END ADDITION]

# --- User List API (Needs to be moved to 'admin' or similar) ---
# --- !!! IMPORTANT: MOVE THIS ROUTE TO A MORE APPROPRIATE BLUEPRINT LATER !!! ---
# [START ADDITION] Basic API to get user list (for teacher names)
@bp.route('/api/users/simple-list') # REMOVE '/teacher' prefix if moved
@login_required
# Add role check if needed, e.g., @admin_required or @academic_required
def get_users_simple_list():
     """ API endpoint to get a simple list of users (ID and Full Name). """
     try:
          users = User.query.order_by(User.first_name, User.last_name).all()
          results = [{'id': user.id, 'full_name': user.full_name} for user in users]
          return jsonify(results)
     except Exception as e:
          current_app.logger.error(f"Error fetching simple user list: {e}")
          return jsonify({'status': 'error', 'message': str(e)}), 500
    
@bp.route('/api/plan/<int:plan_id>/constraints', methods=['GET'])
@login_required
def get_plan_constraints(plan_id):
    plan = LessonPlan.query.get_or_404(plan_id)
    
    # Security check: ตรวจสอบสิทธิ์
    if not any(current_user in c.teachers for c in plan.courses):
        abort(403)

    # 1. ดึงข้อมูล constraints ที่บันทึกไว้ (เหมือนเดิม)
    constraints_obj = LessonPlanConstraint.query.filter_by(lesson_plan_id=plan_id).all()
    constraints_dict = {c.constraint_type: c.value for c in constraints_obj}
    
    # 2. เพิ่มข้อมูลบันทึกส่วนตัว (เหมือนเดิม)
    constraints_dict['manual_notes'] = plan.manual_scheduling_notes
    
    # --- [START] ส่วนที่เพิ่มเข้ามาเพื่อแก้ไข BUG ---
    # 3. ค้นหา room_id ที่บันทึกไว้ใน Course 
    #    (ใช้ตรรกะเดียวกับ search_rooms ที่เราแก้ไป)
    course_with_room = Course.query.filter(
        Course.lesson_plan_id == plan_id,
        Course.teachers.any(id=current_user.id),
        Course.room_id.isnot(None)
    ).first()
    
    if course_with_room:
        constraints_dict['room_id'] = course_with_room.room_id
    else:
        constraints_dict['room_id'] = None
    # --- [END] ส่วนที่เพิ่มเข้ามาเพื่อแก้ไข BUG ---
    
    return jsonify(constraints_dict)

@bp.route('/api/plan/<int:plan_id>/constraints', methods=['POST'])
@login_required
def save_plan_constraints(plan_id):
    plan = LessonPlan.query.get_or_404(plan_id)
    # Security Check: ตรวจสอบสิทธิ์การเข้าถึง (ควรมี)
    if not any(current_user in c.teachers for c in plan.courses):
        abort(403)

    data = request.get_json()

    # 1. แยกและบันทึก "บันทึกส่วนตัว" โดยเฉพาะ
    if 'manual_notes' in data:
        plan.manual_scheduling_notes = data.pop('manual_notes', None)

    # 2. วนลูปเพื่อบันทึก Constraints ที่เหลือ
    for key, value in data.items():
        constraint = LessonPlanConstraint.query.filter_by(
            lesson_plan_id=plan_id, 
            constraint_type=key
        ).first()

        if value: # ตรวจสอบว่ามีค่าส่งมาหรือไม่
            if constraint:
                constraint.value = value
            else:
                new_constraint = LessonPlanConstraint(
                    lesson_plan_id=plan_id,
                    constraint_type=key,
                    value=value
                )
                db.session.add(new_constraint)
        elif constraint: # ถ้าค่าที่ส่งมาเป็นค่าว่าง ให้ลบ constraint เดิมทิ้ง
            db.session.delete(constraint)

    db.session.commit()
    return jsonify({'status': 'success', 'message': 'บันทึกเงื่อนไขเรียบร้อยแล้ว'})

@bp.route('/attendance/<int:entry_id>')
@login_required
def check_attendance(entry_id):
    """
    หน้าศูนย์บัญชาการห้องเรียนสำหรับคาบเรียนนั้นๆ
    ประกอบด้วยการเช็คชื่อ, กรอกคะแนน, และประเมินผล
    """
    # 1. รับค่า date จาก URL, ถ้าไม่มีให้ใช้วันปัจจุบัน
    date_str = request.args.get('date')
    if date_str:
        try:
            attendance_date = date.fromisoformat(date_str)
        except ValueError:
            attendance_date = date.today()
    else:
        attendance_date = date.today()    
    print(f"\n--- LOADING DATA for date: {attendance_date} (type: {type(attendance_date)}) ---")

    # 1. ดึงข้อมูลหลักของคาบเรียน
    entry = TimetableEntry.query.options(
        joinedload(TimetableEntry.slot),
        joinedload(TimetableEntry.course).joinedload(Course.subject),
        joinedload(TimetableEntry.course).joinedload(Course.classroom),
        joinedload(TimetableEntry.course).joinedload(Course.lesson_plan),
        joinedload(TimetableEntry.attendance_records) 
    ).get_or_404(entry_id)

    # Security check: ตรวจสอบว่าครูที่ login อยู่เป็นเจ้าของคอร์สนี้จริง
    if current_user not in entry.course.teachers:
        abort(403)

    # 2. สร้าง String วันที่รูปแบบภาษาไทย
    thai_day = ["จันทร์", "อังคาร", "พุธ", "พฤหัสบดี", "ศุกร์", "เสาร์", "อาทิตย์"]
    thai_month = ["มกราคม", "กุมภาพันธ์", "มีนาคม", "เมษายน", "พฤษภาคม", "มิถุนายน", "กรกฎาคม", "สิงหาคม", "กันยายน", "ตุลาคม", "พฤศจิกายน", "ธันวาคม"]
    day_str = thai_day[attendance_date.weekday()]
    month_str = thai_month[attendance_date.month - 1]
    year_str = attendance_date.year + 543
    date_formatted_thai = f"วันที่ {attendance_date.day} {month_str} พ.ศ. {year_str}"

    # 3. ค้นหาคาบเรียนถัดไป (สำหรับปุ่มคัดลอก)
    next_entry_id = None
    next_entry = TimetableEntry.query.join(WeeklyScheduleSlot).filter(
        TimetableEntry.course_id == entry.course.id,
        WeeklyScheduleSlot.day_of_week == entry.slot.day_of_week,
        WeeklyScheduleSlot.period_number == entry.slot.period_number + 1
    ).first()
    if next_entry:
        next_entry_id = next_entry.id

    # 4. ดึงข้อมูลนักเรียนและสถานะการเข้าเรียนสำหรับวันที่ที่ระบุ
    students = Student.query.join(Enrollment).filter(
        Enrollment.classroom_id == entry.course.classroom_id
    ).order_by(Enrollment.roll_number).all()

    records_on_date = AttendanceRecord.query.filter_by(
        timetable_entry_id=entry.id,
        attendance_date=attendance_date
    ).all()

    print(f"Found {len(records_on_date)} records for this date.")
    if records_on_date:
        print(f"Example record found: Student ID {records_on_date[0].student_id}, Status {records_on_date[0].status}")

    # เพื่อให้ Template สามารถเข้าถึงได้ง่ายและรวดเร็ว
    existing_attendance = {record.student_id: record.status for record in records_on_date}

    # 3. ดึงข้อมูลสำหรับ Tab "กรอกคะแนน" และ "ประเมินผล" (โค้ดส่วนนี้เหมือนเดิม)
    learning_units = []
    assessment_templates = []
    suggested_unit_id = None
    lesson_plan = entry.course.lesson_plan

    # ประกาศตัวแปร classrooms_data ไว้ก่อน
    classrooms_data = []

    if lesson_plan:
        # ดึงหน่วยการเรียนรู้ทั้งหมด เรียงตามลำดับ
        learning_units = LearningUnit.query.filter_by(lesson_plan_id=lesson_plan.id).order_by(LearningUnit.sequence).all()

        # [SMART SELECTION LOGIC] คำนวณหาหน่วยการเรียนรู้ที่แนะนำสำหรับคาบนี้
        # 3.1 นับว่าคาบนี้เป็นคาบที่เท่าไหร่ของเทอมสำหรับคอร์สนี้
        class_session_number = db.session.query(func.count(TimetableEntry.id)).join(WeeklyScheduleSlot).filter(
            TimetableEntry.course_id == entry.course.id,
            WeeklyScheduleSlot.semester_id == entry.course.semester_id, # เพิ่มเงื่อนไข semester
            # เรียงลำดับตามวันและคาบ เพื่อให้การนับแม่นยำ
            (WeeklyScheduleSlot.day_of_week, WeeklyScheduleSlot.period_number) <= (entry.slot.day_of_week, entry.slot.period_number)
        ).scalar() or 1

        # 3.2 วน Loop เพื่อหาว่าคาบที่นับได้ ตกอยู่ในหน่วยใด
        total_hours_so_far = 0
        for unit in learning_units:
            if unit.hours and unit.hours > 0:
                if class_session_number <= total_hours_so_far + unit.hours:
                    suggested_unit_id = unit.id
                    break # เมื่อเจอหน่วยที่ใช่แล้ว ให้ออกจาก Loop ทันที
                total_hours_so_far += unit.hours
        
        # ดึงข้อมูลห้องเรียนสำหรับ Dropdown ใน Gradebook (เหมือนเดิม)
        plan_courses = Course.query.filter(
            Course.lesson_plan_id == lesson_plan.id,
            Course.teachers.any(id=current_user.id)
        ).options(joinedload(Course.classroom)).all()
        
        classroom_map = {}
        for course in plan_courses:
            if course.classroom and course.classroom.id not in classroom_map:
                classroom_map[course.classroom.id] = { 'classroom': course.classroom, 'course_id': course.id }
        
        classrooms_data = sorted(classroom_map.values(), key=lambda x: x['classroom'].name)

    return render_template(
        'teacher/attendance.html',
        title="จัดการห้องเรียน",
        entry=entry,
        students=students,
        learning_units=learning_units,
        assessment_templates=assessment_templates,
        suggested_unit_id=suggested_unit_id,
        attendance_date=attendance_date, # ส่ง object date ไปด้วย
        date_formatted_thai=date_formatted_thai, # ส่ง string ภาษาไทยไปแสดงผล
        next_entry_id=next_entry_id, # ส่ง ID ของคาบถัดไป        
        existing_attendance=existing_attendance,
        plan=lesson_plan,                # <-- ADD THIS
        classrooms_data=classrooms_data,      # <-- ADD THIS        
        # now=datetime.now() # เพิ่มตัวแปร now สำหรับใช้ใน Template
    )

@bp.route('/api/attendance/save', methods=['POST'])
@login_required
def save_attendance():
    data = request.get_json()
    print(f"\n--- SAVING DATA ---")
    print(f"Received payload: {data}")    
    entry_id = data.get('entry_id')
    student_id = data.get('student_id')
    status = data.get('status')
    attendance_date_str = data.get('date') 

    if not all([entry_id, student_id, status, attendance_date_str]):
        return jsonify({'status': 'error', 'message': 'ข้อมูลไม่ครบถ้วน'}), 400

    try:
        attendance_date = date.fromisoformat(attendance_date_str)
    except (ValueError, TypeError):
        return jsonify({'status': 'error', 'message': 'รูปแบบวันที่ไม่ถูกต้อง'}), 400

    entry = TimetableEntry.query.get_or_404(entry_id)
    if current_user not in entry.course.teachers:
        abort(403)

    # Logic: Update if exists, Insert if not (Upsert)
    record = AttendanceRecord.query.filter_by(
        timetable_entry_id=entry_id, 
        student_id=student_id,
        attendance_date=attendance_date
    ).first()

    current_time = datetime.utcnow()

    if record:
        record.status = status
        # Update timestamp to reflect the latest change
        record.recorded_at = current_time       
    else:
        record = AttendanceRecord(
            student_id=student_id,
            timetable_entry_id=entry_id,
            status=status,
            recorder_id=current_user.id,
            recorded_at=current_time,
            attendance_date=attendance_date
        )
        record.timetable_entry = entry
        record.student = db.session.get(Student, student_id)
        db.session.add(record)

    try:
        db.session.commit()
        print("--- SAVE SUCCESSFUL (Committed to DB) ---")
        return jsonify({'status': 'success', 'message': 'บันทึกข้อมูลเรียบร้อย'})
    except Exception as e:
        db.session.rollback()
        print(f"--- SAVE FAILED: Rolling back. Error: {e} ---")
        current_app.logger.error(f"Error in save_attendance: {e}")
        return jsonify({'status': 'error', 'message': str(e)}), 500
    
@bp.route('/api/attendance/save-bulk', methods=['POST'])
@login_required
def save_attendance_bulk():
    data = request.get_json()
    entry_id = data.get('entry_id')
    all_students_status = data.get('status') # e.g., 'PRESENT'

    if not entry_id or not all_students_status:
        return jsonify({'status': 'error', 'message': 'Missing data'}), 400

    entry = TimetableEntry.query.get_or_404(entry_id)
    if current_user not in entry.course.teachers:
        abort(403)

    # ดึงรายชื่อนักเรียนทั้งหมดในห้อง
    enrollments = Enrollment.query.filter_by(classroom_id=entry.course.classroom_id).all()
    student_ids = [e.student_id for e in enrollments]

    # ดึงข้อมูลการเข้าเรียนที่มีอยู่แล้ว
    existing_records = {rec.student_id: rec for rec in AttendanceRecord.query.filter(
        AttendanceRecord.timetable_entry_id == entry_id,
        AttendanceRecord.student_id.in_(student_ids)
    ).all()}

    records_to_add = []
    for student_id in student_ids:
        if student_id in existing_records:
            existing_records[student_id].status = all_students_status
        else:
            records_to_add.append(AttendanceRecord(
                student_id=student_id,
                timetable_entry_id=entry_id,
                status=all_students_status,
                recorder_id=current_user.id
            ))

    if records_to_add:
        db.session.bulk_save_objects(records_to_add)

    try:
        db.session.commit()
        return jsonify({'status': 'success', 'message': 'บันทึกข้อมูลรวบยอดสำเร็จ'})
    except Exception as e:
        db.session.rollback()
        return jsonify({'status': 'error', 'message': str(e)}), 500
    
@bp.route('/api/search-rooms')
@login_required
def search_rooms():
    """API for Tom-Select to search for rooms."""
    
    # --- DEBUGGING PRINT ---
    print("\n--- DEBUG: search_rooms (FIXED VERSION) IS RUNNING ---") 
    # --- END DEBUGGING ---

    query = request.args.get('q', '').strip()
    
    # ดึงข้อมูลห้องเรียนที่ถูกเลือกในปัจจุบันสำหรับแผนการสอนนี้ (เพื่อแสดงผลเริ่มต้น)
    plan_id = request.args.get('plan_id', type=int)
    current_room_id = None
    
    print(f"--- DEBUG: Searching for plan_id: {plan_id} ---") # DEBUG
    
    if plan_id:
        course_with_room = Course.query.filter(
            Course.lesson_plan_id == plan_id,
            Course.teachers.any(id=current_user.id),
            Course.room_id.isnot(None) # <-- เพิ่มเงื่อนไข: ต้องมี room_id
        ).first()
        
        if course_with_room:
            current_room_id = course_with_room.room_id
            print(f"--- DEBUG: Found saved room_id: {current_room_id} ---") # DEBUG
        else:
            print("--- DEBUG: No course with a saved room_id was found. ---") # DEBUG

    # สร้าง Query เริ่มต้น
    room_query = Room.query
    
    # ถ้ามีคำค้นหา ให้กรองตามชื่อ
    if query:
        room_query = room_query.filter(Room.name.ilike(f'%{query}%'))

    # จำกัดผลลัพธ์เพื่อประสิทธิภาพ
    rooms = room_query.limit(20).all()
    
    # แปลงผลลัพธ์เป็น JSON ที่ Tom-Select เข้าใจ
    results = [{'id': room.id, 'name': room.name, 'capacity': room.capacity} for room in rooms]
    
    # ถ้ามีห้องที่ถูกเลือกไว้อยู่แล้ว แต่ไม่ติดมาในผลการค้นหา ให้เพิ่มเข้าไปด้วย
    if current_room_id and not any(r['id'] == current_room_id for r in results):
        current_room = Room.query.get(current_room_id)
        if current_room:
            results.insert(0, {'id': current_room.id, 'name': current_room.name, 'capacity': current_room.capacity})
            print(f"--- DEBUG: Manually added current_room {current_room_id} to results. ---") # DEBUG

    print(f"--- DEBUG: Returning {len(results)} results. ---\n") # DEBUG
    return jsonify(results)

@bp.route('/api/rooms/create', methods=['POST'])
@login_required
def create_room_on_the_fly():
    """API for Tom-Select to create a new room."""
    data = request.json
    name = data.get('name', '').strip()

    if not name:
        return jsonify({'status': 'error', 'message': 'ชื่อห้องห้ามว่างเปล่า'}), 400
    if Room.query.filter_by(name=name).first():
        return jsonify({'status': 'error', 'message': 'มีห้องชื่อนี้อยู่แล้ว'}), 409

    new_room = Room(
        name=name, 
        capacity=data.get('capacity'), 
        room_type=data.get('room_type'), # <--- เพิ่มบรรทัดนี้
        notes=data.get('notes')
    )
    db.session.add(new_room)
    db.session.commit()
    
    # ส่งข้อมูลห้องที่สร้างใหม่กลับไปให้ Tom-Select
    return jsonify({
        'status': 'success',
        'room': {'id': new_room.id, 'name': new_room.name, 'capacity': new_room.capacity}
    })

@bp.route('/api/plan/<int:plan_id>/assign-room', methods=['POST'])
@login_required
def assign_room_to_plan_courses(plan_id):
    """Assigns a room to all courses related to this plan for the current user."""
    data = request.json
    room_id = data.get('room_id')

    # ค้นหาทุก Course ที่เกี่ยวกับแผนการสอนนี้และครูคนนี้
    courses_to_update = Course.query.filter(
        Course.lesson_plan_id == plan_id,
        Course.teachers.any(id=current_user.id)
    ).all()

    if not courses_to_update:
        return jsonify({'status': 'error', 'message': 'ไม่พบรายวิชาที่เกี่ยวข้องกับแผนการสอนนี้'}), 404

    for course in courses_to_update:
        course.room_id = room_id if room_id else None

    db.session.commit()
    return jsonify({'status': 'success', 'message': 'บันทึกห้องเรียนเรียบร้อยแล้ว'})

@bp.route('/api/plan/<int:plan_id>/attendance-overview')
@login_required
def get_attendance_overview(plan_id):
    plan = LessonPlan.query.get_or_404(plan_id)
    if not any(current_user in c.teachers for c in plan.courses):
        abort(403)
    
    # ดึงข้อมูลห้องเรียนทั้งหมดที่เกี่ยวกับแผนการสอนนี้และครูคนนี้
    # เพื่อนำไปสร้าง Dropdown เลือกห้อง
    plan_courses = Course.query.filter(
        Course.lesson_plan_id == plan.id,
        Course.teachers.any(id=current_user.id)
    ).options(joinedload(Course.classroom)).all()
    
    # สร้าง list ของห้องเรียนที่ไม่ซ้ำกัน
    classroom_map = {}
    for course in plan_courses:
        if course.classroom and course.classroom.id not in classroom_map:
            classroom_map[course.classroom.id] = {
                'classroom': course.classroom,
                'course_id': course.id
            }
    
    # เรียงลำดับตามชื่อห้อง
    classrooms_data = sorted(classroom_map.values(), key=lambda x: x['classroom'].name)

    return render_template('teacher/_attendance_overview_tab.html',
                           plan=plan,
                           classrooms_data=classrooms_data)    

@bp.route('/api/course/<int:course_id>/attendance-data')
@login_required
def get_attendance_data(course_id):
    classroom_id = request.args.get('classroom_id')
    if not classroom_id:
        abort(400, 'Missing classroom_id parameter')

    course = Course.query.options(
        joinedload(Course.semester) # <-- FIX 1: Eager load semester
    ).get_or_404(course_id)
    if current_user not in course.teachers:
        abort(403)

    # 1. ดึงข้อมูลนักเรียน
    enrollments = Enrollment.query.join(Student).filter(
        Enrollment.classroom_id == classroom_id
    ).order_by(Enrollment.roll_number).all()
    students_data = [{
        'id': en.student.id, 'roll_number': en.roll_number,
        'name': f"{en.student.first_name} {en.student.last_name}"
    } for en in enrollments]
    student_ids = [s['id'] for s in students_data]

    # 2. สร้างรายการคาบเรียนทั้งหมดในเทอม (20 สัปดาห์)
    
    # --- START OF FIX 2: เปลี่ยนการจัดการ Error ---
    if not course.semester or not course.semester.start_date:
        # เปลี่ยนจาก 500 เป็น 400 (Bad Request)
        # นี่คือปัญหาการตั้งค่า ไม่ใช่เซิร์ฟเวอร์ล่ม
        return jsonify({
            'error': 'ยังไม่ได้ตั้งค่าวันเริ่มภาคเรียน',
            'message': 'ไม่สามารถสร้างตารางเวลาได้เนื่องจาก "วันเริ่มต้นภาคเรียน" ยังไม่ได้ถูกตั้งค่าในระบบ'
        }), 400
    
    semester_start_date = course.semester.start_date
    # --- END OF FIX 2 ---

    # --- FIX 3: Eager load 'slot'
    entries_in_course = TimetableEntry.query.options(
        joinedload(TimetableEntry.slot) 
    ).filter_by(
        course_id=course_id
    ).join(
        WeeklyScheduleSlot
    ).order_by(WeeklyScheduleSlot.day_of_week, WeeklyScheduleSlot.period_number).all()

    sessions_data = []
    thai_days = ["จ.", "อ.", "พ.", "พฤ.", "ศ.", "ส.", "อา."]
    for week in range(20):
        for entry in entries_in_course:
            # FIX 4: เพิ่มการตรวจสอบเผื่อข้อมูล slot มีปัญหา
            if not entry.slot:
                continue # ข้าม entry นี้ไปถ้าไม่มีข้อมูล slot

            # Calculate the specific date for this session
            day_offset = (entry.slot.day_of_week - 1) - semester_start_date.isoweekday() + 1
            session_date = semester_start_date + timedelta(days=week * 7 + day_offset)
            sessions_data.append({
                'entry_id': entry.id,
                'date': session_date.isoformat(),
                'week': week + 1,
                'day': thai_days[entry.slot.day_of_week - 1],
                'period': entry.slot.period_number
            })

    # 3. ดึงข้อมูลการเข้าเรียนที่มีอยู่ทั้งหมด
    records = AttendanceRecord.query.filter(
        AttendanceRecord.student_id.in_(student_ids),
        AttendanceRecord.timetable_entry_id.in_([e.id for e in entries_in_course])
    ).all()

    attendance_map = {
        f"{rec.student_id}-{rec.timetable_entry_id}-{rec.attendance_date.isoformat()}": rec.status
        for rec in records
    }

    return jsonify({
        'students': students_data,
        'sessions': sessions_data,
        'attendance_data': attendance_map
    })

@bp.route('/api/attendance/copy', methods=['POST'])
@login_required
def copy_attendance():
    data = request.get_json()
    target_entry_id = data.get('target_entry_id')
    attendance_date_str = data.get('date')
    source_data = data.get('source_data') # <-- รับข้อมูลจาก Frontend โดยตรง

    if not all([target_entry_id, attendance_date_str, source_data]):
        return jsonify({'status': 'error', 'message': 'ข้อมูลไม่ครบถ้วน'}), 400

    try:
        attendance_date = date.fromisoformat(attendance_date_str)
    except ValueError:
        return jsonify({'status': 'error', 'message': 'รูปแบบวันที่ไม่ถูกต้อง'}), 400

    # ดึงข้อมูลปลายทางที่มีอยู่แล้ว (ถ้ามี) เพื่อทำการ Upsert
    target_records_map = {
        rec.student_id: rec for rec in AttendanceRecord.query.filter_by(
            timetable_entry_id=target_entry_id,
            attendance_date=attendance_date
        ).all()
    }

    for item in source_data:
        student_id = item.get('student_id')
        status = item.get('status')
        
        t_rec = target_records_map.get(student_id)
        if t_rec: # ถ้ามี record ปลายทางอยู่แล้ว -> Update
            t_rec.status = status
            t_rec.recorded_at = datetime.utcnow()
        else: # ถ้ายังไม่มี -> Insert
            new_rec = AttendanceRecord(
                student_id=student_id,
                timetable_entry_id=target_entry_id,
                status=status,
                recorder_id=current_user.id,
                attendance_date=attendance_date
            )
            db.session.add(new_rec)
    
    db.session.commit()
    return jsonify({'status': 'success', 'message': 'คัดลอกข้อมูลการเข้าเรียนเรียบร้อย'})

@bp.route('/grade-submission')
@login_required
def submit_grades_overview():
    if not current_user.has_role('Teacher'):
        abort(403)
    
    current_semester = Semester.query.filter_by(is_current=True).first_or_404()
    
    teacher_courses = Course.query.filter(
        Course.teachers.any(id=current_user.id),
        Course.semester_id == current_semester.id
    ).join(Course.subject).join(Course.classroom).options(
        joinedload(Course.subject),
        joinedload(Course.classroom)
    ).order_by(
        Subject.subject_code.asc(),
        Classroom.name.asc(),
        db.case(
            (Course.grade_submission_status == 'ยังไม่ส่ง', 0),
            (Course.grade_submission_status == 'ต้องการการแก้ไข', 1),
            else_=2
        )
    ).all()

    return render_template('teacher/submit_grades_overview.html',
                           title="ภาพรวมการส่งผลการเรียน",
                           courses=teacher_courses,
                           semester=current_semester)

@bp.route('/course/<int:course_id>/submit-grades', methods=['GET', 'POST'])
@login_required
def submit_course_grades(course_id):
    course = Course.query.options(
        joinedload(Course.subject),
        joinedload(Course.classroom)
    ).get_or_404(course_id)
    form = FlaskForm()

    if current_user not in course.teachers:
        abort(403)
    
    student_grades, max_scores = calculate_final_grades_for_course(course)

    if form.validate_on_submit():
        if course.grade_submission_status not in ['ยังไม่ส่ง', 'ต้องการการแก้ไข', 'pending']:
            flash('ไม่สามารถส่งผลการเรียนได้เนื่องจากสถานะปัจจุบันไม่ถูกต้อง', 'danger')
            return redirect(url_for('teacher.submit_course_grades', course_id=course_id))

        for data in student_grades:
            grade_obj = data['exam_grade_obj']
            if not grade_obj:
                grade_obj = CourseGrade(student_id=data['student'].id, course_id=course.id)
                db.session.add(grade_obj)
            grade_obj.midterm_score = data['midterm_score']
            grade_obj.final_score = data['final_score']                
            grade_obj.final_grade = data['grade']
        
        course.grade_submission_status = 'รอตรวจสอบ (หน.กลุ่มสาระ)'
        course.submitted_by_id = current_user.id
        course.submitted_at = datetime.utcnow()
        db.session.commit()
        try:
            # --- [START] Add Audit Log ---
            log_details = {
                'semester': f"{course.semester.term}/{course.semester.academic_year.year}",
                'subject': f"{course.subject.subject_code} - {course.subject.name}",
                'classroom': course.classroom.name,
                'new_status': course.grade_submission_status # Log the status being set
                # Optional: Add summary like number of students submitted
            }
            audit = AuditLog(
                user_id=current_user.id,
                action_type="Submit Grades", # Use consistent action types
                model_name="Course",
                record_id=course.id,
                details=json.dumps(log_details, ensure_ascii=False) # Store details as JSON
            )
            db.session.add(audit)
            # --- [END] Add Audit Log ---

            db.session.commit() # Commit both grade updates and audit log

            flash(f'ส่งผลการเรียนวิชา {course.subject.name} ห้อง {course.classroom.name} เรียบร้อยแล้ว', 'success')
            return redirect(url_for('teacher.submit_grades_overview'))
        except Exception as e:
            db.session.rollback()
            current_app.logger.error(f"Error submitting grades or logging for course {course_id}: {e}")
            flash('เกิดข้อผิดพลาดในการบันทึกข้อมูล', 'danger')

    return render_template('teacher/submit_course_grades.html',
                           title=f"ยืนยันผลการเรียน: {course.subject.name}",
                           course=course,
                           student_grades=student_grades,
                           max_collected=max_scores['collected'],
                           max_midterm=max_scores['midterm'],
                           max_final=max_scores['final'],
                           grand_max_score=max_scores['grand_total'],
                           form=form)

@bp.route('/subject-summary/select')
@login_required
def subject_summary_selection():
    """ หน้าสำหรับให้ครูเลือกรายวิชาที่จะดูสรุป """
    current_semester = Semester.query.filter_by(is_current=True).first_or_404()
    
    # ค้นหารายวิชา (Subject) ที่ไม่ซ้ำกัน ที่ครูคนนี้สอนในเทอมปัจจุบัน
    subjects = Subject.query.join(Course).filter(
        Course.semester_id == current_semester.id,
        Course.teachers.any(id=current_user.id)
    ).distinct().order_by(Subject.subject_code).all()

    return render_template('teacher/subject_summary_selection.html',
                           title="เลือกรายวิชาเพื่อดูสรุปผล",
                           subjects=subjects,
                           semester=current_semester)

@bp.route('/subject-summary/<int:subject_id>/semester/<int:semester_id>')
@login_required
def view_subject_summary(subject_id, semester_id):
    subject = Subject.query.get_or_404(subject_id)
    semester = Semester.query.get_or_404(semester_id)

    courses = Course.query.filter_by(subject_id=subject.id, semester_id=semester.id).options(
        joinedload(Course.classroom),
        joinedload(Course.teachers)
    ).all()

    if not courses:
        flash('ไม่พบข้อมูลรายวิชาสำหรับภาคเรียนนี้', 'warning')
        return redirect(url_for('teacher.subject_summary_selection'))

    # --- [ยกเครื่อง] ใช้ Service ใหม่ในการคำนวณข้อมูลทั้งหมด ---
    all_student_final_data = []
    grand_max_score = 0
    for course in courses:
        student_grades, max_scores = calculate_final_grades_for_course(course)
        all_student_final_data.extend(student_grades)
        if max_scores['grand_total'] > grand_max_score:
             grand_max_score = max_scores['grand_total']

    # --- 4. จัดกลุ่มข้อมูลและประมวลผลสถิติ (เหมือนเดิมแต่ใช้ข้อมูลใหม่) ---
    summary_data = {
        'overall': {'grades': [], 'scores': []},
        'by_classroom': defaultdict(lambda: {'grades': [], 'scores': []})
    }
    for data in all_student_final_data:
        summary_data['overall']['grades'].append(data['grade'])
        summary_data['overall']['scores'].append(data['total_score'])
        summary_data['by_classroom'][data['classroom_id']]['grades'].append(data['grade'])
        summary_data['by_classroom'][data['classroom_id']]['scores'].append(data['total_score'])

    def process_stats(grade_list, score_list):
        if not grade_list: return {}
        stats = {}
        total_students = len(grade_list)
        grade_counts = Counter(grade_list)
        stats['grade_distribution'] = {str(g): grade_counts.get(str(g), 0) for g in ['4','3.5','3','2.5','2','1.5','1','0','ร','มส']}
        stats['grade_percentages'] = {g: (count / total_students * 100) if total_students > 0 else 0 for g, count in stats['grade_distribution'].items()}
        
        valid_grades_for_gpa = [float(g) for g in grade_list if g not in ['ร', 'มส', None]]
        stats['gpa'] = np.mean(valid_grades_for_gpa) if valid_grades_for_gpa else 0
        stats['sd'] = np.std(valid_grades_for_gpa, ddof=1) if len(valid_grades_for_gpa) > 1 else 0
        
        valid_scores = score_list or [0]
        stats['min_score'] = np.min(valid_scores)
        stats['max_score'] = np.max(valid_scores)
        stats['median_score'] = np.median(valid_scores)

        stats['passed_count'] = len([g for g in valid_grades_for_gpa if g >= 1])
        stats['good_excellent_count'] = len([g for g in valid_grades_for_gpa if g >= 3])
        stats['failed_count'] = total_students - stats['passed_count']
        stats['total_students'] = total_students
        stats['passed_percent'] = (stats['passed_count'] / total_students * 100) if total_students > 0 else 0
        stats['good_excellent_percent'] = (stats['good_excellent_count'] / total_students * 100) if total_students > 0 else 0
        stats['failed_percent'] = (stats['failed_count'] / total_students * 100) if total_students > 0 else 0
        return stats

    summary_data['overall']['stats'] = process_stats(summary_data['overall']['grades'], summary_data['overall']['scores'])
    for cid, data in summary_data['by_classroom'].items():
        classroom_obj = db.session.get(Classroom, cid)
        summary_data['by_classroom'][cid]['name'] = classroom_obj.name
        summary_data['by_classroom'][cid]['stats'] = process_stats(data['grades'], data['scores'])

    chart_labels = ['4','3.5','3','2.5','2','1.5','1','0','ร','มส']
    chart_data = [summary_data['overall']['stats'].get('grade_distribution', {}).get(label, 0) for label in chart_labels]
    summary_data['chart_data'] = {'labels': chart_labels, 'data': chart_data}
    
    all_teachers = list(set(teacher for course in courses for teacher in course.teachers))

    return render_template('teacher/subject_summary.html',
                           title=f"สรุปผลรายวิชา: {subject.name}",
                           subject=subject,
                           semester=semester,
                           summary_data=summary_data,
                           all_teachers=all_teachers,
                           courses_count=len(courses),
                           grand_max_score=grand_max_score)

@bp.route('/remediation')
@login_required
def remediation_courses():
    semester = Semester.query.filter_by(is_current=True).first_or_404()

    # --- REVISED QUERY V2 ---
    grades_in_process_q = db.session.query(CourseGrade, Enrollment).join(
        Course, CourseGrade.course_id == Course.id
    ).join(
        Enrollment, and_(
            CourseGrade.student_id == Enrollment.student_id,
            Course.classroom_id == Enrollment.classroom_id # Ensure enrollment matches course classroom
        )
    ).filter(
        Course.teachers.any(id=current_user.id),
        Course.semester_id == semester.id,
        # --- NEW FILTER LOGIC V2 ---
        or_(
            # Case 1: Currently failing grade AND no remediation started yet
            and_(
                CourseGrade.final_grade.in_(['0', 'ร', 'มส']),
                # Ensure original_final_grade is ALSO None OR matches final_grade here
                # to avoid picking up already remediated students who might temporarily have a failing grade saved
                or_(CourseGrade.original_final_grade.is_(None), CourseGrade.original_final_grade == CourseGrade.final_grade),
                CourseGrade.remediation_status == 'None' # Status before starting remediation
            ),
            # Case 2: Originally failed AND remediation is in progress/completed/submitted, but not yet fully approved
            and_(
                 CourseGrade.original_final_grade.in_(['0', 'ร', 'มส']),
                 CourseGrade.remediation_status != 'Approved', # Exclude fully approved cases
                 CourseGrade.remediation_status != 'None' # Exclude cases captured by Case 1
            )
        )
        # --- END NEW FILTER LOGIC V2 ---
    ).options(
        joinedload(CourseGrade.student),
        joinedload(CourseGrade.course).joinedload(Course.subject),
        joinedload(CourseGrade.course).joinedload(Course.classroom)
    ).order_by(
         Course.subject_id, Course.classroom_id, Enrollment.roll_number
    ).all()
    # --- END REVISED QUERY V2 ---

    # --- REVISED GROUPING LOGIC V2 ---
    awaiting_ms = []  # Grade is 'มส', Status is 'None'
    awaiting_r = []   # Grade is 'ร', Status is 'None'
    awaiting_zero = []# Grade is '0', Status is 'None'
    in_progress = []  # Status 'In Progress'
    completed = []    # Status 'Completed'
    submitted = []    # Status starts with 'Submitted' or 'Pending Director Approval'

    for grade_obj, enrollment_obj in grades_in_process_q:
        item_tuple = (grade_obj, enrollment_obj)
        status = grade_obj.remediation_status
        # Use current grade for initial status 'None' because original_final_grade might not be set yet
        current_grade = grade_obj.final_grade
        # Use original grade for grouping other statuses if available, fallback to current
        original_grade = grade_obj.original_final_grade or current_grade

        if status.startswith('Submitted') or status == 'Pending Director Approval':
             submitted.append(item_tuple)
        elif status == 'Completed':
             completed.append(item_tuple)
        elif status == 'In Progress':
             in_progress.append(item_tuple)
        elif status == 'None': # Should only be records caught by Case 1 of the query
             if current_grade == 'มส':
                 awaiting_ms.append(item_tuple)
             elif current_grade == 'ร':
                 awaiting_r.append(item_tuple)
             elif current_grade == '0':
                 awaiting_zero.append(item_tuple)
        # 'Approved' are excluded by the query

    # --- RECALCULATE STATS V2 ---
    total_awaiting = len(awaiting_ms) + len(awaiting_r) + len(awaiting_zero)
    total_in_progress = len(in_progress)
    total_completed = len(completed)
    total_submitted = len(submitted)
    total_all = total_awaiting + total_in_progress + total_completed + total_submitted

    stats = {
        'awaiting_action': total_awaiting,
        'in_progress': total_in_progress,
        'completed': total_completed,
        'submitted': total_submitted,
        'total': total_all, # Total currently visible/actionable by teacher
        'percent_awaiting': (total_awaiting / total_all * 100) if total_all > 0 else 0,
        'percent_in_progress': (total_in_progress / total_all * 100) if total_all > 0 else 0,
        'percent_completed': (total_completed / total_all * 100) if total_all > 0 else 0,
        'percent_submitted': (total_submitted / total_all * 100) if total_all > 0 else 0,
    }

    # Debug print (optional, can be removed after verification)
    print("-" * 20)
    print(f"Total records found by query V2: {len(grades_in_process_q)}")
    print(f"Awaiting MS: {len(awaiting_ms)}")
    print(f"Awaiting R: {len(awaiting_r)}")
    print(f"Awaiting 0: {len(awaiting_zero)}")
    print(f"In Progress: {len(in_progress)}")
    print(f"Completed: {len(completed)}")
    print(f"Submitted: {len(submitted)}")
    print("-" * 20)


    return render_template('teacher/remediation_courses.html',
                           title="ศูนย์บัญชาการสอนซ่อม",
                           semester=semester,
                           stats=stats,
                           # Pass the CORRECTED lists to the template
                           ms_students_original=awaiting_ms,  # Renamed variable in template context
                           r_students_original=awaiting_r,   # Renamed variable in template context
                           zero_students_original=awaiting_zero,# Renamed variable in template context
                           in_progress_students=in_progress, # New list for template
                           completed_students=completed,
                           submitted_students=submitted)

@bp.route('/remediation/course/<int:course_id>')
@login_required
def remediation_details(course_id):
    """ Displays the list of students needing remediation for a specific course. """
    course = Course.query.options(
        joinedload(Course.subject),
        joinedload(Course.classroom)
    ).get_or_404(course_id)

    if current_user not in course.teachers:
        abort(403)

    # 1. Calculate ALL student grades for the course first to get the complete data structure.
    all_student_grades, max_scores = calculate_final_grades_for_course(course)

    # 2. Now, filter this complete list to find only the students who need remediation.
    students_to_remediate = [
        student_data for student_data in all_student_grades
        if student_data['grade'] in ['0', 'ร', 'มส'] or (student_data.get('exam_grade_obj') and student_data['exam_grade_obj'].remediation_status != 'None')
    ]

    return render_template('teacher/remediation_details.html',
                           title=f"จัดการนักเรียนที่ต้องซ่อม: {course.subject.name}",
                           course=course,
                           students=students_to_remediate,
                           max_scores=max_scores)

@bp.route('/api/remediation/course/<int:course_id>/student/<int:student_id>')
@login_required
def get_remediation_details_api(course_id, student_id):
    """ API endpoint to fetch detailed data for a student's remediation using live calculation. """
    course = db.session.get(Course, course_id)
    student = db.session.get(Student, student_id)
    if not course or not student or current_user not in course.teachers:
        abort(404)

    all_student_grades, max_scores = calculate_final_grades_for_course(course)
    student_data = next((s for s in all_student_grades if s['student'].id == student_id), None)
    
    if not student_data:
        return jsonify({'error': 'Student data not found in this course calculation.'}), 404
        
    live_final_grade = student_data.get('grade')
    original_course_grade_obj = student_data.get('exam_grade_obj')

    response_data = {
        'student': {
            'id': student.id,
            'full_name': student_data.get('full_name')
        },
        'course_grade': {
            'final_grade': live_final_grade,
            # --- NEW: Pass the original saved grade for context ---
            'original_final_grade': original_course_grade_obj.original_final_grade if original_course_grade_obj else live_final_grade,
            'collected_score': student_data.get('collected_score', 0),
            'midterm_score': student_data.get('midterm_score'),
            'final_score': student_data.get('final_score'),
            'total_score': student_data.get('total_score'),
            
            # --- ADDED FIELD as per instructions ---
            'midterm_remediated_score': original_course_grade_obj.midterm_remediated_score if original_course_grade_obj else None,
            # --- END ADDED FIELD ---
        },
        'max_scores': max_scores,
        'remediation_data': {}
    }

    if live_final_grade == 'มส':
        # ... (Existing 'ms' logic) ...
        # --- REVISED: Use data directly from the service ---
        total_periods = student_data.get('total_periods', 0)
        absent_count = student_data.get('absent_count', 0)
        response_data['remediation_data']['ms'] = {
            'total_periods': total_periods,
            'absent_count': absent_count,
            'absence_percentage': (absent_count / total_periods * 100) if total_periods > 0 else 0
        }

    elif live_final_grade in ['ร', '0']:
        # ... (Existing 'r'/'0' logic) ...
        incomplete_items = []
        all_scores_q = Score.query.join(GradedItem).join(LearningUnit).filter(
            LearningUnit.lesson_plan_id == course.lesson_plan_id,
            Score.student_id == student_id
        )
        scores_map = {s.graded_item_id: s.score for s in all_scores_q}
        
        all_items = GradedItem.query.join(LearningUnit).options(
            joinedload(GradedItem.dimension)
        ).filter(LearningUnit.lesson_plan_id == course.lesson_plan.id).all()

        for item in all_items:
            incomplete_items.append({
                'id': item.id, 
                'name': item.name, 
                'max_score': item.max_score, 
                'score': scores_map.get(item.id),
                'dimension_code': item.dimension.code if item.dimension else None,
                'indicator_type': item.indicator_type
            })
            
        if max_scores['midterm'] > 0:
            incomplete_items.append({'id': 'exam_midterm', 'name': 'สอบกลางภาค', 'max_score': max_scores['midterm'], 'score': student_data.get('midterm_score')})
        if max_scores['final'] > 0:
            incomplete_items.append({'id': 'exam_final', 'name': 'สอบปลายภาค', 'max_score': max_scores['final'], 'score': student_data.get('final_score')})

        if incomplete_items:
            response_data['remediation_data']['scores'] = incomplete_items

    return jsonify(response_data)


@bp.route('/api/remediation/save', methods=['POST'])
@login_required
def save_remediation_data():
    """ [REVISED] API endpoint to save remediation scores with robust "Upsert" logic. """
    try:
        data = request.get_json()
        if not data:
            return jsonify({'status': 'error', 'message': 'Invalid request data'}), 400

        course_id = data.get('course_id')
        student_id = data.get('student_id')
        scores_payload = data.get('scores', [])
        ms_remediated = data.get('ms_remediated', False)
        
        # --- ADDED FIELD as per instructions ---
        midterm_remediated_score_str = data.get('midterm_remediated_score')
        # --- END ADDED FIELD ---

        if not all([course_id, student_id]):
            return jsonify({'status': 'error', 'message': 'Missing course or student ID'}), 400

        course = db.session.get(Course, course_id)
        if not course:
            return jsonify({'status': 'error', 'message': 'Course not found'}), 404
        if current_user not in course.teachers:
            return jsonify({'status': 'error', 'message': 'Permission denied'}), 403

        # --- THE "UPSERT" FIX IS HERE ---
        # Find the record. If it doesn't exist, create it in memory.
        course_grade = CourseGrade.query.filter_by(course_id=course_id, student_id=student_id).first()
        if not course_grade:
            course_grade = CourseGrade(course_id=course_id, student_id=student_id)
            db.session.add(course_grade)
        # --- END OF FIX ---

        if course_grade.original_final_grade is None:
            # Use the grade from the payload if it's the very first save
            original_grade_from_frontend = data.get('original_grade')
            course_grade.original_final_grade = original_grade_from_frontend or course_grade.final_grade

        if ms_remediated:
            course_grade.ms_remediated_status = True

        for item in scores_payload:
            item_id = item.get('id')
            score_value = item.get('score')
            score = float(score_value) if score_value not in [None, ''] else None

            if str(item_id).startswith('exam_'):
                if item_id == 'exam_midterm':
                    course_grade.midterm_score = score
                elif item_id == 'exam_final':
                    course_grade.final_score = score
            else: # GradedItem score
                score_obj = Score.query.filter_by(student_id=student_id, graded_item_id=item_id).first()
                if score_obj:
                    score_obj.score = score
                else:
                    db.session.add(Score(student_id=student_id, graded_item_id=item_id, score=score))

        # --- NEW BLOCK: Save remediated midterm score (for Export) ---
        try:
            # Convert to float if not empty/None, otherwise set to None
            score_val = float(midterm_remediated_score_str) if midterm_remediated_score_str is not None and midterm_remediated_score_str != '' else None
            # (Validation logic could be added here if needed)
            course_grade.midterm_remediated_score = score_val
        except (ValueError, TypeError):
            current_app.logger.warning(f"Invalid format for midterm_remediated_score: {midterm_remediated_score_str}")
            pass # Silently ignore invalid format
        # --- END NEW BLOCK ---

        # --- Existing logic to recalculate grade based on *performance* scores ---
        # This logic is INTENTIONALLY kept, as it does NOT use the new
        # 'midterm_remediated_score' field in its calculation.
        all_grades, max_scores = calculate_final_grades_for_course(course)
        new_student_data = next((s for s in all_grades if s['student'].id == student_id), None)

        if new_student_data:
            new_grade = new_student_data['grade']
            course_grade.final_grade = new_grade
            
            # --- THE FIX IS HERE: Stricter definition of "Completed" ---
            # Define what a truly passing remediated grade is
            passing_grades = ['1', '1.5', '2', '2.5', '3', '3.5', '4']

            # Stage 1: If this is the first time saving, always set to 'In Progress'.
            if course_grade.remediation_status == 'None':
                course_grade.remediation_status = 'In Progress'
            
            # Stage 2: Only promote to 'Completed' if it's already in progress AND 
            # the new grade is a valid passing grade.
            if course_grade.remediation_status == 'In Progress' and new_grade in passing_grades:
                course_grade.remediation_status = 'Completed'
                course_grade.remediated_at = datetime.utcnow()
        # --- End of existing recalculation logic ---

        db.session.commit()
        return jsonify({'status': 'success', 'message': 'บันทึกผลการแก้ไขเรียบร้อยแล้ว'})

    except Exception as e:
        db.session.rollback()
        current_app.logger.error(f"Critical error in save_remediation_data: {e}")
        return jsonify({'status': 'error', 'message': 'เกิดข้อผิดพลาดร้ายแรงในระบบ: ' + str(e)}), 500

@bp.route('/api/remediation/submit', methods=['POST'])
@login_required
def submit_remediated_grades():
    """ API to submit all 'Completed' remediation records for a course. """
    data = request.get_json()
    course_id = data.get('course_id')
    if not course_id:
        return jsonify({'status': 'error', 'message': 'Missing course_id'}), 400

    course = db.session.get(Course, course_id)
    if not course or current_user not in course.teachers:
        return jsonify({'status': 'error', 'message': 'Permission denied'}), 403

    # Find all students with 'Completed' status for this course and update them
    updated_count = CourseGrade.query.filter_by(
        course_id=course_id,
        remediation_status='Completed'
    ).update({'remediation_status': 'Submitted to Dept. Head'})

    if updated_count > 0:
        db.session.commit()
        try:
            # 1. ค้นหาหัวหน้ากลุ่มสาระของวิชานี้
            subject_group = course.subject.subject_group
            if subject_group and subject_group.head:
                head_user = subject_group.head

                # 2. สร้างข้อความและ URL
                title = "แจ้งเตือนการส่งผลการซ่อม"
                message = f"ครู {current_user.full_name} ได้ส่งผลการซ่อมของนักเรียน {updated_count} คน สำหรับวิชา {course.subject.name} ({course.classroom.name}) เพื่อรอการตรวจสอบ"
                # TODO: ต้องสร้าง URL ปลายทางสำหรับ Head (ถ้ามี)
                # url_for_head = url_for('department.review_remediation', course_id=course.id, _external=True) 

                # 3. สร้าง Notification
                new_notif = Notification(
                    user_id=head_user.id,
                    title=title,
                    message=message,
                    url=None, # ใส่ URL ที่ถูกต้องที่นี่
                    notification_type='REMEDIATION_SUBMIT'
                )
                db.session.add(new_notif)
                db.session.commit()
        except Exception as e:
            current_app.logger.error(f"Failed to send remediation notification for course {course_id}: {e}")
            # ไม่ต้อง rollback เพราะการส่งเกรดหลักสำเร็จไปแล้ว
            pass
    
    return jsonify({
        'status': 'success', 
        'message': f'ส่งผลการซ่อมของนักเรียน {updated_count} คนเรียบร้อยแล้ว'
    })        

@bp.route('/remediation/submit-all', methods=['POST'])
@login_required
def submit_all_remediated_grades():
    semester = Semester.query.filter_by(is_current=True).first_or_404()
    
    updated_count = CourseGrade.query.join(Course).filter(
        Course.teachers.any(id=current_user.id),
        Course.semester_id == semester.id,
        CourseGrade.remediation_status == 'Completed'
    ).update({'remediation_status': 'Submitted to Dept. Head'}, synchronize_session=False)

    if updated_count > 0:
        db.session.commit()
    
    return jsonify({
        'status': 'success',
        'message': f'ส่งผลการซ่อมของนักเรียน {updated_count} คนให้หัวหน้ากลุ่มสาระฯ เรียบร้อยแล้ว'
    })

@bp.route('/course/<int:course_id>/export/pator05')
@login_required
# @teacher_required # Add permission decorator if you have one
def export_pator05(course_id):
    """Generates and returns the Pator05 PDF for the given course."""

    # 1. Get data using the service function
    pator05_data = get_pator05_data(course_id)
    if not pator05_data:
        flash('ไม่พบข้อมูลสำหรับสร้าง ปถ.05 หรือเกิดข้อผิดพลาดในการดึงข้อมูล', 'danger')
        return redirect(url_for('teacher.dashboard'))

    try:
        # --- Render ALL HTML Templates ---
        html_cover = render_template('exports/pator05/cover.html', **pator05_data)
        html_criteria = render_template('exports/pator05/criteria.html', **pator05_data)
        html_score_structure = render_template('exports/pator05/score_structure.html', **pator05_data)
        html_attendance = render_template('exports/pator05/attendance.html', **pator05_data)
        html_scores = render_template('exports/pator05/scores.html', **pator05_data)
        html_summary = render_template('exports/pator05/summary.html', **pator05_data)
        html_back = render_template('exports/pator05/back.html', **pator05_data)

        # --- Create ALL WeasyPrint HTML objects ---
        doc_cover = HTML(string=html_cover)
        doc_criteria = HTML(string=html_criteria)
        doc_structure = HTML(string=html_score_structure)
        doc_attendance = HTML(string=html_attendance)
        doc_scores = HTML(string=html_scores)
        doc_summary = HTML(string=html_summary)
        doc_back = HTML(string=html_back)

        # --- Render ALL documents ---
        rendered_cover = doc_cover.render()
        rendered_criteria = doc_criteria.render()
        rendered_structure = doc_structure.render()
        rendered_attendance = doc_attendance.render()
        rendered_scores = doc_scores.render()
        rendered_summary = doc_summary.render()
        rendered_back = doc_back.render()
        # --- [เพิ่มส่วนนี้เพื่อ Debug] ---
        current_app.logger.debug("--- Pator05 Page Count Debug ---")
        current_app.logger.debug(f"Cover pages: {len(rendered_cover.pages)}")
        current_app.logger.debug(f"Criteria pages: {len(rendered_criteria.pages)}")
        current_app.logger.debug(f"Structure pages: {len(rendered_structure.pages)}")
        current_app.logger.debug(f"Attendance pages: {len(rendered_attendance.pages)}")
        current_app.logger.debug(f"Scores pages: {len(rendered_scores.pages)}")
        current_app.logger.debug(f"Summary pages: {len(rendered_summary.pages)}")
        current_app.logger.debug(f"Back pages: {len(rendered_back.pages)}")
        current_app.logger.debug("------------------------------------")
        # ------------------------------------
        # --- Collect all pages from rendered documents ---
        all_pages = []

        # [NEW] ใช้วิธีวนลูปและ .append() ทีละหน้าแทน .extend()
        for page in rendered_cover.pages:
            all_pages.append(page)
        for page in rendered_criteria.pages:
            all_pages.append(page)
        for page in rendered_structure.pages:
            all_pages.append(page)
        for page in rendered_attendance.pages:
            all_pages.append(page)
        for page in rendered_scores.pages:
            all_pages.append(page)
        for page in rendered_summary.pages:
            all_pages.append(page)
        for page in rendered_back.pages:
            all_pages.append(page)

        current_app.logger.debug(f"Total pages in all_pages list: {len(all_pages)}")
        # --- Generate the final PDF in memory using ALL pages ---
        # --- Generate the final PDF in memory using .copy() ---
        pdf_bytes = io.BytesIO()

        # [NEW] สร้างเอกสารใหม่โดยการ copy หน้าทั้งหมดใน all_pages
        # เราจะใช้ rendered_cover (หรือตัวไหนก็ได้) เป็นฐาน
        final_document = rendered_cover.copy(all_pages)

        # สั่ง .write_pdf() จากเอกสารใหม่ที่รวมทุกหน้าแล้ว
        final_document.write_pdf(target=pdf_bytes) 
        pdf_bytes.seek(0)

        # --- Create filename ---
        timestamp = datetime.now().strftime("%Y%m%d_%H%M")
        safe_subject_name = "".join(c if c.isalnum() else "_" for c in pator05_data['course_info']['subject_name'])
        safe_classroom_name = "".join(c if c.isalnum() else "_" for c in pator05_data['course_info']['classroom_name'])
        filename = f"Pator05_{pator05_data['course_info']['academic_year']}_{pator05_data['course_info']['semester_term']}_{safe_subject_name}_{safe_classroom_name}_{timestamp}.pdf"

        # --- Send the file to the user ---
        return send_file(
            pdf_bytes,
            mimetype='application/pdf',
            as_attachment=True,
            download_name=filename
        )

    except Exception as e:
        current_app.logger.error(f"Error generating Pator05 for course {course_id}: {e}", exc_info=True)
        flash(f'เกิดข้อผิดพลาดร้ายแรงขณะสร้างไฟล์ PDF: {e}', 'danger')
        return redirect(url_for('teacher.dashboard')) # Adjust redirect
    
@bp.route('/course/<int:course_id>/export/pator05/excel')
@login_required
def export_pator05_excel(course_id):
    """Generates and returns the Pator05 Excel file for the given course."""

    # 1. Get data using the existing service function
    pator05_data = get_pator05_data(course_id)
    if not pator05_data:
        flash('ไม่พบข้อมูลสำหรับสร้าง ปถ.05 หรือเกิดข้อผิดพลาดในการดึงข้อมูล', 'danger')
        return redirect(url_for('teacher.dashboard'))

    try:
        output = io.BytesIO()
        writer = pd.ExcelWriter(output, engine='openpyxl')

        # --- Sheet 1: Cover Info (ข้อมูลหน้าปก) ---
        cover_info_data = {
            "รายการ": [
                "รหัสวิชา", "รายวิชา", "หน่วยกิต", "ชม./สัปดาห์", "กลุ่มสาระฯ",
                "ระดับชั้น", "ห้องเรียน", "ภาคเรียน", "ปีการศึกษา",
                "ครูผู้สอน", "ครูที่ปรึกษา"
            ],
            "ข้อมูล": [
                pator05_data['course_info']['subject_code'],
                pator05_data['course_info']['subject_name'],
                pator05_data['course_info']['credit'],
                pator05_data['course_info']['hours_per_week'],
                pator05_data['course_info']['subject_group'],
                pator05_data['course_info']['grade_level'],
                pator05_data['course_info']['classroom_name'],
                pator05_data['course_info']['semester_term'],
                pator05_data['course_info']['academic_year'],
                ", ".join(pator05_data['course_info']['teachers']),
                ", ".join(pator05_data['course_info']['advisors']) if pator05_data['course_info']['advisors'] else '-'
            ]
        }
        df_cover = pd.DataFrame(cover_info_data)
        df_cover.to_excel(writer, sheet_name='ข้อมูลปก', index=False)

        # --- Sheet 2: Score Structure (โครงสร้างคะแนน) ---
        structure_rows = []
        for unit in pator05_data['score_structure']['units']:
            unit_index = pator05_data['score_structure']['units'].index(unit) + 1
            structure_rows.append({'ประเภท': f"หน่วยที่ {unit_index}", 'รายละเอียด': unit['title'], 'K': '', 'P': '', 'A': '', 'รวมหน่วย': '', 'กลางภาค': '', 'ปลายภาค': '', 'งานสำคัญ (ร)': ''})
            if unit['items']:
                 for item in unit['items']:
                     structure_rows.append({'ประเภท': '', 'รายละเอียด': item['indicator_description'], 'K': '', 'P': '', 'A': '', 'รวมหน่วย': '', 'กลางภาค': '', 'ปลายภาค': '', 'งานสำคัญ (ร)': ''})
            structure_rows.append({'ประเภท': 'รวมคะแนนหน่วย', 'รายละเอียด': '', 'K': unit['k_total'], 'P': unit['p_total'], 'A': unit['a_total'], 'รวมหน่วย': unit['unit_collected_total'], 'กลางภาค': unit['midterm_score'], 'ปลายภาค': unit['final_score'], 'งานสำคัญ (ร)': ", ".join([f"{si['dimension_code']} {si['name']}" for si in unit['summative_items_info']])})

        # Add Total Rows
        total_k = sum(u['k_total'] for u in pator05_data['score_structure']['units'])
        total_p = sum(u['p_total'] for u in pator05_data['score_structure']['units'])
        total_a = sum(u['a_total'] for u in pator05_data['score_structure']['units'])
        structure_rows.append({'ประเภท': 'รวมคะแนนทั้งหมด', 'รายละเอียด': '', 'K': total_k, 'P': total_p, 'A': total_a, 'รวมหน่วย': pator05_data['score_structure']['collected_total'], 'กลางภาค': pator05_data['score_structure']['midterm_total'], 'ปลายภาค': pator05_data['score_structure']['final_total'], 'งานสำคัญ (ร)': ''})
        structure_rows.append({'ประเภท': 'คะแนนรวมปรับตามสัดส่วน', 'รายละเอียด': '', 'K': '', 'P': '', 'A': '', 'รวมหน่วย': pator05_data['score_structure']['ratio_collected'], 'กลางภาค': pator05_data['score_structure']['ratio_midterm'], 'ปลายภาค': pator05_data['score_structure']['ratio_final'], 'งานสำคัญ (ร)': ''})

        df_structure = pd.DataFrame(structure_rows)
        # Reorder columns slightly for clarity
        df_structure = df_structure[['ประเภท', 'รายละเอียด', 'K', 'P', 'A', 'รวมหน่วย', 'กลางภาค', 'ปลายภาค', 'งานสำคัญ (ร)']]
        df_structure.to_excel(writer, sheet_name='โครงสร้างคะแนน', index=False)

        # --- Sheet 3: Attendance (เวลาเรียน) ---
        att_columns = ['เลขที่', 'เลข ป.ต.', 'ชื่อ-สกุล', 'สถานะ'] + [f"ชม.{i}" for i in range(1, pator05_data['course_info']['hours_per_week'] * 20 + 1)] + ['รวมมา', 'รวมขาด', 'รวมลา', 'รวมสาย']
        att_data = []
        att_map = {'PRESENT': '/', 'ABSENT': 'ข', 'LATE': 'ส', 'LEAVE': 'ล'}
        total_possible_hours = pator05_data.get('total_possible_hours', 0)

        for student in pator05_data['students_data']:
            row = {
                'เลขที่': student['roll_number'],
                'เลข ป.ต.': student['student_id'],
                'ชื่อ-สกุล': student['full_name'],
                'สถานะ': student['status']
            }
            att_summary = student['attendance_summary']

            # Calculate attendance status text
            status_text = 'ปกติ'
            if student['status'] != 'กำลังศึกษา':
                status_text = student['status']
            elif total_possible_hours > 0:
                attended = att_summary.get('present', 0) + att_summary.get('late', 0)
                att_percent = (attended / total_possible_hours) * 100
                if att_percent < 80:
                    status_text = f"มส {att_percent:.0f}%"
            row['สถานะ'] = status_text # Overwrite status with calculated text

            for i in range(1, pator05_data['course_info']['hours_per_week'] * 20 + 1):
                hour_key = f"H{i}"
                status = student['attendance'].get(hour_key, 'PRESENT')
                row[f"ชม.{i}"] = att_map.get(status, status)

            row['รวมมา'] = att_summary.get('present', 0)
            row['รวมขาด'] = att_summary.get('absent', 0)
            row['รวมลา'] = att_summary.get('leave', 0)
            row['รวมสาย'] = att_summary.get('late', 0)
            att_data.append(row)

        df_attendance = pd.DataFrame(att_data, columns=att_columns) # Ensure column order
        df_attendance.to_excel(writer, sheet_name='เวลาเรียน', index=False)


        # --- Sheet 4: Scores (บันทึกคะแนน) ---
        score_columns = ['เลขที่', 'เลข ป.ต.', 'ชื่อ-สกุล']
        unit_item_map = {} # Store items for each unit header
        for unit in pator05_data['score_structure']['units']:
             unit_id = unit['unit_id']
             unit_index = pator05_data['score_structure']['units'].index(unit) + 1
             items_in_unit = unit.get('graded_items_structure', [])
             unit_item_map[f'หน่วย{unit_index}_รวม'] = [item['id'] for item in items_in_unit] # Store item IDs for later sum
             for item in items_in_unit:
                 score_columns.append(f"หน่วย{unit_index}_{item['dimension']}_{item['id']}")
             score_columns.append(f"หน่วย{unit_index}_รวม") # Add column for unit total

        score_data = []
        for student in pator05_data['students_data']:
            row = {
                'เลขที่': student['roll_number'],
                'เลข ป.ต.': student['student_id'],
                'ชื่อ-สกุล': student['full_name']
            }
            student_scores = student.get('scores', {})
            unit_totals = student.get('unit_totals', {})

            for unit_idx, unit in enumerate(pator05_data['score_structure']['units']):
                unit_id = unit['unit_id']
                unit_index = unit_idx + 1
                items_in_unit = unit.get('graded_items_structure', [])
                for item in items_in_unit:
                    col_name = f"หน่วย{unit_index}_{item['dimension']}_{item['id']}"
                    row[col_name] = student_scores.get(item['id'], 0) # Default to 0 if no score

                # Add unit total score
                total_col_name = f"หน่วย{unit_index}_รวม"
                row[total_col_name] = unit_totals.get(unit_id, 0) # Use pre-calculated total

            score_data.append(row)

        df_scores = pd.DataFrame(score_data, columns=score_columns) # Ensure column order
        df_scores.to_excel(writer, sheet_name='บันทึกคะแนน', index=False)


        # --- Sheet 5: Summary (สรุปผล) ---
        # Reuse logic from summary.html template for scaled scores
        summary_columns = ['เลขที่', 'เลข ป.ต.', 'ชื่อ-สกุล', 'คะแนนเก็บ(ดิบ)', 'คะแนนเก็บ(จริง)', 'กลางภาค(ดิบ)', 'กลางภาค(จริง)', 'รวมระหว่างภาค', 'ปลายภาค(ดิบ)', 'ปลายภาค(จริง)', 'รวมคะแนน', 'เกรด', 'ผลแก้ตัว']
        summary_data = []

        show_midterm = pator05_data['score_structure']['midterm_total'] > 0
        show_final = pator05_data['score_structure']['final_total'] > 0
        max_coll_raw = pator05_data['score_structure']['collected_total'] or 0
        ratio_coll = pator05_data['score_structure']['ratio_collected'] or 0
        max_mid_raw = pator05_data['score_structure']['midterm_total'] or 0
        ratio_mid = pator05_data['score_structure']['ratio_midterm'] or 0
        max_final_raw = pator05_data['score_structure']['final_total'] or 0
        ratio_final = pator05_data['score_structure']['ratio_final'] or 0

        for student in pator05_data['students_data']:
            coll_raw = student.get('total_collected', 0) or 0
            mid_raw = student.get('midterm_score', 0) or 0 if show_midterm else 0
            final_raw = student.get('final_score', 0) or 0 if show_final else 0

            scaled_coll = (coll_raw / max_coll_raw * ratio_coll) if max_coll_raw > 0 else 0
            scaled_mid = (mid_raw / max_mid_raw * ratio_mid) if max_mid_raw > 0 and show_midterm else 0
            mid_period_total = scaled_coll + scaled_mid
            scaled_final = (final_raw / max_final_raw * ratio_final) if max_final_raw > 0 and show_final else 0
            grand_total_scaled = mid_period_total + scaled_final

            current_grade = student.get('final_grade', '') or ''
            original_grade = student.get('original_final_grade') or current_grade
            remedial_grade = ''
            if student.get('remediation_status') not in ['None', 'In Progress'] and current_grade in ['1', '1.5', '2', '2.5', '3', '3.5', '4']:
                remedial_grade = current_grade

            row = {
                'เลขที่': student['roll_number'],
                'เลข ป.ต.': student['student_id'],
                'ชื่อ-สกุล': student['full_name'],
                'คะแนนเก็บ(ดิบ)': coll_raw,
                'คะแนนเก็บ(จริง)': scaled_coll,
                'กลางภาค(ดิบ)': mid_raw if show_midterm else None,
                'กลางภาค(จริง)': scaled_mid if show_midterm else None,
                'รวมระหว่างภาค': mid_period_total,
                'ปลายภาค(ดิบ)': final_raw if show_final else None,
                'ปลายภาค(จริง)': scaled_final if show_final else None,
                'รวมคะแนน': grand_total_scaled,
                'เกรด': original_grade,
                'ผลแก้ตัว': remedial_grade or '-' # Show '-' if not remediated or passed originally
            }
            summary_data.append(row)

        df_summary = pd.DataFrame(summary_data, columns=summary_columns) # Ensure column order
        df_summary.to_excel(writer, sheet_name='สรุปผล', index=False, float_format="%.2f")


        # --- Save the Excel file ---
        writer.close() # Use close() instead of save() for ExcelWriter with BytesIO
        output.seek(0)

        # --- Create filename ---
        timestamp = datetime.now().strftime("%Y%m%d_%H%M")
        safe_subject_name = "".join(c if c.isalnum() else "_" for c in pator05_data['course_info']['subject_name'])
        safe_classroom_name = "".join(c if c.isalnum() else "_" for c in pator05_data['course_info']['classroom_name'])
        filename = f"Pator05_{pator05_data['course_info']['academic_year']}_{pator05_data['course_info']['semester_term']}_{safe_subject_name}_{safe_classroom_name}_{timestamp}.xlsx"

        # --- Send the file to the user ---
        return send_file(
            output,
            mimetype='application/vnd.openxmlformats-officedocument.spreadsheetml.sheet',
            as_attachment=True,
            download_name=filename
        )

    except Exception as e:
        current_app.logger.error(f"Error generating Pator05 Excel for course {course_id}: {e}", exc_info=True)
        flash(f'เกิดข้อผิดพลาดร้ายแรงขณะสร้างไฟล์ Excel: {e}', 'danger')
        return redirect(url_for('teacher.dashboard')) # Adjust redirect    
    
# --- [NEW] Helper Function for Google Credentials ---
def _get_google_creds(user):
    """
    Loads, refreshes, and saves Google credentials for a user.
    Returns refreshed credentials object or None on failure.
    """
    if not user.google_credentials_json:
        current_app.logger.error(f"User {user.id} has no Google credentials saved.")
        return None

    try:
        creds_data = json.loads(user.google_credentials_json)
        # เพิ่ม client_id/secret จาก config เข้าไปใน creds_data
        # เพราะ Credentials.from_authorized_user_info ต้องการมัน
        creds_data['client_id'] = current_app.config['GOOGLE_CLIENT_ID']
        creds_data['client_secret'] = current_app.config['GOOGLE_CLIENT_SECRET']
        
        creds = google.oauth2.credentials.Credentials.from_authorized_user_info(creds_data)

        if creds.expired and creds.refresh_token:
            current_app.logger.info(f"Refreshing Google token for user {user.id}...")
            creds.refresh(GoogleRequest())
            # บันทึก Token ใหม่ที่ refresh แล้วกลับลง DB
            user.google_credentials_json = creds.to_json()
            db.session.commit()
            current_app.logger.info(f"Token refreshed and saved for user {user.id}.")

        return creds
    except Exception as e:
        db.session.rollback()
        current_app.logger.error(f"Error refreshing Google token for user {user.id}: {e}", exc_info=True)
        # ถ้า Token มีปัญหา (เช่น ถูก revoke), บังคับให้ user login ใหม่
        # โดยการล้าง credentials เก่า
        user.google_credentials_json = None
        user.google_id = None # อาจจะต้องล้าง google_id ด้วยเพื่อให้ flow login ทำงานใหม่
        db.session.commit()
        return None
# --- [END NEW] ---
#     
@bp.route('/plan/<int:plan_id>/export/pdf')
@login_required
def export_lesson_plan_pdf(plan_id):
    """
    [REVISED v6 - Single Cover] Generates PDF with a single cover page
    followed by one page per LearningUnit.
    """
    plan = db.session.query(LessonPlan).options(
        joinedload(LessonPlan.subject), # Need subject for cover
        joinedload(LessonPlan.academic_year), # Need year for cover
        selectinload(LessonPlan.courses).selectinload(Course.teachers) # Need teachers for cover
    ).get(plan_id)
    if not plan or not any(current_user in c.teachers for c in plan.courses): abort(403)

    # --- Fetch Cover Data ---
    settings_keys = ['school_name', 'school_logo_path', 'school_affiliation', 'school_district', 'school_province']
    settings_q = Setting.query.filter(Setting.key.in_(settings_keys)).all()
    school_info_cover = {s.key: s.value for s in settings_q}
    if school_info_cover.get('school_logo_path'):
         school_info_cover['school_logo_url'] = url_for('static', filename=f"uploads/{school_info_cover['school_logo_path']}", _external=True)
    else: school_info_cover['school_logo_url'] = None

    teachers_cover = list(set(teacher for course in plan.courses for teacher in course.teachers)) # Get unique teachers
    teacher_names_cover = ", ".join([t.full_name for t in teachers_cover]) or '-'

    cover_data = {
        'school_info': school_info_cover,
        'plan_info': {
            'subject_name': plan.subject.name,
            'subject_code': plan.subject.subject_code,
            'academic_year': plan.academic_year.year,
            'teacher_names': teacher_names_cover
        }
    }
    # --- End Fetch Cover Data ---

    # --- Fetch Unit Data ---
    all_unit_render_data = get_lesson_plan_export_data(plan_id) # This now returns one dict per unit
    if not all_unit_render_data:
        flash('ไม่พบข้อมูลแผนการสอน หรือเกิดข้อผิดพลาด', 'danger')
        return redirect(url_for('teacher.workspace', plan_id=plan_id))

    try:
        all_rendered_pages = []
        base_doc = None

        # 1. Render Cover Page
        html_cover = render_template('exports/lesson_plan/plan_cover.html', data=cover_data)
        doc_cover = HTML(string=html_cover)
        rendered_cover = doc_cover.render()
        all_rendered_pages.extend(rendered_cover.pages)
        if not base_doc and rendered_cover.pages: base_doc = rendered_cover

        # 2. Render Unit Pages
        for unit_data in all_unit_render_data:
            html_string = render_template('exports/lesson_plan/learning_unit_plan.html', data=unit_data)
            doc = HTML(string=html_string)
            rendered_doc = doc.render()
            all_rendered_pages.extend(rendered_doc.pages)
            if not base_doc and rendered_doc.pages: base_doc = rendered_doc # Fallback

        if not all_rendered_pages or not base_doc:
             flash('ไม่สามารถสร้างหน้า PDF ได้ (ไม่มีเนื้อหา)', 'danger')
             return redirect(url_for('teacher.workspace', plan_id=plan_id))

        # Combine all pages
        pdf_bytes = io.BytesIO()
        base_doc.copy(all_rendered_pages).write_pdf(target=pdf_bytes)
        pdf_bytes.seek(0)

        # Create filename
        timestamp = datetime.now().strftime("%Y%m%d_%H%M")
        safe_subject_name = "".join(c if c.isalnum() else "_" for c in plan.subject.name)
        filename = f"FullLessonPlan_{plan.subject.subject_code}_{timestamp}.pdf"

        return send_file(pdf_bytes, mimetype='application/pdf', as_attachment=True, download_name=filename)

    except Exception as e:
        current_app.logger.error(f"Error generating Full Lesson Plan PDF for plan {plan_id}: {e}", exc_info=True)
        flash(f'เกิดข้อผิดพลาดขณะสร้างไฟล์ PDF รวม: {e}', 'danger')
        return redirect(url_for('teacher.workspace', plan_id=plan_id))


@bp.route('/plan/<int:plan_id>/export/docx')
@login_required
def export_lesson_plan_docx(plan_id):
    """
    [REVISED v7 - Single Cover & Merged Cells] Generates DOCX with cover page,
    unit pages using multi-row table, merged activity cell, and sub-topics.
    """
    plan = db.session.query(LessonPlan).options(
        joinedload(LessonPlan.subject),
        joinedload(LessonPlan.academic_year),
        selectinload(LessonPlan.courses).selectinload(Course.teachers)
    ).get(plan_id)
    if not plan or not any(current_user in c.teachers for c in plan.courses): abort(403)

    # --- Fetch Cover Data ---
    settings_keys = ['school_name', 'school_logo_path', 'school_affiliation', 'school_district', 'school_province']
    settings_q = Setting.query.filter(Setting.key.in_(settings_keys)).all()
    school_info_cover = {s.key: s.value for s in settings_q}
    # Logo handling for DOCX is complex, skipping for now, add placeholder
    teachers_cover = list(set(teacher for course in plan.courses for teacher in course.teachers))
    teacher_names_cover = ", ".join([t.full_name for t in teachers_cover]) or '-'
    # --- End Fetch Cover Data ---

    # --- Fetch Unit Data ---
    all_unit_render_data = get_lesson_plan_export_data(plan_id)
    if not all_unit_render_data:
        flash('ไม่พบข้อมูลแผนการสอน หรือเกิดข้อผิดพลาด', 'danger')
        return redirect(url_for('teacher.workspace', plan_id=plan_id))

    try:
        document = Document()
        # --- Helper function ---
        def add_para(text, bold=False, size=10, indent_first=False, indent_left_pt=0, parent=document, align=None, style=None):
            p = parent.add_paragraph(style=style); run = p.add_run(text); run.font.size = Pt(size)
            run.font.name = 'Sarabun';
            if bold: run.bold = True
            if indent_first: p.paragraph_format.first_line_indent = Pt(36)
            if indent_left_pt > 0: p.paragraph_format.left_indent = Pt(indent_left_pt)
            if align: p.alignment = align
            p.paragraph_format.space_after = Pt(0)
            return p
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        from docx.shared import Cm

        # --- 1. Add Cover Page Content ---
        # TODO: Add Logo image if path exists
        add_para(f"โรงเรียน{school_info_cover.get('school_name', '...')}", bold=True, size=16, align=WD_ALIGN_PARAGRAPH.CENTER)
        add_para(f"สังกัด {school_info_cover.get('school_affiliation', '...')}", size=12, align=WD_ALIGN_PARAGRAPH.CENTER)
        add_para(f"อำเภอ {school_info_cover.get('school_district', '...')} จังหวัด {school_info_cover.get('school_province', '...')}", size=12, align=WD_ALIGN_PARAGRAPH.CENTER)
        document.add_paragraph().add_run().add_break(docx.enum.text.WD_BREAK.LINE) # Spacer
        add_para("แผนการจัดการเรียนรู้", bold=True, size=20, align=WD_ALIGN_PARAGRAPH.CENTER)
        document.add_paragraph().add_run().add_break(docx.enum.text.WD_BREAK.LINE) # Spacer
        add_para(f"รายวิชา {plan.subject.name} ({plan.subject.subject_code})", size=14, align=WD_ALIGN_PARAGRAPH.CENTER)
        add_para(f"ปีการศึกษา {plan.academic_year.year}", size=14, align=WD_ALIGN_PARAGRAPH.CENTER)
        document.add_paragraph().add_run().add_break(docx.enum.text.WD_BREAK.LINE) # Spacer
        add_para("จัดทำโดย", size=12, align=WD_ALIGN_PARAGRAPH.CENTER)
        add_para(teacher_names_cover, size=12, align=WD_ALIGN_PARAGRAPH.CENTER)
        add_para("ตำแหน่ง ครู", size=12, align=WD_ALIGN_PARAGRAPH.CENTER) # TODO: Fetch actual position
        document.add_page_break()
        # --- End Cover Page ---

        # --- 2. Loop through each LearningUnit dictionary ---
        for i, unit_data in enumerate(all_unit_render_data):
            # No page break needed before first unit page
            plan_info = unit_data['plan_info']

            # --- Render Header (Always use 'ที่') ---
            add_para(f"โรงเรียน{unit_data['school_info'].get('school_name', '...')}", bold=True, size=11, align=WD_ALIGN_PARAGRAPH.CENTER)
            add_para(f"แผนการจัดการเรียนรู้ที่ {unit_data['plan_number']}", bold=True, size=14, align=WD_ALIGN_PARAGRAPH.CENTER)
            add_para(f"กลุ่มสาระการเรียนรู้ {plan_info['subject_group']} รายวิชา {plan_info['subject_name']} รหัสวิชา {plan_info['subject_code']}", size=10, align=WD_ALIGN_PARAGRAPH.CENTER)
            add_para(f"ชั้น {plan_info['grade_level']} หน่วยการเรียนรู้ที่ {plan_info['unit_sequence']} เรื่อง {plan_info['unit_title']} เวลา {plan_info['total_unit_hours']} ชั่วโมง", size=10, align=WD_ALIGN_PARAGRAPH.CENTER)
            # No subunit title needed here as per new logic
            add_para(f"ครูผู้สอน {plan_info['teacher_names']}", size=10, align=WD_ALIGN_PARAGRAPH.CENTER)
            document.add_paragraph()

            # --- Render Multi-Row Content Table with Merged Cell ---
            table = document.add_table(rows=6, cols=2) # Create all 6 rows
            table.style = 'Table Grid'

            # Merge cells for Activity (Cell 0,1 merges with 1,1 and 2,1)
            activity_cell = table.cell(0, 1)
            activity_cell.merge(table.cell(2, 1))

            # --- Populate Cells using Helper ---
            def populate_cell(cell, section_title, content_list=[], content_text="", list_prefix="- ", use_pre_wrap=False, sub_indent_pt=36, is_list=False, structured_list=None):
                # Clear existing paragraph(s) in the cell first
                for p in cell.paragraphs:
                    p._element.getparent().remove(p._element)
                    
                add_para(section_title, bold=True, parent=cell, size=10)
                if structured_list: # For Competencies/Characteristics
                     for item in structured_list:
                        add_para(f"{list_prefix}{item['main']}", parent=cell, size=10, indent_left_pt=sub_indent_pt)
                        if item['subs']:
                            for sub_item in item['subs']:
                                add_para(f"- {sub_item}", parent=cell, size=10, indent_left_pt=sub_indent_pt + 18)
                elif content_list:
                    for item_idx, item in enumerate(content_list):
                        indent = sub_indent_pt
                        prefix = list_prefix if is_list else ""
                        if isinstance(item, str) and item.startswith("มาตรฐาน"): indent = sub_indent_pt # No prefix for standard
                        elif isinstance(item, str) and item.startswith("ตัวชี้วัด"): indent = sub_indent_pt + 18; prefix="" # Deeper indent for indicator
                        add_para(f"{prefix}{item}", parent=cell, size=10, indent_left_pt=indent)
                elif content_text:
                    if use_pre_wrap:
                         lines = content_text.split('\n')
                         for line_idx, line in enumerate(lines):
                              # Add paragraph even if line is empty to preserve spacing
                              add_para(line if line.strip() else " ", parent=cell, size=10, indent_left_pt=sub_indent_pt if line_idx > 0 or not line.startswith("ชั่วโมงที่") else 0 ) # Indent lines after the first, unless it's "ชั่วโมงที่..."
                    else:
                         add_para(content_text, parent=cell, size=10, indent_left_pt=sub_indent_pt)

            # Row 1
            populate_cell(table.cell(0, 0), "1. มาตรฐานการเรียนรู้ และตัวชี้วัด", content_list=unit_data['learning_standard']['standards_text'] + unit_data['learning_standard']['indicators_text'])
            populate_cell(activity_cell, "7. กิจกรรมการเรียนรู้", content_text=unit_data.get('activities',''), use_pre_wrap=True, sub_indent_pt=0) # No initial indent for activity

            # Row 2
            populate_cell(table.cell(1, 0), "2. จุดประสงค์การเรียนรู้", content_text=unit_data.get('learning_objectives',''), use_pre_wrap=True)
            # Cell 1,1 is merged

            # Row 3
            populate_cell(table.cell(2, 0), "3. สาระสำคัญ", content_text=unit_data.get('core_concepts',''), use_pre_wrap=True)
            # Cell 2,1 is merged

            # Row 4
            populate_cell(table.cell(3, 0), "4. สาระการเรียนรู้", content_text=unit_data.get('learning_content',''), use_pre_wrap=True)
            populate_cell(table.cell(3, 1), "8. การวัดและประเมินผลการเรียนรู้", content_list=unit_data.get('assessment_methods',[]), is_list=True)

            # Row 5
            populate_cell(table.cell(4, 0), "5. สมรรถนะสำคัญของผู้เรียน", structured_list=unit_data.get('competencies_structured', []), list_prefix="") # Use structured data
            populate_cell(table.cell(4, 1), "9. สื่อและแหล่งเรียนรู้", content_text=unit_data.get('media_sources',''), use_pre_wrap=True)

            # Row 6
            populate_cell(table.cell(5, 0), "6. คุณลักษณะอันพึงประสงค์", structured_list=unit_data.get('desired_characteristics_structured', []), list_prefix="") # Use structured data
            # Populate Log cell
            log_cell = table.cell(5, 1)
            for p in log_cell.paragraphs: p._element.getparent().remove(p._element) # Clear cell
            log_data = unit_data.get('post_teaching_log', {})
            add_para("10. บันทึกผลหลังการจัดการเรียนรู้", bold=True, parent=log_cell, size=10)
            add_para("ผลการจัดการเรียนรู้:", indent_left_pt=36, parent=log_cell, size=10)
            add_para(log_data.get('log_content') or '...', indent_left_pt=36, parent=log_cell, size=10)
            if log_data.get('problems_obstacles'):
                add_para("ปัญหาและอุปสรรค:", indent_left_pt=36, parent=log_cell, size=10)
                add_para(log_data['problems_obstacles'], indent_left_pt=36, parent=log_cell, size=10)
            if log_data.get('solutions'):
                add_para("ข้อเสนอแนะและแนวทางแก้ไข:", indent_left_pt=36, parent=log_cell, size=10)
                add_para(log_data['solutions'], indent_left_pt=36, parent=log_cell, size=10)

            # --- Signature ---
            document.add_paragraph()
            sig_para = add_para("ลงชื่อ.............................. ครูผู้สอน", size=10, align=WD_ALIGN_PARAGRAPH.RIGHT)
            add_para(f"({plan_info['teacher_names']})", size=10, align=WD_ALIGN_PARAGRAPH.RIGHT)
            add_para("ตำแหน่ง ครู", size=10, align=WD_ALIGN_PARAGRAPH.RIGHT)

            # Add page break after each unit except the last one
            if i < len(all_unit_render_data) - 1:
                 document.add_page_break()


        # --- Save DOCX ---
        docx_bytes = io.BytesIO(); document.save(docx_bytes); docx_bytes.seek(0)
        # --- Create filename ---
        timestamp = datetime.now().strftime("%Y%m%d_%H%M")
        safe_subject_name = "".join(c if c.isalnum() else "_" for c in plan.subject.name)
        filename = f"FullLessonPlan_{plan.subject.subject_code}_{timestamp}.docx"
        # --- Send file ---
        return send_file(docx_bytes, mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document', as_attachment=True, download_name=filename)

    except Exception as e:
        current_app.logger.error(f"Error generating Full Lesson Plan DOCX for plan {plan_id}: {e}", exc_info=True)
        flash(f'เกิดข้อผิดพลาดขณะสร้างไฟล์ Word รวม: {e}', 'danger')
        return redirect(url_for('teacher.workspace', plan_id=plan_id))

@bp.route('/api/subject/<int:subject_id>/previous-plans')
@login_required
def get_previous_plans(subject_id):
    """ Finds previous lesson plans for the same subject. """
    target_year_id = request.args.get('target_year_id', type=int)
    if not target_year_id:
        return jsonify([]) # Return empty if target year is missing

    target_year = db.session.get(AcademicYear, target_year_id)
    if not target_year:
         return jsonify([]) # Return empty if target year not found

    # Find plans for the same subject in previous years
    previous_plans_q = LessonPlan.query.join(AcademicYear).options(
        joinedload(LessonPlan.academic_year),
        # Eager load teachers associated through ANY course linked to the plan
        selectinload(LessonPlan.courses).selectinload(Course.teachers)
    ).filter(
        LessonPlan.subject_id == subject_id,
        AcademicYear.year < target_year.year # Only look at years BEFORE the target year
    ).order_by(AcademicYear.year.desc()).all()

    results = []
    for plan in previous_plans_q:
         # Get unique teachers for this plan
        teachers = list(set(teacher for course in plan.courses for teacher in course.teachers))
        teacher_names = ", ".join([t.full_name for t in teachers]) if teachers else None

        # Check if there are any logs associated with this plan's units
        has_logs = db.session.query(PostTeachingLog.id).join(LearningUnit).filter(
            LearningUnit.lesson_plan_id == plan.id
        ).limit(1).scalar() is not None

        results.append({
            'plan_id': plan.id,
            'year': plan.academic_year.year,
            'teacher_names': teacher_names,
            'has_logs': has_logs
        })

    return jsonify(results)

@bp.route('/api/plan/<int:plan_id>/teaching-logs')
@login_required
def get_teaching_logs(plan_id):
    """ Gets teaching logs for a specific plan, rendered as HTML. """
    plan = db.session.get(LessonPlan, plan_id)

    # Basic permission check - Adjust as needed
    is_involved = False
    if plan:
        # Check if the current user is assigned to teach the SAME subject in the CURRENT semester
        current_semester = Semester.query.filter_by(is_current=True).first()
        if current_semester:
            is_involved = Course.query.filter(
                Course.subject_id == plan.subject_id,
                Course.semester_id == current_semester.id,
                Course.teachers.any(id=current_user.id)
            ).limit(1).scalar() is not None
        # OR allow admin/academic roles
        # is_allowed = is_involved or current_user.has_role('Admin') or current_user.has_role('Academic')


    if not plan or not is_involved: # Use the determined permission
        return "ไม่พบแผนการสอน หรือไม่มีสิทธิ์เข้าถึง", 404

    # Fetch logs grouped by LearningUnit
    logs_by_unit = db.session.query(PostTeachingLog).join(LearningUnit).options(
         joinedload(PostTeachingLog.author), # Load teacher info
         joinedload(PostTeachingLog.unit) # Load unit info
    ).filter(
         LearningUnit.lesson_plan_id == plan_id
    ).order_by(LearningUnit.sequence, PostTeachingLog.created_at).all()

    # Group logs manually for the template
    grouped_logs = defaultdict(list)
    for log in logs_by_unit:
         grouped_logs[log.unit].append(log)

    # Render a partial template (ensure '_previous_logs_content.html' exists)
    return render_template('teacher/_previous_logs_content.html', grouped_logs=grouped_logs)


@bp.route('/api/lesson-plan/import', methods=['POST'])
@login_required
def import_or_create_plan():
    """ Handles importing an existing plan or creating a blank one. """
    data = request.get_json()
    subject_id = data.get('subject_id')
    target_academic_year_id = data.get('target_academic_year_id')
    source_plan_id = data.get('source_plan_id') # Can be null

    if not subject_id or not target_academic_year_id:
        return jsonify({'status': 'error', 'message': 'ข้อมูลไม่ครบถ้วน (Subject ID หรือ Target Year ID)'}), 400

    # --- Permission Check: Does the current user teach this subject in the target year? ---
    target_year = db.session.get(AcademicYear, target_academic_year_id)
    if not target_year:
         return jsonify({'status': 'error', 'message': 'ปีการศึกษาเป้าหมายไม่ถูกต้อง'}), 400

    # Find the current semester within the target year (or any if current not set for future)
    target_semester = Semester.query.filter_by(academic_year_id=target_year.id, is_current=True).first()
    if not target_semester:
         target_semester = Semester.query.filter_by(academic_year_id=target_year.id).order_by(Semester.term).first()

    if not target_semester:
         return jsonify({'status': 'error', 'message': 'ไม่พบภาคเรียนสำหรับปีการศึกษาเป้าหมาย'}), 400


    # Check if user is assigned to teach this subject in ANY classroom during the target semester
    is_assigned = Course.query.filter(
        Course.subject_id == subject_id,
        Course.semester_id == target_semester.id,
        Course.teachers.any(id=current_user.id)
    ).limit(1).scalar() is not None

    if not is_assigned:
         # Consider allowing Admin/Academic roles to bypass this
         # if not current_user.has_role('Admin') and not current_user.has_role('Academic'):
         return jsonify({'status': 'error', 'message': 'คุณไม่ได้รับมอบหมายให้สอนวิชานี้ในปีการศึกษาเป้าหมาย'}), 403
    # --- End Permission Check ---


    # --- Check if plan already exists ---
    existing_plan = LessonPlan.query.filter_by(
        subject_id=subject_id,
        academic_year_id=target_academic_year_id
    ).first()
    if existing_plan:
         return jsonify({'status': 'error', 'message': 'มีแผนการสอนสำหรับวิชา/ปีการศึกษานี้อยู่แล้ว'}), 400
    # --- End Check ---

    new_plan_id = None
    error_message = None

    if source_plan_id:
        # --- Import (Copy) ---
        # Basic check if source plan exists (more robust permission might be needed)
        source_plan = db.session.get(LessonPlan, source_plan_id)
        if not source_plan:
            return jsonify({'status': 'error', 'message': 'ไม่พบแผนต้นทาง'}), 404

        success, result_data = copy_lesson_plan(
            source_plan_id=source_plan_id,
            target_academic_year_id=target_academic_year_id,
            current_user_id=current_user.id
        )
        if success:
            new_plan_id = result_data
        else:
            error_message = result_data # Contains the error message from copy_lesson_plan
    else:
        # --- Create Blank ---
        success, result_data = create_blank_lesson_plan(
            subject_id=subject_id,
            academic_year_id=target_academic_year_id
        )
        if success:
            new_plan_id = result_data
        else:
            error_message = result_data # Contains the error message

    if new_plan_id:
        # --- Link the new plan to the relevant Course(s) ---
        courses_to_link = Course.query.filter(
             Course.subject_id == subject_id,
             Course.semester_id == target_semester.id, # Use target_semester found earlier
             Course.teachers.any(id=current_user.id)
         ).all()
        linked_count = 0
        for course in courses_to_link:
            if course.lesson_plan_id is None: # Only link if the course doesn't already have a plan
                 course.lesson_plan_id = new_plan_id
                 linked_count += 1
        if linked_count > 0:
             try:
                  db.session.commit()
             except Exception as e:
                   db.session.rollback()
                   current_app.logger.error(f"Error linking new plan {new_plan_id} to courses: {e}")
                   # Proceed to return success, linking is secondary for now

        return jsonify({'status': 'success', 'new_plan_id': new_plan_id})
    else:
        # Use 400 Bad Request for logical errors during creation/import
        return jsonify({'status': 'error', 'message': error_message or 'ไม่สามารถสร้าง/นำเข้าแผนได้'}), 400

def copy_lesson_plan(source_plan_id: int, target_academic_year_id: int, current_user_id: int):
    """
    Creates a copy of a lesson plan for a new academic year.

    Args:
        source_plan_id: ID of the LessonPlan to copy.
        target_academic_year_id: ID of the AcademicYear for the new plan.
        current_user_id: ID of the user performing the copy.

    Returns:
        Tuple (bool, Union[int, str]): (True, new_plan_id) on success,
                                       (False, error_message) on failure.
    """
    try:
        # 1. Fetch Source Plan with related data eagerly
        source_plan = db.session.query(LessonPlan).options(
            selectinload(LessonPlan.subject), # Load subject for error message
            selectinload(LessonPlan.learning_units).options(
                selectinload(LearningUnit.graded_items).selectinload(GradedItem.dimension), # Load dimension too
                selectinload(LearningUnit.assessment_items).selectinload(AssessmentItem.topic), # Load topic
                selectinload(LearningUnit.indicators).selectinload(Indicator.standard), # Standard indicators + standard
                selectinload(LearningUnit.sub_units).options(
                     selectinload(SubUnit.indicators).selectinload(Indicator.standard),
                     selectinload(SubUnit.graded_items).selectinload(GradedItem.dimension),
                     selectinload(SubUnit.assessment_items).selectinload(AssessmentItem.topic)
                )
            ),
            selectinload(LessonPlan.custom_indicators).selectinload(Indicator.standard), # Custom indicators + standard
            selectinload(LessonPlan.constraints)
        ).get(source_plan_id)

        if not source_plan:
            return False, "ไม่พบแผนการสอนต้นทาง"

        # 2. Check for Existing Target Plan
        target_year = db.session.get(AcademicYear, target_academic_year_id) # Get target year object for message
        if not target_year:
             return False, "ไม่พบปีการศึกษาเป้าหมาย"

        existing_target_plan = LessonPlan.query.filter_by(
            subject_id=source_plan.subject_id,
            academic_year_id=target_academic_year_id
        ).first()
        if existing_target_plan:
            subject_name = source_plan.subject.name if source_plan.subject else f"ID {source_plan.subject_id}"
            return False, f"มีแผนการสอนสำหรับวิชา {subject_name} ในปีการศึกษา {target_year.year} อยู่แล้ว"

        # 3. Create New Plan Object
        new_plan = LessonPlan(
            subject_id=source_plan.subject_id,
            academic_year_id=target_academic_year_id,
            target_mid_ratio=source_plan.target_mid_ratio,
            target_final_ratio=source_plan.target_final_ratio,
            status='ฉบับร่าง', # Always start as draft
            revision_notes=None, # Clear revision notes
            manual_scheduling_notes=source_plan.manual_scheduling_notes # Keep manual notes
        )
        db.session.add(new_plan)
        # Flush to get the new_plan.id needed for custom indicators
        db.session.flush()

        # --- Mappings to track old IDs to new objects ---
        unit_map = {} # old_unit_id -> new_unit_object
        graded_item_map = {} # old_graded_item_id -> new_graded_item_object
        assessment_item_map = {} # old_assessment_item_id -> new_assessment_item_object

        # 4. Deep Copy Learning Units and their contents
        for source_unit in sorted(source_plan.learning_units, key=lambda u: u.sequence):
            new_unit = LearningUnit(
                lesson_plan=new_plan, # Associate with new plan
                title=source_unit.title,
                sequence=source_unit.sequence,
                midterm_score=source_unit.midterm_score,
                final_score=source_unit.final_score,
                topic=source_unit.topic,
                hours=source_unit.hours,
                learning_objectives=source_unit.learning_objectives,
                learning_content=source_unit.learning_content,
                learning_activities=source_unit.learning_activities,
                core_concepts=source_unit.core_concepts,
                # activities=source_unit.activities, # Deprecated
                media_sources=source_unit.media_sources
                # DO NOT COPY reflections or PostTeachingLog
            )
            db.session.add(new_unit)
            db.session.flush() # Get new_unit.id
            unit_map[source_unit.id] = new_unit

            # 4.1 Copy Standard Indicator relationships (link to existing ones)
            # Filter out custom indicators before assigning
            standard_indicators = [ind for ind in source_unit.indicators if ind.creator_type == 'ADMIN']
            new_unit.indicators = standard_indicators

            # 4.2 Deep Copy Graded Items
            for source_item in source_unit.graded_items:
                new_item = GradedItem(
                    learning_unit=new_unit, # Associate with new unit
                    name=source_item.name,
                    max_score=source_item.max_score,
                    indicator_type=source_item.indicator_type,
                    assessment_type=source_item.assessment_type,
                    assessment_dimension_id=source_item.assessment_dimension_id,
                    is_group_assignment=source_item.is_group_assignment
                )
                db.session.add(new_item)
                db.session.flush() # Get new_item.id
                graded_item_map[source_item.id] = new_item

            # 4.3 Deep Copy Assessment Items
            for source_assess_item in source_unit.assessment_items:
                new_assess_item = AssessmentItem(
                    unit=new_unit, # Associate with new unit
                    assessment_topic_id=source_assess_item.assessment_topic_id
                )
                db.session.add(new_assess_item)
                db.session.flush() # Get new_assess_item.id
                assessment_item_map[source_assess_item.id] = new_assess_item

            # 4.4 Deep Copy SubUnits (if any)
            for source_sub_unit in sorted(source_unit.sub_units, key=lambda su: su.hour_sequence):
                 new_sub_unit = SubUnit(
                      learning_unit=new_unit, # Associate with new unit
                      title=source_sub_unit.title,
                      hour_sequence=source_sub_unit.hour_sequence,
                      activities=source_sub_unit.activities
                 )
                 db.session.add(new_sub_unit) # Add first to allow relationship assignment
                 db.session.flush() # Get new_sub_unit ID

                 # Copy relationships using the maps created
                 sub_standard_indicators = [ind for ind in source_sub_unit.indicators if ind.creator_type == 'ADMIN']
                 new_sub_unit.indicators = sub_standard_indicators # Link to existing standard ones
                 new_sub_unit.graded_items = [graded_item_map[gi.id] for gi in source_sub_unit.graded_items if gi.id in graded_item_map]
                 new_sub_unit.assessment_items = [assessment_item_map[ai.id] for ai in source_sub_unit.assessment_items if ai.id in assessment_item_map]
                 # No need to add again, already added

        # 5. Deep Copy Custom Indicators (if any) linked directly to the PLAN
        # Also copy custom indicators linked to UNITS/SUBUNITS
        all_source_custom_indicators = list(source_plan.custom_indicators) # Plan level
        for su in source_plan.learning_units:
             all_source_custom_indicators.extend([ind for ind in su.indicators if ind.creator_type == 'TEACHER']) # Unit level
             for sub in su.sub_units:
                  all_source_custom_indicators.extend([ind for ind in sub.indicators if ind.creator_type == 'TEACHER']) # SubUnit level

        # Use a set to handle potential duplicates if an indicator is linked multiple times
        unique_source_custom_indicators = {ind.id: ind for ind in all_source_custom_indicators}.values()

        custom_indicator_map = {} # old_custom_id -> new_custom_object
        for source_custom_ind in unique_source_custom_indicators:
            # Check if this custom indicator was linked to a unit/subunit that was copied
            # We need to decide where the *new* custom indicator should be linked.
            # Simplest approach: Link all copied custom indicators ONLY to the new PLAN.
            new_custom_ind = Indicator(
                lesson_plan=new_plan, # Link to new plan ONLY
                code=source_custom_ind.code,
                description=source_custom_ind.description,
                standard_id=source_custom_ind.standard_id, # Assume custom standard exists
                creator_type='TEACHER',
                creator_id=current_user_id # Assign to the user doing the copy
            )
            db.session.add(new_custom_ind)
            db.session.flush()
            custom_indicator_map[source_custom_ind.id] = new_custom_ind

        # 6. Copy Constraints
        for source_constraint in source_plan.constraints:
            new_constraint = LessonPlanConstraint(
                lesson_plan=new_plan, # Link to new plan
                constraint_type=source_constraint.constraint_type,
                value=source_constraint.value
            )
            db.session.add(new_constraint)

        # 7. Commit all changes
        db.session.commit()

        # 8. Return success and the new plan ID
        return True, new_plan.id

    except Exception as e:
        db.session.rollback()
        # Use Flask's logger
        current_app.logger.error(f"Error copying lesson plan {source_plan_id}: {e}", exc_info=True)
        return False, f"เกิดข้อผิดพลาดในการคัดลอก: {str(e)}"

@bp.route('/history')
@login_required
def teaching_history():
    # ดึงค่า filter จาก URL query string
    selected_year_id = request.args.get('year_id', type=int)
    selected_term = request.args.get('semester_term', type=int)

    # Query ปี/เทอม ทั้งหมดสำหรับ Dropdown filter
    all_years = AcademicYear.query.order_by(AcademicYear.year.desc()).all()
    # Note: We query semesters dynamically based on year selection below

    courses_in_period = []
    selected_semester = None

    if selected_year_id and selected_term:
        selected_semester = Semester.query.filter_by(academic_year_id=selected_year_id, term=selected_term).options(
             joinedload(Semester.academic_year) # Load year info for display
        ).first()
        if selected_semester:
            courses_in_period = Course.query.filter(
                Course.teachers.any(id=current_user.id),
                Course.semester_id == selected_semester.id
            ).join(Course.subject).join(Course.classroom).options(
                joinedload(Course.subject),
                joinedload(Course.classroom)
            ).order_by(Subject.subject_code, Classroom.name).all()
        else:
            flash('ไม่พบข้อมูลภาคเรียนที่ระบุ', 'warning')


    return render_template('teacher/history.html',
                           title="ประวัติการสอนและผลการเรียน",
                           all_years=all_years,
                           courses=courses_in_period,
                           selected_year_id=selected_year_id,
                           selected_term=selected_term,
                           selected_semester=selected_semester) # Pass selected semester object

@bp.route('/history/course/<int:course_id>/grades')
@login_required
def view_historical_grades(course_id):
    # semester_id จำเป็นสำหรับการดึงข้อมูลที่ถูกต้องและ permission check
    semester_id = request.args.get('semester_id', type=int)
    if not semester_id:
         abort(400, "Missing semester_id parameter")

    course = Course.query.options(
        joinedload(Course.subject),
        joinedload(Course.classroom),
        joinedload(Course.semester).joinedload(Semester.academic_year) # Load semester/year info
    ).get_or_404(course_id)

    # Permission check: Did the current user teach this course IN THIS SEMESTER?
    if current_user not in course.teachers or course.semester_id != semester_id:
         abort(403) # Forbidden

    # ใช้ Service เดิมในการคำนวณ/ดึงข้อมูลเกรด
    # This service fetches the CourseGrade object which contains the final grades
    student_grades, max_scores = calculate_final_grades_for_course(course)

    return render_template('teacher/history_grades_view.html',
                           title=f"ผลการเรียนย้อนหลัง: {course.subject.name}",
                           course=course, # Pass the course object with semester info
                           student_grades=student_grades,
                           max_collected=max_scores['collected'],
                           max_midterm=max_scores['midterm'],
                           max_final=max_scores['final'],
                           grand_max_score=max_scores['grand_total'])

#
# ROUTE: Mobile Classroom Hub
#
@bp.route('/mobile/entry/<int:entry_id>')
@login_required
# @initial_setup_required
def mobile_entry(entry_id):
    
    # --- 1. ดึงข้อมูลคาบเรียน, วันที่, และนักเรียน ---
    date_iso = request.args.get('date', date.today().isoformat())
    attendance_date = date.fromisoformat(date_iso)

    entry = db.session.get(TimetableEntry, entry_id)
    if not entry: abort(404)
        
    course = entry.course
    classroom = entry.course.classroom
    semester = course.semester 
    lesson_plan = course.lesson_plan 

    if current_user not in course.teachers:
        abort(403)

    enrollments = Enrollment.query.join(
        Student, Enrollment.student_id == Student.id
    ).filter(
        Enrollment.classroom_id == classroom.id
    ).options(
        contains_eager(Enrollment.student)
    ).order_by(
        Enrollment.roll_number,
        Student.student_id 
    ).all()
    student_ids = [e.student_id for e in enrollments]

    # --- 2. คำนวณลำดับคาบเรียน (Hour Sequence) ---
    hour_sequence = 1 
    if not semester.start_date:
        current_app.logger.warning(f"Course {course.id}: ไม่ได้ตั้งค่าวันเริ่มเทอม (Semester Start Date)")
    else:
        all_entries_for_course = TimetableEntry.query.filter_by(
            course_id=course.id
        ).join(
            WeeklyScheduleSlot, TimetableEntry.weekly_schedule_slot_id == WeeklyScheduleSlot.id
        ).options(
            contains_eager(TimetableEntry.slot)
        ).all()
        
        slots_by_day = defaultdict(list)
        for entry_in_course in all_entries_for_course:
            if entry_in_course.slot: 
                slots_by_day[entry_in_course.slot.day_of_week].append(entry_in_course.slot.period_number)
        
        teaching_days_of_week = set(slots_by_day.keys())
        temp_hour_count = 0
        current_date_loop = semester.start_date
        
        while current_date_loop <= attendance_date:
            day_iso = current_date_loop.isoweekday()
            if day_iso in teaching_days_of_week:
                periods_on_this_day = sorted(slots_by_day[day_iso])
                if current_date_loop == attendance_date:
                    current_period_num = entry.slot.period_number
                    try:
                        period_index_on_day = periods_on_this_day.index(current_period_num)
                        temp_hour_count += (period_index_on_day + 1) 
                    except ValueError:
                        current_app.logger.error(f"Period {current_period_num} not in schedule for day {day_iso}!")
                        temp_hour_count += 1 # Failsafe
                    break 
                else:
                    temp_hour_count += len(periods_on_this_day)
            current_date_loop += timedelta(days=1) 
        
        hour_sequence = temp_hour_count if temp_hour_count > 0 else 1
        
    current_app.logger.info(f"Mobile Entry: Course {course.id} on {attendance_date} is Hour Sequence: {hour_sequence}")

    # --- 3. [FIXED] ค้นหา "หน่วยแม่" (Parent Unit) โดยใช้ Hour Sequence ---
    parent_unit = None
    suggested_unit_id = None 

    if lesson_plan:
        # 3.1 [FIX] ค้นหา "หน่วยแม่" โดยการนับชั่วโมงสะสม
        all_units_in_plan = LearningUnit.query.filter_by(
            lesson_plan_id=lesson_plan.id
        ).order_by(LearningUnit.sequence).all()
        
        total_hours_so_far = 0
        for unit in all_units_in_plan:
            if unit.hours and unit.hours > 0:
                if hour_sequence <= total_hours_so_far + unit.hours:
                    parent_unit = unit # Found it!
                    suggested_unit_id = unit.id
                    current_app.logger.info(f"Calculated Parent Unit {unit.id} for Hour {hour_sequence}")
                    break # Exit loop
                total_hours_so_far += unit.hours
        
        if not parent_unit:
            current_app.logger.warning(f"Could not calculate parent unit for Hour {hour_sequence}.")

    # --- 4. ดึงข้อมูลการเช็คชื่อ (เหมือนเดิม) ---
    existing_records_list = AttendanceRecord.query.filter(
        AttendanceRecord.timetable_entry_id == entry_id,
        AttendanceRecord.attendance_date == attendance_date
    ).all()
    attendance_records = {rec.student_id: rec.status for rec in existing_records_list}
    
    # --- 5. [FIXED] ดึงข้อมูลคะแนนเก็บ (Graded Item) จาก "หน่วยแม่" ---
    graded_items = []
    if parent_unit:
        # Load all graded items from the PARENT unit
        parent_with_items = LearningUnit.query.options(
            selectinload(LearningUnit.graded_items)
        ).get(parent_unit.id)
        graded_items = parent_with_items.graded_items if parent_with_items else []
    
    existing_scores_list = Score.query.filter(
        Score.graded_item_id.in_([item.id for item in graded_items]),
        Score.student_id.in_([e.student_id for e in enrollments]) # 👈 [FIX] ใช้ student_ids
    ).all()
    scores = {f"{s.student_id}-{s.graded_item_id}": s.score for s in existing_scores_list}

    # --- 6. ดึงข้อมูลกลุ่มของนักเรียน (Student Group Map) (เหมือนเดิม) ---
    student_group_map_json = json.dumps({}) 
    if lesson_plan:
        all_groups = StudentGroup.query.filter_by(
            lesson_plan_id=lesson_plan.id, 
            course_id=course.id
        ).options(
            selectinload(StudentGroup.enrollments)
        ).all()
        student_group_map = {
            en.student_id: group.id 
            for group in all_groups 
            for en in group.enrollments 
            if en.student_id in student_ids 
        }
        student_group_map_json = json.dumps(student_group_map)

    # --- 6.1 สร้างรายการประเมินสำหรับ Dropdown (เหมือนเดิม) ---
    assessment_topics_for_dropdown = []
    final_topic_list = []
    topic_key_map = {}
    unit_names_map = {} 
    
    if lesson_plan:
        unit_names_map = { unit.id: unit.title for unit in all_units_in_plan } # ใช้ all_units_in_plan ที่ดึงมาแล้ว
        
        all_plan_assessment_items = AssessmentItem.query.join(
            LearningUnit, AssessmentItem.learning_unit_id == LearningUnit.id
        ).filter(
            LearningUnit.lesson_plan_id == lesson_plan.id
        ).options(
            joinedload(AssessmentItem.topic).joinedload(AssessmentTopic.parent), 
            joinedload(AssessmentItem.unit)
        ).all()

        all_parent_ids = set()
        for item in all_plan_assessment_items:
            if item.topic and item.topic.parent_id:
                all_parent_ids.add(item.topic.parent_id)
                if item.topic.parent and item.topic.parent.parent_id:
                     all_parent_ids.add(item.topic.parent.parent_id)
        
        parent_topic_objects = {}
        if all_parent_ids:
            parents = AssessmentTopic.query.filter(AssessmentTopic.id.in_(all_parent_ids)).all()
            parent_topic_objects = {p.id: p for p in parents}

        for item in all_plan_assessment_items:
            if not (item.topic and item.unit):
                continue
            current_topic = item.topic
            current_unit_id = item.unit.id
            current_unit_name = unit_names_map.get(current_unit_id, 'N/A')
            item_key = f"{current_unit_id}-{current_topic.id}"
            if item_key not in topic_key_map:
                final_topic_list.append({
                    'topic_id': current_topic.id,
                    'topic_name': current_topic.name,
                    'unit_id': current_unit_id,
                    'unit_name': current_unit_name,
                    'parent_id': current_topic.parent_id
                })
                topic_key_map[item_key] = True
            if current_topic.parent_id and current_topic.parent_id in parent_topic_objects:
                parent = parent_topic_objects[current_topic.parent_id]
                parent_key = f"{current_unit_id}-{parent.id}"
                if parent_key not in topic_key_map:
                    final_topic_list.append({
                        'topic_id': parent.id,
                        'topic_name': parent.name,
                        'unit_id': current_unit_id,
                        'unit_name': current_unit_name,
                        'parent_id': parent.parent_id
                    })
                    topic_key_map[parent_key] = True
                    if parent.parent_id and parent.parent_id in parent_topic_objects:
                        grand_parent = parent_topic_objects[parent.parent_id]
                        grand_parent_key = f"{current_unit_id}-{grand_parent.id}"
                        if grand_parent_key not in topic_key_map:
                            final_topic_list.append({
                                'topic_id': grand_parent.id,
                                'topic_name': grand_parent.name,
                                'unit_id': current_unit_id,
                                'unit_name': current_unit_name,
                                'parent_id': grand_parent.parent_id
                            })
                            topic_key_map[grand_parent_key] = True
    
    assessment_topics_for_dropdown = sorted(
        final_topic_list, 
        key=lambda x: (
            x['unit_id'] or 999,
            x['parent_id'] or 0,
            x['topic_id']
        )
    )
    assessment_topics_json = json.dumps(assessment_topics_for_dropdown)
    
    # --- 7. [FIXED] ดึงข้อมูลการประเมินเชิงคุณภาพ (Qualitative Assessment) จาก "หน่วยแม่" ---
    templates = AssessmentTemplate.query.options(
        selectinload(AssessmentTemplate.topics).options(
            selectinload(AssessmentTopic.children).selectinload(AssessmentTopic.children)
        ),
        selectinload(AssessmentTemplate.rubric_levels)
    ).order_by(AssessmentTemplate.display_order).all()
    
    existing_qual_scores_list = QualitativeScore.query.filter(
        QualitativeScore.course_id == course.id
    ).all()
    existing_qual_scores = {
        f"{score.student_id}-{score.assessment_topic_id}": score.score_value
        for score in existing_qual_scores_list
    }

    # [FIX 2] ดึงข้อมูลจาก "หน่วยแม่" (Parent Unit)
    valid_topic_ids = set()
    if parent_unit:
        # Load all assessment items from the PARENT unit
        parent_with_assessment_items = LearningUnit.query.options(
            selectinload(LearningUnit.assessment_items).selectinload(AssessmentItem.topic)
        ).get(parent_unit.id)
        
        if parent_with_assessment_items:
            valid_topic_ids = {item.topic.id for item in parent_with_assessment_items.assessment_items if item.topic}

    # 7.3 สร้าง JSON ก้อนใหญ่ (แบบกรองแล้ว)
    qualitative_assessment_data = {
        'templates': [], 
        'existing_scores': existing_qual_scores
    }

    for tpl in templates:
        if not tpl.rubric_levels:
            continue
            
        template_dict = {
            'id': tpl.id,
            'name': tpl.name,
            'topics': [], 
            'rubrics': sorted(
                [{'label': r.label, 'value': r.value} for r in tpl.rubric_levels],
                key=lambda x: x['value'], reverse=True
            )
        }
        
        def build_filtered_tree(topic):
            filtered_children = []
            if topic.children:
                for child in topic.children:
                    filtered_child = build_filtered_tree(child)
                    if filtered_child: 
                        filtered_children.append(filtered_child)
            
            # [FIXED] เงื่อนไข: Topic นี้ ถูกเลือกใน "หน่วยแม่" หรือไม่
            if topic.id in valid_topic_ids or filtered_children:
                return {
                    'id': topic.id,
                    'name': topic.name,
                    'children': filtered_children
                }
            return None

        top_level_topics = [t for t in tpl.topics if t.parent_id is None]
        top_level_topics.sort(key=lambda x: x.id) 
        
        for topic in top_level_topics:
            filtered_topic_tree = build_filtered_tree(topic)
            if filtered_topic_tree: 
                template_dict['topics'].append(filtered_topic_tree)
                
        if template_dict['topics']:
            qualitative_assessment_data['templates'].append(template_dict)

    qualitative_assessment_data_json = json.dumps(qualitative_assessment_data)

    # --- [THE FIX] ---
    # 8. สร้าง Form เปล่าเพื่อส่ง CSRF Token
    form = FlaskForm()
    # --- [END FIX] ---

    # --- 8. Render Template (ส่งข้อมูลใหม่ทั้งหมด) ---
    return render_template('teacher/mobile_entry.html',
                           entry=entry,
                           course=course,
                           classroom=classroom,
                           enrollments=enrollments,
                           attendance_records=attendance_records,
                           graded_items=graded_items, 
                           scores=scores,
                           date_iso=date_iso,
                           attendance_date=attendance_date,
                           
                           # --- ข้อมูลที่อัปเดต ---
                           assessment_topics_json=assessment_topics_json,
                           qualitative_assessment_data_json=qualitative_assessment_data_json, 
                           suggested_unit_id=suggested_unit_id,
                           student_group_map_json=student_group_map_json,
                           lesson_plan=lesson_plan, 
                           classroom_id_js=classroom.id,
                           form=form # 👈 [THE FIX] ส่ง form ไปยัง template
                           )

# --- [ V28.2 START: เพิ่ม API 1 ] ---
@bp.route('/api/attendance/set-status', methods=['POST'])
@login_required
def api_set_attendance_status():
    """
    API Endpoint for saving a single attendance record from the mobile UI (click-to-cycle).
    """
    data = request.get_json()
    if not data:
        return jsonify({'status': 'error', 'message': 'Invalid data'}), 400

    try:
        student_id = int(data.get('student_id'))
        entry_id = int(data.get('entry_id'))
        status = str(data.get('status'))
        date_iso = data.get('date', date.today().isoformat())
        attendance_date = date.fromisoformat(date_iso)
        
        # --- ตรวจสอบสิทธิ์ (สำคัญมาก) ---
        entry = db.session.get(TimetableEntry, entry_id)
        if not entry:
            return jsonify({'status': 'error', 'message': 'Timetable entry not found'}), 404
        if current_user not in entry.course.teachers:
            return jsonify({'status': 'error', 'message': 'Forbidden'}), 403
        
        # --- ค้นหาหรือสร้าง Record ---
        record = AttendanceRecord.query.filter_by(
            student_id=student_id,
            timetable_entry_id=entry_id,
            attendance_date=attendance_date
        ).first()

        current_time = datetime.utcnow()

        if record:
            # Update existing record
            if record.status != status:
                record.status = status
                record.recorded_at = current_time
                record.recorder_id = current_user.id
                db.session.add(record)
        else:
            # Create new record
            record = AttendanceRecord(
                student_id=student_id,
                timetable_entry_id=entry_id,
                status=status,
                recorder_id=current_user.id,
                attendance_date=attendance_date,
                recorded_at=current_time
            )
            db.session.add(record)

        # --- ตรวจสอบการแจ้งเตือน (เรียก Service) ---
        if status in ['ABSENT', 'LATE', 'TARDY']:
             # เราต้องส่ง object ของ record ไป (ต้อง commit ก่อนเพื่อให้มี id หรือ flush)
             # หรือส่ง object ที่เพิ่งสร้าง/อัปเดตไป
             db.session.flush() # Flush เพื่อให้ record มี state
             check_and_create_attendance_warnings(record)

        db.session.commit()
        return jsonify({'status': 'success', 'message': 'Attendance updated'})

    except Exception as e:
        db.session.rollback()
        current_app.logger.error(f"Error in api_set_attendance_status: {e}", exc_info=True)
        return jsonify({'status': 'error', 'message': str(e)}), 500
# --- [ V28.2 END: API 1 ] ---

# --- [ V28.2 START: เพิ่ม API 2 ] ---
@bp.route('/api/qualitative/set-score', methods=['POST'])
@login_required
def api_set_qualitative_score():
    """
    API Endpoint for saving a single qualitative score from the mobile UI.
    """
    data = request.get_json()
    if not data:
        return jsonify({'status': 'error', 'message': 'Invalid data'}), 400

    try:
        student_id = int(data.get('student_id'))
        course_id = int(data.get('course_id'))
        topic_id = int(data.get('topic_id'))
        score_value = data.get('score_value') # รับค่ามา (อาจเป็น null หรือ int)

        # --- ตรวจสอบสิทธิ์ (สำคัญมาก) ---
        course = db.session.get(Course, course_id)
        if not course:
            return jsonify({'status': 'error', 'message': 'Course not found'}), 404
        if current_user not in course.teachers:
            return jsonify({'status': 'error', 'message': 'Forbidden'}), 403
        
        # --- ค้นหาหรือสร้าง Record ---
        score_entry = QualitativeScore.query.filter_by(
            student_id=student_id,
            course_id=course_id,
            assessment_topic_id=topic_id
        ).first()

        if score_value is None:
            # ถ้าค่ที่ส่งมาเป็น null/None ให้ลบ record ที่มีอยู่ (ถ้ามี)
            if score_entry:
                db.session.delete(score_entry)
        else:
            # ถ้ามีค่าส่งมา (เป็นตัวเลข)
            score_value = int(score_value)
            if score_entry:
                # Update existing score
                score_entry.score_value = score_value
            else:
                # Create new score
                score_entry = QualitativeScore(
                    student_id=student_id,
                    course_id=course_id,
                    assessment_topic_id=topic_id,
                    score_value=score_value
                )
                db.session.add(score_entry)
        
        db.session.commit()
        return jsonify({'status': 'success', 'message': 'Qualitative score updated'})

    except Exception as e:
        db.session.rollback()
        current_app.logger.error(f"Error in api_set_qualitative_score: {e}", exc_info=True)
        return jsonify({'status': 'error', 'message': str(e)}), 500
# --- [ V28.2 END: API 2 ] ---
    
#
# API: Create Student Group
#
@bp.route('/api/student-groups', methods=['POST'])
@login_required
# @teacher_required
def create_student_group():
    """
    API endpoint to create a new student group.
    """
    data = request.get_json()
    name = data.get('name')
    course_id = data.get('course_id')
    lesson_plan_id = data.get('lesson_plan_id') # Optional

    if not name or not course_id:
        return jsonify({'success': False, 'error': 'Missing name or course_id'}), 400

    # Validate that the teacher owns this course
    course = Course.query.get_or_404(course_id)
    
    if getattr(course, "teacher_id", None) and course.teacher_id != current_user.id:
        return jsonify({'success': False, 'error': 'Forbidden'}), 403

    try:
        new_group = StudentGroup(
            name=name,
            course_id=course_id,
            lesson_plan_id=lesson_plan_id
        )
        db.session.add(new_group)
        db.session.commit()
        
        return jsonify({
            'success': True, 
            'id': new_group.id, 
            'name': new_group.name
        }), 201

    except Exception as e:
        db.session.rollback()
        return jsonify({'success': False, 'error': str(e)}), 500

@bp.route('/api/attendance/copy-between-entries', methods=['POST'])
@login_required
def copy_attendance_between_entries():
    """ API to copy attendance records from a source entry to a target entry on the same date. """
    data = request.get_json()
    source_entry_id = data.get('source_entry_id')
    target_entry_id = data.get('target_entry_id')
    date_str = data.get('date')

    if not all([source_entry_id, target_entry_id, date_str]):
        return jsonify({'status': 'error', 'message': 'ข้อมูลไม่ครบถ้วน'}), 400

    try:
        attendance_date = date.fromisoformat(date_str)
    except ValueError:
        return jsonify({'status': 'error', 'message': 'รูปแบบวันที่ไม่ถูกต้อง'}), 400

    # Security check: Ensure user teaches both courses (implicitly checked via source/target fetch)
    source_entry = TimetableEntry.query.get_or_404(source_entry_id)
    target_entry = TimetableEntry.query.get_or_404(target_entry_id)

    # Basic check: Are they the same course and teacher?
    if source_entry.course_id != target_entry.course_id:
         return jsonify({'status': 'error', 'message': 'ไม่สามารถคัดลอกระหว่างรายวิชาต่างกันได้'}), 400
    if current_user not in source_entry.course.teachers: # Check one is enough
         abort(403)

    try:
        # 1. Fetch source records for the given date
        source_records = AttendanceRecord.query.filter_by(
            timetable_entry_id=source_entry_id,
            attendance_date=attendance_date
        ).all()

        if not source_records:
            return jsonify({'status': 'warning', 'message': 'ไม่พบข้อมูลการเข้าเรียนในคาบต้นทาง'}), 200 # Not an error, just nothing to copy

        # 2. Fetch target records (if any) to perform upsert
        target_records_map = {
            rec.student_id: rec for rec in AttendanceRecord.query.filter_by(
                timetable_entry_id=target_entry_id,
                attendance_date=attendance_date
            ).all()
        }

        current_time = datetime.utcnow()
        records_processed = 0

        # 3. Loop through source records and create/update target records
        for s_rec in source_records:
            t_rec = target_records_map.get(s_rec.student_id)
            if t_rec: # Update existing target record
                if t_rec.status != s_rec.status: # Only update if status is different
                    t_rec.status = s_rec.status
                    t_rec.recorded_at = current_time # Update timestamp
                    t_rec.recorder_id = current_user.id # Update recorder
                    records_processed += 1
            else: # Insert new target record
                new_rec = AttendanceRecord(
                    student_id=s_rec.student_id,
                    timetable_entry_id=target_entry_id,
                    status=s_rec.status,
                    recorder_id=current_user.id,
                    attendance_date=attendance_date,
                    recorded_at=current_time # Set initial timestamp
                )
                db.session.add(new_rec)
                records_processed += 1

        if records_processed > 0:
            db.session.commit()
            print(f"--- Copied/Updated {records_processed} attendance records ---")
            return jsonify({'status': 'success', 'message': f'คัดลอกข้อมูล {records_processed} รายการเรียบร้อย'})
        else:
             print("--- No attendance changes needed during copy ---")
             return jsonify({'status': 'info', 'message': 'ข้อมูลคาบถัดไปตรงกับคาบนี้อยู่แล้ว (ไม่มีอะไรเปลี่ยนแปลง)'})


    except Exception as e:
        db.session.rollback()
        print(f"--- ERROR during attendance copy: {e} ---")
        current_app.logger.error(f"Error copying attendance: {e}")
        return jsonify({'status': 'error', 'message': str(e)}), 500
    
# --- [NEW] API: The "Factory" for creating Google Forms ---
@bp.route('/api/graded-item/<int:item_id>/create-google-form', methods=['POST'])
@login_required
def create_google_form_for_item(item_id):
    item = GradedItem.query.get_or_404(item_id)
    course = item.learning_unit.lesson_plan.courses[0] # Get associated course

    # 1. Permission Check
    if current_user not in course.teachers:
        abort(403)
    if item.google_form_id:
        return jsonify({'status': 'error', 'message': 'This item is already linked to a Google Form.'}), 400

    # 2. Get and Refresh Credentials
    creds = _get_google_creds(current_user)
    if not creds:
        return jsonify({'status': 'error', 'message': 'ไม่สามารถยืนยันสิทธิ์ Google ได้ กรุณา Logout และ Login ด้วย Google ใหม่อีกครั้ง', 'reauth': True}), 401

    try:
        # 3. Build API Services
        forms_service = googleapiclient.discovery.build('forms', 'v1', credentials=creds, cache_discovery=False)
        sheets_service = googleapiclient.discovery.build('sheets', 'v4', credentials=creds, cache_discovery=False)
        script_service = googleapiclient.discovery.build('script', 'v1', credentials=creds, cache_discovery=False)

        # 4. Create the Google Form
        form_title = f"{course.subject.name} - {item.name}"
        form_body = {'info': {'title': form_title, 'documentTitle': form_title}}
        form_result = forms_service.forms().create(body=form_body).execute()
        form_id = form_result['formId']
        form_url = form_result['responderUri']

        # 5. Add "Student ID" question and set as Quiz
        update_request_body = {
            'requests': [
                {'createItem': {'item': {'title': 'รหัสนักเรียน (Student ID)',
                                        'questionItem': {'question': {'required': True, 'textQuestion': {}}}},
                                'location': {'index': 0}}},
                {'updateFormInfo': {'info': {'description': f'สำหรับ {course.subject.name} ห้อง {course.classroom.name}\nรายการคะแนน: {item.name} ({item.max_score} คะแนน)'},
                                    'updateMask': 'description'}},
                {'updateSettings': {'settings': {'quizSettings': {'isQuiz': True}},
                                    'updateMask': 'quizSettings.isQuiz'}}
            ]
        }
        forms_service.forms().batchUpdate(formId=form_id, body=update_request_body).execute()

        # 6. Create the Google Sheet to store responses
        sheet_body = {'properties': {'title': f'[RESPONSES] {form_title}'}}
        sheet_result = sheets_service.spreadsheets().create(body=sheet_body).execute()
        sheet_id = sheet_result['spreadsheetId']
        sheet_url = sheet_result['spreadsheetUrl']

        # 7. Link Form to Sheet
        link_request_body = {'requests': [{'createWatch': {'watch': {'target': {'spreadsheetId': sheet_id}, 'eventType': 'RESPONSES'}}}]}
        # หมายเหตุ: การ Link Sheet นี้ใช้ API คนละตัวกับการตั้งค่า "Create spreadsheet" ในหน้า Form UI
        # เราจะใช้ Apps Script trigger แทนการ link โดยตรง
        # forms_service.forms().watches().create(formId=form_id, body=link_request_body).execute()
        # --- [NEW] Set form destination to our new sheet ---
        # link_sheet_body = { "requests": [ { "updateFormInfo": { "info": { "linkedSheetId": sheet_id }, "updateMask": "linkedSheetId" } } ] }
        # forms_service.forms().batchUpdate(formId=form_id, body=link_sheet_body).execute()


        # 8. Create and Install the Apps Script "Trigger" (The Spy)
        # 8.1 Create a secure token for our webhook
        webhook_token = jwt.encode(
            {'type': 'graded_item', 'item_id': item.id, 'course_id': course.id, 'user_id': current_user.id}, # 👈 [FIX]
            current_app.config['SECRET_KEY'],
            algorithm='HS256'
        )
        webhook_url = url_for('teacher.google_form_submit_score', _external=True)

        # 8.2 Define the Apps Script code
        script_content = f"""
function onFormSubmit(e) {{
  var formResponse = e.response;
  var itemResponses = formResponse.getItemResponses();
  var studentId = '';
  
  // 1. [FIX] Find the Student ID by question title (MUCH safer)
  for (var i = 0; i < itemResponses.length; i++) {{
    // Check if title *includes* 'รหัสนักเรียน'
    if (itemResponses[i].getQuestion().includes('รหัสนักเรียน')) {{ 
      studentId = itemResponses[i].getResponse();
      break;
    }}
  }}
  // [END FIX]

  // 2. Get the score
  var score = formResponse.getGradableResponse().getScore();
  
  // 3. Prepare data to send to our app
  var payload = {{
    'token': '{webhook_token}',
    'student_id': studentId,
    'score': score,
    'max_score': {item.max_score},
    'form_id': '{form_id}',
    'response_id': formResponse.getId()
  }};
  
  var options = {{
    'method' : 'post',
    'contentType': 'application/json',
    'payload' : JSON.stringify(payload)
  }};
  
  // 4. Send the data to our app's webhook
  try {{
    UrlFetchApp.fetch('{webhook_url}', options);
  }} catch (error) {{
    Logger.log('Error sending score to EdHub: ' + error.toString());
  }}
}}
"""
        script_manifest = {
            "timeZone": "Asia/Bangkok",
            "exceptionLogging": "STACKDRIVER",
            "runtimeVersion": "V8"
        }
        
        # 8.3 Create the script project
        script_request_body = {
            'title': f'EdHub Trigger for {form_title}',
            'parentId': sheet_id # [FIX] ผูก Script กับ Sheet
        }
        script_project = script_service.projects().create(body=script_request_body).execute()
        script_id = script_project['scriptId']

        # 8.4 Update the script content
        script_files = [
            {'name': 'Code', 'type': 'SERVER_JS', 'source': script_content},
            {'name': 'appsscript', 'type': 'JSON', 'source': json.dumps(script_manifest)}
        ]
        script_service.projects().updateContent(
            scriptId=script_id, 
            body={'files': script_files}
        ).execute()

        # 8.5 Install the 'onFormSubmit' trigger
        trigger_request_body = {
            'function': 'onFormSubmit',
            'event': {'type': 'ON_FORM_SUBMIT'}
        }
        # [FIX] ใช้ script_id ที่เพิ่งสร้าง และตั้งค่า trigger ให้กับ Sheet
        script_service.projects().triggers().create(
            scriptId=script_id,
            body=trigger_request_body
        ).execute()


        # 9. Save IDs to our Database
        item.google_form_id = form_id
        item.google_sheet_id = sheet_id
        db.session.commit()

        return jsonify({
            'status': 'success',
            'message': 'สร้าง Google Form และติดตั้ง Trigger สำเร็จ!',
            'form_url': form_url,
            'sheet_url': sheet_url
        })

    except Exception as e:
        db.session.rollback()
        current_app.logger.error(f"Error creating Google Form for item {item_id}: {e}", exc_info=True)
        # พยายามลบ Form ที่อาจสร้างค้างไว้
        try:
            if 'form_id' in locals():
                forms_service.forms().delete(formId=form_id).execute()
            if 'sheet_id' in locals():
                sheets_service.spreadsheets().delete(spreadsheetId=sheet_id).execute() # (ระวัง: Sheet API ไม่มี delete)
        except:
            pass
        return jsonify({'status': 'error', 'message': f'เกิดข้อผิดพลาด: {e}'}), 500

# --- [NEW] API: The "Factory" for creating Google Forms for EXAMS ---
@bp.route('/api/course/<int:course_id>/create-google-form/<string:exam_type>', methods=['POST'])
@login_required
def create_google_form_for_exam(course_id, exam_type):
    if exam_type not in ['midterm', 'final']:
        abort(400, 'Invalid exam type')

    course = Course.query.get_or_404(course_id)
    
    # 1. Permission Check
    if current_user not in course.teachers:
        abort(403)
    
    exam_name = "สอบกลางภาค" if exam_type == 'midterm' else "สอบปลายภาค"
    max_score = 0
    if course.lesson_plan:
         all_units = course.lesson_plan.learning_units
         if exam_type == 'midterm':
             max_score = sum(u.midterm_score or 0 for u in all_units)
         else:
             max_score = sum(u.final_score or 0 for u in all_units)

    # 2. Check if already exists
    if (exam_type == 'midterm' and course.midterm_google_form_id) or \
       (exam_type == 'final' and course.final_google_form_id):
        return jsonify({'status': 'error', 'message': f'This course is already linked to a Google Form for {exam_type}.'}), 400

    # 3. Get and Refresh Credentials
    creds = _get_google_creds(current_user)
    if not creds:
        return jsonify({'status': 'error', 'message': 'ไม่สามารถยืนยันสิทธิ์ Google ได้ กรุณา Logout และ Login ด้วย Google ใหม่อีกครั้ง', 'reauth': True}), 401

    try:
        # 4. Build API Services
        forms_service = googleapiclient.discovery.build('forms', 'v1', credentials=creds, cache_discovery=False)
        sheets_service = googleapiclient.discovery.build('sheets', 'v4', credentials=creds, cache_discovery=False)
        script_service = googleapiclient.discovery.build('script', 'v1', credentials=creds, cache_discovery=False)

        # 5. Create the Google Form
        form_title = f"{course.subject.name} - {exam_name}"
        form_body = {'info': {'title': form_title, 'documentTitle': form_title}}
        form_result = forms_service.forms().create(body=form_body).execute()
        form_id = form_result['formId']
        form_url = form_result['responderUri']

        # 6. Add "Student ID" question and set as Quiz
        update_request_body = {
            'requests': [
                {'createItem': {'item': {'title': 'รหัสนักเรียน (Student ID)',
                                        'questionItem': {'question': {'required': True, 'textQuestion': {}}}},
                                'location': {'index': 0}}},
                {'updateFormInfo': {'info': {'description': f'สำหรับ {course.subject.name} ห้อง {course.classroom.name}\nการสอบ: {exam_name} ({max_score} คะแนน)'},
                                    'updateMask': 'description'}},
                {'updateSettings': {'settings': {'quizSettings': {'isQuiz': True}},
                                    'updateMask': 'quizSettings.isQuiz'}}
            ]
        }
        forms_service.forms().batchUpdate(formId=form_id, body=update_request_body).execute()

        # 7. Create the Google Sheet
        sheet_body = {'properties': {'title': f'[RESPONSES] {form_title}'}}
        sheet_result = sheets_service.spreadsheets().create(body=sheet_body).execute()
        sheet_id = sheet_result['spreadsheetId']
        sheet_url = sheet_result['spreadsheetUrl']

        # 8. Link Form to Sheet
        # link_sheet_body = { "requests": [ { "updateFormInfo": { "info": { "linkedSheetId": sheet_id }, "updateMask": "linkedSheetId" } } ] }
        # forms_service.forms().batchUpdate(formId=form_id, body=link_sheet_body).execute()

        # 9. Create and Install the Apps Script "Trigger"
        # 9.1 Create a secure token (THIS IS THE KEY DIFFERENCE)
        webhook_token = jwt.encode(
            {'type': exam_type, 'course_id': course.id, 'user_id': current_user.id}, # 👈 [FIX] Payload for Exams
            current_app.config['SECRET_KEY'],
            algorithm='HS256'
        )
        webhook_url = url_for('teacher.google_form_submit_score', _external=True)

        # 9.2 Define the (robust) Apps Script code
        script_content = f"""
function onFormSubmit(e) {{
  var formResponse = e.response;
  var itemResponses = formResponse.getItemResponses();
  var studentId = '';
  
  // [FIX] Find the Student ID by question title
  for (var i = 0; i < itemResponses.length; i++) {{
    if (itemResponses[i].getQuestion().includes('รหัสนักเรียน')) {{
      studentId = itemResponses[i].getResponse();
      break;
    }}
  }}

  var score = formResponse.getGradableResponse().getScore();
  
  var payload = {{
    'token': '{webhook_token}',
    'student_id': studentId,
    'score': score,
    'max_score': {max_score},
    'form_id': '{form_id}',
    'response_id': formResponse.getId()
  }};
  
  var options = {{ 'method' : 'post', 'contentType': 'application/json', 'payload' : JSON.stringify(payload) }};
  
  try {{
    UrlFetchApp.fetch('{webhook_url}', options);
  }} catch (error) {{
    Logger.log('Error sending score to EdHub: ' + error.toString());
  }}
}}
"""
        script_manifest = {"timeZone": "Asia/Bangkok", "exceptionLogging": "STACKDRIVER", "runtimeVersion": "V8"}
        
        # 9.3 Create the script project
        script_project = script_service.projects().create(
            body={'title': f'EdHub Trigger for {form_title}', 'parentId': sheet_id}
        ).execute()
        script_id = script_project['scriptId']

        # 9.4 Update the script content
        script_files = [
            {'name': 'Code', 'type': 'SERVER_JS', 'source': script_content},
            {'name': 'appsscript', 'type': 'JSON', 'source': json.dumps(script_manifest)}
        ]
        script_service.projects().updateContent(scriptId=script_id, body={'files': script_files}).execute()

        # 9.5 Install the 'onFormSubmit' trigger
        trigger_request_body = {'function': 'onFormSubmit', 'event': {'type': 'ON_FORM_SUBMIT'}}
        script_service.projects().triggers().create(scriptId=script_id, body=trigger_request_body).execute()

        # 10. Save IDs to *Course* Model
        if exam_type == 'midterm':
            course.midterm_google_form_id = form_id
            course.midterm_google_sheet_id = sheet_id
        else: # final
            course.final_google_form_id = form_id
            course.final_google_sheet_id = sheet_id
            
        db.session.commit()

        return jsonify({
            'status': 'success', 'message': 'สร้าง Google Form และติดตั้ง Trigger สำเร็จ!',
            'form_url': form_url, 'sheet_url': sheet_url
        })

    except Exception as e:
        db.session.rollback()
        current_app.logger.error(f"Error creating Google Form for exam {exam_type} (Course {course_id}): {e}", exc_info=True)
        return jsonify({'status': 'error', 'message': f'เกิดข้อผิดพลาด: {e}'}), 500
    
# --- [REVISED] API: The "Webhook" to receive ALL scores from Google ---
@bp.route('/api/google-form/submit-score', methods=['POST'])
def google_form_submit_score():
    data = request.get_json()
    token = data.get('token')
    
    if not token:
        current_app.logger.warning("Google Webhook received NO TOKEN.")
        abort(401)

    try:
        # 1. Decode the token to identify the item AND TYPE
        payload = jwt.decode(token, current_app.config['SECRET_KEY'], algorithms=['HS256'])
        score_type = payload.get('type') # 👈 [NEW] 'graded_item', 'midterm', or 'final'
        course_id = payload['course_id']
        
        # 2. Get data from Google
        student_id_str = data.get('student_id')
        score_value = data.get('score')
        score_float = float(score_value) if score_value is not None else None

        if not student_id_str or score_value is None:
            current_app.logger.warning(f"Webhook (Type: {score_type}): Missing student_id or score.")
            return jsonify({'status': 'error', 'message': 'Missing student_id or score'}), 400

        # 3. Find the student in our DB
        student = Student.query.filter_by(student_id=student_id_str).first()
        if not student:
            current_app.logger.error(f"Webhook (Type: {score_type}): Student ID '{student_id_str}' NOT FOUND.")
            return jsonify({'status': 'error', 'message': 'Student ID not found'}), 404

        # 4. --- [NEW] Smart Upsert Logic ---
        if score_type == 'graded_item':
            item_id = payload['item_id']
            score_obj = Score.query.filter_by(student_id=student.id, graded_item_id=item_id).first()
            if score_obj:
                score_obj.score = score_float
            else:
                score_obj = Score(student_id=student.id, graded_item_id=item_id, score=score_float)
                db.session.add(score_obj)
            log_message = f"Webhook SUCCESS: Saved GradedItem score {score_float} for Student {student.id} on Item {item_id}."

        elif score_type in ['midterm', 'final']:
            grade_obj = CourseGrade.query.filter_by(student_id=student.id, course_id=course_id).first()
            if not grade_obj:
                grade_obj = CourseGrade(student_id=student.id, course_id=course_id)
                db.session.add(grade_obj)
            
            if score_type == 'midterm':
                grade_obj.midterm_score = score_float
                log_message = f"Webhook SUCCESS: Saved Midterm score {score_float} for Student {student.id} on Course {course_id}."
            else: # final
                grade_obj.final_score = score_float
                log_message = f"Webhook SUCCESS: Saved Final score {score_float} for Student {student.id} on Course {course_id}."
        
        else:
            current_app.logger.error(f"Webhook: Unknown score type '{score_type}' in token.")
            return jsonify({'status': 'error', 'message': 'Unknown score type'}), 400
        
        db.session.commit()
        current_app.logger.info(log_message)
        return jsonify({'status': 'success'})

    except jwt.ExpiredSignatureError:
        current_app.logger.error("Google Webhook: Expired token.")
        abort(401)
    except jwt.InvalidTokenError:
        current_app.logger.error("Google Webhook: Invalid token.")
        abort(401)
    except Exception as e:
        db.session.rollback()
        current_app.logger.error(f"Error in google_form_submit_score: {e}", exc_info=True)
        return jsonify({'status': 'error', 'message': str(e)}), 500